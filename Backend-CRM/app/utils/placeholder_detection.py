"""
Placeholder Detection Utility

Automatically detects placeholders (e.g., {{PI_NAME}}) in template content
and extracts unique placeholder names.
"""

from typing import Any, Dict, List, Set
import io
import re

# Liberal placeholder pattern:
#   - Allows any case: {{site_name}}, {{Site_Name}}, {{SITE_NAME}}
#   - Allows whitespace padding inside braces: {{ SITE_NAME }}
#   - Allows spaces/hyphens inside the name: {{Site Name}}, {{site-name}}
# Names are normalized to UPPER_SNAKE_CASE so downstream config keys are consistent.
_PLACEHOLDER_PATTERN = re.compile(r'\{\{\s*([A-Za-z0-9_][A-Za-z0-9_\- ]*?)\s*\}\}')


def _normalize_placeholder_name(raw: str) -> str:
    """Normalize a placeholder token to UPPER_SNAKE_CASE."""
    cleaned = re.sub(r'[\s\-]+', '_', raw.strip())
    cleaned = re.sub(r'_+', '_', cleaned).strip('_')
    return cleaned.upper()


def detect_placeholders_in_text(text: str) -> Set[str]:
    """
    Detect all placeholders in a text string.

    Args:
        text: The text string to scan

    Returns:
        Set of placeholder names (without {{}}, normalized to UPPER_SNAKE_CASE)
    """
    if not text or not isinstance(text, str):
        return set()

    matches = _PLACEHOLDER_PATTERN.findall(text)
    return {_normalize_placeholder_name(m) for m in matches if m.strip()}


def detect_placeholders_in_docx_bytes(content: bytes) -> Set[str]:
    """
    Detect placeholders inside a DOCX file's body (paragraphs + tables).

    Args:
        content: Raw bytes of the .docx file

    Returns:
        Set of normalized placeholder names found anywhere in the document.
    """
    if not content:
        return set()

    try:
        from docx import Document
    except ImportError:
        return set()

    try:
        doc = Document(io.BytesIO(content))
    except Exception:
        return set()

    text_blob_parts: List[str] = []
    for paragraph in doc.paragraphs:
        text_blob_parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_blob_parts.append(cell.text)

    return detect_placeholders_in_text("\n".join(text_blob_parts))


def detect_placeholders_in_json(content: Any) -> Set[str]:
    """
    Recursively detect all placeholders in a JSON structure (TipTap document format).
    
    Args:
        content: The JSON content (dict, list, or primitive)
        
    Returns:
        Set of unique placeholder names (without {{}})
    """
    placeholders = set()
    
    if isinstance(content, dict):
        for key, value in content.items():
            if key == "text" and isinstance(value, str):
                # This is a text node in TipTap format - detect placeholders
                placeholders.update(detect_placeholders_in_text(value))
            elif isinstance(value, (dict, list)):
                # Recursively process nested structures
                placeholders.update(detect_placeholders_in_json(value))
            elif isinstance(value, str):
                # Detect placeholders in any string value
                placeholders.update(detect_placeholders_in_text(value))
    elif isinstance(content, list):
        for item in content:
            placeholders.update(detect_placeholders_in_json(item))
    elif isinstance(content, str):
        placeholders.update(detect_placeholders_in_text(content))
    
    return placeholders


def create_default_placeholder_config(placeholders: Set[str]) -> Dict[str, Dict[str, bool]]:
    """
    Create default placeholder configuration with all placeholders editable.
    
    Args:
        placeholders: Set of placeholder names
        
    Returns:
        Configuration dict: {"PLACEHOLDER_NAME": {"editable": true}}
    """
    config = {}
    for placeholder in placeholders:
        config[placeholder] = {"editable": True}
    return config


def detect_and_create_config(template_content: Dict[str, Any]) -> Dict[str, Dict[str, bool]]:
    """
    Detect placeholders in template content and create default configuration.
    
    Args:
        template_content: The template JSON content
        
    Returns:
        Configuration dict with all detected placeholders set to editable=True
    """
    placeholders = detect_placeholders_in_json(template_content)
    return create_default_placeholder_config(placeholders)

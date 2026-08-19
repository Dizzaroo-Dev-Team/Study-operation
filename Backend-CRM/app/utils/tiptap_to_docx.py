"""Convert TipTap JSON content into a .docx file using python-docx.

Used when an agreement is created from a CLAUSE-COMPOSED template (or any
template whose content was edited in the Clause Builder). The builder stores
the full document as TipTap JSON in StudyTemplate.template_content; this module
renders that JSON to a real .docx so the existing placeholder-replacement and
signing pipeline can consume it unchanged.

Supported node types: doc, heading, paragraph, bulletList, orderedList,
listItem, hardBreak, text (with bold/italic/underline marks). Tables are
rendered as plain paragraphs row-by-row (best effort) — the clause builder
rarely emits tables, and a faithful table port is out of scope here.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)


def _apply_marks(run, marks: Optional[List[Dict[str, Any]]]) -> None:
    """Apply TipTap marks (bold/italic/underline) to a python-docx run."""
    if not marks:
        return
    for mark in marks:
        mtype = mark.get("type")
        if mtype == "bold":
            run.bold = True
        elif mtype == "italic":
            run.italic = True
        elif mtype == "underline":
            run.underline = True


def _render_inline(paragraph, nodes: Optional[List[Dict[str, Any]]]) -> None:
    """Render an array of inline nodes (text / hardBreak) into a paragraph."""
    if not nodes:
        return
    for node in nodes:
        ntype = node.get("type")
        if ntype == "text":
            run = paragraph.add_run(node.get("text", ""))
            _apply_marks(run, node.get("marks"))
        elif ntype == "hardBreak":
            paragraph.add_run().add_break()


def _render_list(doc: Document, node: Dict[str, Any], ordered: bool) -> None:
    """Render bulletList / orderedList by emitting list-styled paragraphs."""
    style = "List Number" if ordered else "List Bullet"
    for item in node.get("content", []):
        if item.get("type") != "listItem":
            continue
        # A listItem contains block nodes (usually paragraphs)
        for block in item.get("content", []):
            if block.get("type") == "paragraph":
                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:
                    # Style not present in the default template — fall back
                    p = doc.add_paragraph()
                _render_inline(p, block.get("content"))


def _render_block(doc: Document, node: Dict[str, Any]) -> None:
    """Render a single top-level block node into the document."""
    ntype = node.get("type")

    if ntype == "heading":
        level = int(node.get("attrs", {}).get("level", 2) or 2)
        level = max(1, min(level, 9))
        p = doc.add_heading(level=level)
        _render_inline(p, node.get("content"))

    elif ntype == "paragraph":
        p = doc.add_paragraph()
        _render_inline(p, node.get("content"))

    elif ntype == "bulletList":
        _render_list(doc, node, ordered=False)

    elif ntype == "orderedList":
        _render_list(doc, node, ordered=True)

    elif ntype == "blockquote":
        for block in node.get("content", []):
            p = doc.add_paragraph()
            p.style = doc.styles["Quote"] if "Quote" in [s.name for s in doc.styles] else p.style
            _render_inline(p, block.get("content"))

    elif ntype == "table":
        # Best-effort: flatten each row's cell text onto a tab-separated line.
        for row in node.get("content", []):
            cells_text: List[str] = []
            for cell in row.get("content", []):
                parts: List[str] = []
                for block in cell.get("content", []):
                    for inline in block.get("content", []) or []:
                        if inline.get("type") == "text":
                            parts.append(inline.get("text", ""))
                cells_text.append(" ".join(parts))
            doc.add_paragraph("\t".join(cells_text))

    elif ntype == "horizontalRule":
        doc.add_paragraph("_" * 40)

    else:
        # Unknown node — try to recurse into its content so nothing is silently lost
        for child in node.get("content", []) or []:
            _render_block(doc, child)


def tiptap_json_to_docx(json_content: Dict[str, Any], output_path: Path) -> Path:
    """Render a TipTap JSON document to a .docx file at output_path.

    Returns the output_path. Raises ValueError if the content is empty.
    """
    if not json_content or not isinstance(json_content, dict):
        raise ValueError("tiptap_json_to_docx: empty or invalid content")

    content = json_content.get("content")
    if not content:
        raise ValueError("tiptap_json_to_docx: document has no content")

    doc = Document()

    # Reasonable default body font
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
    except Exception:
        pass

    for node in content:
        try:
            _render_block(doc, node)
        except Exception as exc:  # noqa: BLE001 — never let one node abort the doc
            logger.warning("tiptap_json_to_docx: failed to render node %s: %s", node.get("type"), exc)

    doc.save(str(output_path))
    return output_path


def has_renderable_content(json_content: Optional[Dict[str, Any]]) -> bool:
    """True if the TipTap JSON has at least one non-empty block."""
    if not json_content or not isinstance(json_content, dict):
        return False
    content = json_content.get("content")
    if not content or not isinstance(content, list):
        return False
    # Treat a single empty paragraph as "no content"
    if len(content) == 1:
        only = content[0]
        if only.get("type") == "paragraph" and not only.get("content"):
            return False
    return True

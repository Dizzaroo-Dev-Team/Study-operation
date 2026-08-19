"""
DOCX Placeholder Replacement Utility

Replaces placeholders in DOCX files using python-docx library.
Preserves formatting, tables, images, and layout.
Supports dynamic field mappings from template configuration.

Also provides helper utilities for enforcing locked (non-editable)
placeholder values across document versions.
"""

import io
import logging
from pathlib import Path
from typing import Optional, Dict, Any, Tuple

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.models import SiteProfile, Agreement

logger = logging.getLogger(__name__)

# Legacy mapping for backward compatibility (used for signature blocks)
LEGACY_PLACEHOLDER_MAP = {
    "SITE_NAME": "site_name",
    "HOSPITAL_NAME": "hospital_name",
    "PI_NAME": "pi_name",
    "PI_EMAIL": "pi_email",
    "PI_PHONE": "pi_phone",
    "PRIMARY_CONTRACTING_ENTITY": "primary_contracting_entity",
    "AUTHORIZED_SIGNATORY_NAME": "authorized_signatory_name",
    "AUTHORIZED_SIGNATORY_EMAIL": "authorized_signatory_email",
    "AUTHORIZED_SIGNATORY_TITLE": "authorized_signatory_title",
    "ADDRESS_LINE_1": "address_line_1",
    "CITY": "city",
    "STATE": "state",
    "COUNTRY": "country",
    "POSTAL_CODE": "postal_code",
}


def _sanitize_replacement_value(text: str) -> str:
    """Strip ``\\n`` and ``\\r`` from a replacement value so they never become
    ``<w:br/>`` or ``<w:cr/>`` elements inside the generated document."""
    return text.replace('\r', '').replace('\n', ' ').strip()


def _normalize_placeholder_name(name: str) -> str:
    """Normalize placeholder name for resilient field_mappings lookup."""
    text = str(name or "").replace("\u00a0", " ").strip()
    text = " ".join(text.split())
    return text.upper()


def get_placeholder_value_from_mapping(
    placeholder_name: str,
    field_mappings: Optional[Dict[str, str]],
    site_profile: Optional[SiteProfile] = None,
    agreement: Optional[Agreement] = None,
) -> Optional[str]:
    """
    Get the value for a placeholder using dynamic field mappings.
    
    Args:
        placeholder_name: The placeholder name (e.g., "SITE_NAME")
        field_mappings: Dict mapping placeholder names to data source paths (e.g., {"SITE_NAME": "site_profile.site_name"})
        site_profile: The SiteProfile instance
        agreement: The Agreement instance
        
    Returns:
        The value from the data source as a string, or None if:
        - No field_mappings dict is provided
        - No mapping exists for this placeholder
        - The data source object is missing
        - The resolved field value is null
        
        When no explicit mapping exists the placeholder is left unchanged.
    """
    mapping_path: Optional[str] = None

    # Only resolve values when explicit field_mappings are provided.
    # Without a mapping, placeholders are preserved unchanged in the document.
    if not field_mappings:
        logger.debug(
            "No field_mappings provided; placeholder '%s' will be preserved unchanged.",
            placeholder_name,
        )
        return None

    # Case-insensitive lookup: stored mapping keys may be lower/mixed case (AI
    # suggestions keep the document's case; manual adds upper-case them), so match
    # on an UPPER-cased view rather than the raw dict.
    fm_upper = {str(k).upper(): v for k, v in field_mappings.items()}
    normalized_placeholder = _normalize_placeholder_name(placeholder_name)
    mapping_path = fm_upper.get(normalized_placeholder)
    if not mapping_path and " " in normalized_placeholder:
        mapping_path = fm_upper.get(normalized_placeholder.replace(" ", "_"))
    if not mapping_path and "_" in normalized_placeholder:
        mapping_path = fm_upper.get(normalized_placeholder.replace("_", " "))
    if not mapping_path:
        logger.debug(
            "No explicit field_mappings entry for placeholder '%s'; "
            "placeholder will be preserved unchanged.",
            placeholder_name,
        )
        return None
    
    # Parse the mapping path (format: "data_source.field_name")
    parts = mapping_path.split(".", 1)
    if len(parts) != 2:
        logger.warning(
            "Invalid mapping path format for placeholder %s: %s. Expected format: 'data_source.field_name'",
            placeholder_name,
            mapping_path,
        )
        return None
    
    data_source, field_name = parts[0].lower(), parts[1]
    
    # Get value from appropriate data source
    value = None
    if data_source == "site_profile":
        if not site_profile:
            logger.warning("SiteProfile not provided for placeholder: %s", placeholder_name)
            return None
        value = getattr(site_profile, field_name, None)
    elif data_source == "agreement":
        if not agreement:
            logger.warning("Agreement not provided for placeholder: %s", placeholder_name)
            return None
        value = getattr(agreement, field_name, None)
    else:
        logger.warning(
            "Unknown data source '%s' for placeholder: %s (mapping_path=%s)",
            data_source,
            placeholder_name,
            mapping_path,
        )
        return None
    
    if value is None:
        logger.debug(
            "Field %s from %s is null for placeholder: %s (placeholder left unchanged)",
            field_name,
            data_source,
            placeholder_name,
        )
        return None

    sanitized = _sanitize_replacement_value(str(value))
    # Empty / whitespace-only mapped values must behave like "no data": keep
    # {{PLACEHOLDER}} visible for reviewers (otherwise content controls show blank).
    if not sanitized:
        logger.debug(
            "Field %s from %s is empty for placeholder: %s (placeholder left unchanged)",
            field_name,
            data_source,
            placeholder_name,
        )
        return None

    return sanitized


def get_placeholder_value(placeholder_name: str, profile: SiteProfile) -> str:
    """
    Legacy function for backward compatibility.
    Get the value for a placeholder from SiteProfile using legacy mapping.
    
    Args:
        placeholder_name: The placeholder name (e.g., "PI_NAME")
        profile: The SiteProfile instance
        
    Returns:
        The value from SiteProfile, or empty string if not found
    """
    field_name = LEGACY_PLACEHOLDER_MAP.get(placeholder_name.upper())
    if not field_name:
        return ""
    
    value = getattr(profile, field_name, None)
    return _sanitize_replacement_value(str(value)) if value is not None else ""


def replace_placeholders_in_docx(
    docx_path: Path,
    profile: SiteProfile,
    output_path: Path,
    sponsor_signatory_name: Optional[str] = None,
    sponsor_signatory_email: Optional[str] = None,
    current_user_email: Optional[str] = None,
    field_mappings: Optional[Dict[str, str]] = None,
    agreement: Optional[Agreement] = None,
    template_id: Optional[str] = None,
    placeholder_config: Optional[Dict[str, Dict[str, bool]]] = None,
    resolved_values: Optional[Dict[str, str]] = None,
) -> Path:
    """
    Replace placeholders in a DOCX file and save to output path.
    Uses dynamic field_mappings if provided, otherwise falls back to legacy mapping.

    Editability rules (driven purely by placeholder_config):
      - If editable == True  → insert an EDITABLE content control (w:sdt) for the value.
      - If editable == False → insert the value as plain static text (no content control).

    When a template has a placeholder_config, the generated DOCX is placed into
    Word "forms" protection mode so that ONLY the editable content controls can
    be modified in ONLYOFFICE / Word. All other text is read-only.

    Args:
        docx_path: Path to source DOCX file
        profile: SiteProfile instance
        output_path: Path to save the modified DOCX
        sponsor_signatory_name: Optional sponsor signatory name
        sponsor_signatory_email: Optional sponsor signatory email
        current_user_email: Optional current user email for fallback
        field_mappings: Optional dict mapping placeholder names to data source paths
        agreement: Optional Agreement instance for agreement-based placeholders
        template_id: Optional template ID for logging
        placeholder_config: Optional dict mapping placeholder names to config:
            {"PLACEHOLDER": {"editable": true/false}}

    Returns:
        Path to the output DOCX file
    """
    try:
        logger.info(f"Starting placeholder replacement for template_id: {template_id}, output: {output_path}")
        
        doc = Document(docx_path)

        # Normalise runs: split any <w:r> that has multiple content children
        # (e.g. <w:t>…</w:t><w:br/><w:t>…</w:t>) into one-element-per-run.
        # This guarantees that <w:br/> elements live in their own runs and
        # cannot be lost when a neighbouring <w:t> is edited.
        _normalize_document_runs(doc)

        # Extract all placeholders from the document once; replacement behavior is
        # decided per placeholder from configurable_editable, independent of mapping.
        placeholders_found = _extract_placeholders_from_docx(doc)
        logger.info(f"Found {len(placeholders_found)} unique placeholders in document")
        if field_mappings:
            logger.info(f"Using dynamic field_mappings: {field_mappings}")
        else:
            logger.info("No field_mappings provided; editable placeholders will still be editable.")

        replacements_made = 0
        replacements_skipped = 0

        for placeholder_name in placeholders_found:
            # Skip signature blocks - handled separately
            if placeholder_name in ["SITE_SIGNATURE_BLOCK", "SPONSOR_SIGNATURE_BLOCK"]:
                continue

            # Per-signer signature placeholders ({{SIGNATURE_<ROLE>}}, e.g.
            # SIGNATURE_DIRECTOR / SIGNATURE_PI / SIGNATURE_VP) are PRESERVED verbatim:
            # they are not values to fill at create-time but precise anchors the signing
            # stamper finds (and covers) when each signer signs. Leaving them as literal
            # text — not wrapped in an editable content control — keeps them stable
            # through document protection and DOCX->PDF conversion.
            if placeholder_name.upper().startswith("SIGNATURE_"):
                logger.info("Preserving per-signer signature placeholder %s for signing-time stamping", placeholder_name)
                continue

            # Prefer pre-resolved values (built in the route via
            # cta.service.resolve_field_mapping_values — handles study/irb sources,
            # case-insensitive keys, AND the shared DEFAULT_TOKEN_MAP so well-known
            # placeholders auto-fill even without an explicit template mapping). Fall
            # back to the in-line resolver (site_profile/agreement only) when not given.
            value = None
            if resolved_values:
                value = resolved_values.get(_normalize_placeholder_name(placeholder_name))
            if value is None:
                value = get_placeholder_value_from_mapping(
                    placeholder_name,
                    field_mappings,
                    site_profile=profile,
                    agreement=agreement,
                )

            placeholder_text = f"{{{{{placeholder_name}}}}}"
            placeholder_tokens = _placeholder_token_variants(placeholder_name)

            # Editability is driven only by template config.
            is_editable = True  # Default to editable if no config
            if placeholder_config:
                placeholder_cfg = placeholder_config.get(placeholder_name, {})
                is_editable = placeholder_cfg.get("editable", True)

            if is_editable:
                # Editable fields are always wrapped in editable content controls.
                # If mapping/value is missing or empty, keep the placeholder token so users
                # can still edit it in the generated document.
                control_value = (
                    placeholder_text
                    if value is None or not str(value).strip()
                    else value
                )
                logger.info(
                    "Replacing editable placeholder %s with content control (value=%s)",
                    placeholder_name,
                    "mapped" if value is not None else "placeholder",
                )
                for token in placeholder_tokens:
                    _replace_with_editable_content_control(
                        doc, token, control_value, placeholder_name
                    )
                replacements_made += 1
                logger.debug(
                    "Replaced editable %s with value: %s",
                    placeholder_name,
                    control_value,
                )
                continue

            # Non-editable fields: only replace when a mapped non-empty value exists.
            if value is not None and str(value).strip():
                logger.info(f"Replacing locked placeholder {placeholder_name} as plain text (non-editable)")
                for paragraph in _iter_all_paragraphs(doc):
                    for token in placeholder_tokens:
                        if token in paragraph.text:
                            if _replace_placeholder_in_paragraph_preserving_runs(
                                paragraph, token, value
                            ):
                                replacements_made += 1
                logger.debug(f"Replaced locked {placeholder_name} with value: {value}")
            else:
                replacements_skipped += 1
                logger.warning(
                    "Locked placeholder %s left unchanged (no value/mapping)",
                    placeholder_name,
                )

        logger.info(
            "Placeholder replacement complete: %s replacements made, %s skipped",
            replacements_made,
            replacements_skipped,
        )
        
        # Handle signature block placeholders (always use legacy logic)
        _replace_signature_blocks_in_docx(
            doc,
            profile,
            sponsor_signatory_name,
            sponsor_signatory_email,
            current_user_email
        )
        
        # Apply document protection always:
        # - Entire document is put into "forms" protection mode.
        # - ONLY content controls (editable placeholders) remain editable.
        # Legacy templates with no placeholder_config still get protected as
        # defense-in-depth so signers cannot mutate non-placeholder content.
        editable_field_names = [
            name for name, cfg in (placeholder_config or {}).items()
            if cfg.get("editable", True)
        ]
        logger.info(
            "Applying document protection for template_id=%s. "
            "Editable placeholders: %s",
            template_id,
            ", ".join(editable_field_names) if editable_field_names else "none",
        )
        try:
            _apply_restricted_editing_protection(doc)
            logger.info("Document protection applied successfully")
        except Exception as e:
            logger.error(
                "Failed to apply document protection; document will remain fully editable. "
                "Error: %s",
                e,
                exc_info=True,
            )
        
        # Save the modified document
        doc.save(output_path)
        logger.info(f"Successfully replaced placeholders in DOCX: {output_path}")
        
        return output_path
        
    except Exception as e:
        logger.error(f"Failed to replace placeholders in DOCX: {str(e)}", exc_info=True)
        raise Exception(f"Failed to process DOCX file: {str(e)}")


def _iter_all_paragraphs(doc: Document):
    """Yield every paragraph in the document: body, tables, headers, footers,
    and text boxes (w:txbxContent).  This ensures placeholders placed anywhere
    in the template are found and replaced."""
    from docx.text.paragraph import Paragraph as _Para

    # Body paragraphs
    for paragraph in doc.paragraphs:
        yield paragraph

    # Tables in body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    # Headers and footers
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph

    # Text boxes (w:txbxContent) – scan whole doc XML so header/footer text boxes are included
    for txbx in doc.element.iter(qn('w:txbxContent')):
        for p_elem in txbx.findall(qn('w:p')):
            yield _Para(p_elem, txbx)


def _extract_placeholders_from_docx(doc: Document) -> set:
    """
    Extract all placeholder names from a DOCX document.
    Searches body paragraphs, tables, headers, footers, and text boxes.
    
    Args:
        doc: python-docx Document object
        
    Returns:
        Set of placeholder names (without {{ }})
    """
    import re
    placeholders = set()
    # Allow mixed-case, spaces, and punctuation inside placeholders:
    # {{PLACEHOLDER}} or {{ Place Holder }} or {{Protocol #}}
    pattern = r'\{\{\s*([^{}]+?)\s*\}\}'
    
    for paragraph in _iter_all_paragraphs(doc):
        matches = re.findall(pattern, paragraph.text)
        placeholders.update(matches)
    
    return placeholders


def _placeholder_token_variants(placeholder_name: str) -> list:
    """Return common token variants for a placeholder name."""
    name = (placeholder_name or "").strip()
    if not name:
        return []
    return [
        f"{{{{{name}}}}}",
        f"{{{{ {name} }}}}",
        f"{{{{ {name}}}}}",
        f"{{{{{name} }}}}",
    ]


def _normalize_paragraph_runs(paragraph) -> None:
    """Split any ``<w:r>`` that contains more than one content child element
    (``<w:t>``, ``<w:br/>``, ``<w:tab/>``, ``<w:cr/>``, etc.) into separate
    atomic runs, each holding **exactly one** content element.

    After normalisation every ``<w:br/>`` lives in its own dedicated ``<w:r>``
    and cannot be displaced or lost when a neighbouring ``<w:t>`` is edited
    during placeholder replacement.

    Formatting (``<w:rPr>``) is deep-copied into each new run so the visual
    appearance is unchanged.
    """
    from copy import deepcopy

    for run in list(paragraph.runs):
        r_elem = run._element
        r_pr = r_elem.find(qn('w:rPr'))

        content_children = [
            child for child in r_elem if child.tag != qn('w:rPr')
        ]

        if len(content_children) <= 1:
            continue

        parent = r_elem.getparent()
        idx = list(parent).index(r_elem)

        for child in content_children:
            new_r = OxmlElement('w:r')
            if r_pr is not None:
                new_r.append(deepcopy(r_pr))
            new_r.append(child)          # lxml *moves* the child
            parent.insert(idx, new_r)
            idx += 1

        parent.remove(r_elem)


def _normalize_document_runs(doc: Document) -> None:
    """Normalise every paragraph in the document (body + tables)."""
    for paragraph in doc.paragraphs:
        _normalize_paragraph_runs(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _normalize_paragraph_runs(paragraph)


def _get_run_t_text(r_elem) -> str:
    """Return the concatenated text of all ``<w:t>`` children in a ``<w:r>``."""
    return ''.join((t.text or '') for t in r_elem.findall(qn('w:t')))


def _set_run_t_text(r_elem, text: str) -> None:
    """Set the text of a ``<w:r>`` by editing its *first* ``<w:t>`` child.

    Unlike the python-docx ``Run.text`` setter this does **not** remove
    ``<w:br/>``, ``<w:tab/>`` or any other non-text child elements, so the
    original paragraph structure (line breaks, tabs, etc.) is preserved.

    Assumes runs have been normalised (one ``<w:t>`` per run).
    """
    t_elems = r_elem.findall(qn('w:t'))
    if t_elems:
        t_elems[0].text = text
        t_elems[0].set(qn('xml:space'), 'preserve')
    else:
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r_elem.append(t)


def _create_text_run_from(source_r_elem, text: str):
    """Build a **new** ``<w:r>`` element that copies *only* the formatting
    (``<w:rPr>``) from *source_r_elem* plus a single ``<w:t>`` with *text*.

    ``<w:br/>``, ``<w:tab/>`` and every other content element in the source
    are intentionally **not** copied so they cannot leak extra line breaks
    into the output.
    """
    from copy import deepcopy

    new_r = OxmlElement('w:r')
    r_pr = source_r_elem.find(qn('w:rPr'))
    if r_pr is not None:
        new_r.append(deepcopy(r_pr))
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_r.append(t)
    return new_r


def _replace_placeholder_in_paragraph_preserving_runs(
    paragraph, placeholder_text: str, replacement_text: str
) -> bool:
    """Replace *placeholder_text* with *replacement_text* inside a paragraph
    by editing individual ``<w:t>`` XML elements within runs.

    Key guarantees:

    * ``<w:br/>``, ``<w:tab/>``, bookmarks, formatting and every other
      non-text XML element in the paragraph are **untouched**.
    * Handles placeholders that Word has split across several ``<w:r>``
      elements.
    * Replacement values are sanitised (``\\n``, ``\\r`` stripped) so
      database values with stray line-endings cannot inject ``<w:br/>``.

    Returns ``True`` when a replacement was made.
    """
    clean_value = _sanitize_replacement_value(replacement_text)

    runs = list(paragraph.runs)
    if not runs:
        return False

    run_spans = []
    pos = 0
    for run in runs:
        t_text = _get_run_t_text(run._element)
        run_spans.append((run, t_text, pos, pos + len(t_text)))
        pos += len(t_text)

    full_text = "".join(info[1] for info in run_spans)
    ph_start = full_text.find(placeholder_text)
    if ph_start == -1:
        return False
    ph_end = ph_start + len(placeholder_text)

    affected = [
        (run, t_text, rs, re_)
        for run, t_text, rs, re_ in run_spans
        if re_ > ph_start and rs < ph_end
    ]
    if not affected:
        return False

    if len(affected) == 1:
        run, t_text, rs, _ = affected[0]
        ls = ph_start - rs
        le = ph_end - rs
        new_text = t_text[:ls] + clean_value + t_text[le:]
        _set_run_t_text(run._element, new_text)
        return True

    first_run, first_text, first_rs, _ = affected[0]
    last_run, last_text, last_rs, _ = affected[-1]

    before = first_text[: ph_start - first_rs]
    after = last_text[ph_end - last_rs :]
    _set_run_t_text(first_run._element, before + clean_value + after)

    last_run._element.getparent().remove(last_run._element)
    for run, _, _, _ in affected[1:-1]:
        run._element.getparent().remove(run._element)

    return True


def _apply_restricted_editing_protection(doc: Document) -> None:
    """
    Apply Word document protection to restrict editing to content controls only.
    
    This makes the entire document read-only except for editable content controls.
    Users can only edit text inside content controls; all other text is protected.
    
    Args:
        doc: Document object to protect
    """
    try:
        # Get the document's settings part
        # Document protection is stored in settings.xml
        # python-docx stores settings in part.related_parts
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        
        # Get settings part from related parts
        settings_part = None
        for rel in doc.part.rels.values():
            if rel.reltype == RT.SETTINGS:
                settings_part = rel.target_part
                break
        
        if settings_part is None:
            logger.warning("Settings part not found - document protection cannot be applied. Document will work but editing may not be restricted.")
            return
        
        settings_xml = settings_part.element
        
        # Check if documentProtection already exists
        doc_protection = settings_xml.find(qn('w:documentProtection'))
        
        if doc_protection is None:
            # Create new documentProtection element
            doc_protection = OxmlElement('w:documentProtection')
            settings_xml.append(doc_protection)
        
        # Set protection attributes:
        # - w:enforcement="1" - Enable protection
        # - w:edit="forms" - Allow editing only in form fields (content controls)
        # - w:cryptProviderType="rsaFull" and w:cryptAlgorithmClass="hash" - Encryption settings
        # - w:cryptAlgorithmType="typeAny" - Algorithm type
        # - w:cryptAlgorithmSid="4" - Algorithm SID
        # - w:cryptSpinCount="100000" - Spin count for encryption
        # - w:hash="..." - Hash value (can be empty for forms protection)
        
        doc_protection.set(qn('w:enforcement'), '1')
        doc_protection.set(qn('w:edit'), 'forms')  # allow editing only in content controls
        doc_protection.set(qn('w:cryptProviderType'), 'rsaFull')
        doc_protection.set(qn('w:cryptAlgorithmClass'), 'hash')
        doc_protection.set(qn('w:cryptAlgorithmType'), 'typeAny')
        doc_protection.set(qn('w:cryptAlgorithmSid'), '4')
        doc_protection.set(qn('w:cryptSpinCount'), '100000')
        
        # For forms protection, we can use a simple hash or leave it empty
        # Word will accept empty hash for forms protection
        doc_protection.set(qn('w:hash'), '')
        
        logger.info("Applied document protection: restricted editing mode (editing allowed only in content controls)")
        
    except Exception as e:
        logger.error(f"Failed to apply document protection: {str(e)}", exc_info=True)
        # Don't fail the entire operation if protection fails
        # Document will still work, just without protection


def remove_document_protection(docx_path: Path, output_path: Path) -> None:
    """
    Remove document protection from a DOCX file.
    
    This function removes the "forms" protection that restricts editing to content controls only.
    After removing protection, the entire document becomes fully editable.
    
    This is used when creating review copies for external reviewers, so they can edit
    the entire document without restrictions (including seeing the full menu bar in OnlyOffice).
    
    Args:
        docx_path: Path to the protected DOCX file
        output_path: Path to save the unprotected DOCX file
    """
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        
        # Load the document
        doc = Document(docx_path)
        
        # Get settings part from related parts
        settings_part = None
        for rel in doc.part.rels.values():
            if rel.reltype == RT.SETTINGS:
                settings_part = rel.target_part
                break
        
        if settings_part is None:
            logger.warning("Settings part not found - document may not have protection to remove.")
            # Still save the document even if no settings part exists
            doc.save(output_path)
            return
        
        settings_xml = settings_part.element
        
        # Find and remove documentProtection element
        doc_protection = settings_xml.find(qn('w:documentProtection'))
        
        if doc_protection is not None:
            # Remove the protection element
            settings_xml.remove(doc_protection)
            logger.info(f"Removed document protection from {docx_path}")
        else:
            logger.debug(f"No document protection found in {docx_path}")
        
        # Save the unprotected document
        doc.save(output_path)
        logger.info(f"Saved unprotected document to {output_path}")
        
    except Exception as e:
        logger.error(f"Failed to remove document protection: {str(e)}", exc_info=True)
        # If removal fails, just copy the file as-is (better than failing completely)
        import shutil
        shutil.copy2(docx_path, output_path)
        logger.warning(f"Document protection removal failed, copied file as-is: {output_path}")


def _replace_with_editable_content_control(
    doc: Document,
    placeholder_text: str,
    value: str,
    tag_name: str
) -> None:
    """
    Replace placeholder text with an editable content control (w:sdt element).
    
    The content control will be editable - used for fields marked as editable.
    Document will be in restricted editing mode, allowing editing only in content controls.
    
    Args:
        doc: Document object
        placeholder_text: The placeholder text to replace (e.g., "{{SITE_NAME}}")
        value: The value to insert into the content control
        tag_name: The tag name for the content control (usually the placeholder name)
    """
    try:
        for paragraph in _iter_all_paragraphs(doc):
            if placeholder_text in paragraph.text:
                _replace_placeholder_in_paragraph_with_editable_content_control(
                    paragraph, placeholder_text, value, tag_name
                )
        
        logger.debug(f"Replaced {placeholder_text} with editable content control (tag: {tag_name})")
    except Exception as e:
        logger.error(f"Failed to create editable content control for {placeholder_text}: {str(e)}", exc_info=True)
        for paragraph in _iter_all_paragraphs(doc):
            if placeholder_text in paragraph.text:
                _replace_placeholder_in_paragraph_preserving_runs(
                    paragraph, placeholder_text, value
                )


def _replace_placeholder_in_paragraph_with_editable_content_control(
    paragraph,
    placeholder_text: str,
    value: str,
    tag_name: str,
) -> None:
    """Replace placeholder text in a paragraph with an editable content
    control (``w:sdt``) WITHOUT locking attributes.

    Handles placeholders that span multiple ``<w:r>`` elements.  Before/after
    text runs are built from scratch (only ``<w:rPr>`` + ``<w:t>``) so that
    ``<w:br/>``, ``<w:tab/>`` and other content elements from the original
    runs are never accidentally duplicated.
    """
    from copy import deepcopy

    try:
        if placeholder_text not in paragraph.text:
            return

        clean_value = _sanitize_replacement_value(value)

        p_elem = paragraph._element
        runs = list(paragraph.runs)
        if not runs:
            return

        # -- Build position map from <w:t> text only ----------------------
        run_spans = []
        pos = 0
        for run in runs:
            t_text = _get_run_t_text(run._element)
            run_spans.append((run, t_text, pos, pos + len(t_text)))
            pos += len(t_text)

        full_text = "".join(info[1] for info in run_spans)
        ph_start = full_text.find(placeholder_text)
        if ph_start == -1:
            return
        ph_end = ph_start + len(placeholder_text)

        affected = [
            (run, t_text, rs, re_)
            for run, t_text, rs, re_ in run_spans
            if re_ > ph_start and rs < ph_end
        ]
        if not affected:
            return

        first_run = affected[0][0]
        first_r_elem = first_run._element
        first_rs = affected[0][2]
        first_text = affected[0][1]
        last_run = affected[-1][0]
        last_rs = affected[-1][2]
        last_text = affected[-1][1]

        before_text = first_text[: ph_start - first_rs]
        after_text = last_text[ph_end - last_rs :]

        # -- Build the content control (w:sdt) ----------------------------
        sdt = OxmlElement('w:sdt')

        sdt_pr = OxmlElement('w:sdtPr')
        tag_el = OxmlElement('w:tag')
        tag_el.set(qn('w:val'), tag_name)
        sdt_pr.append(tag_el)
        sdt_pr.append(OxmlElement('w:richText'))

        sdt_content = OxmlElement('w:sdtContent')
        new_run = OxmlElement('w:r')

        r_pr = first_r_elem.find(qn('w:rPr'))
        if r_pr is not None:
            new_run.append(deepcopy(r_pr))

        t_elem = OxmlElement('w:t')
        t_elem.set(qn('xml:space'), 'preserve')
        t_elem.text = clean_value
        new_run.append(t_elem)

        sdt_content.append(new_run)
        sdt.append(sdt_pr)
        sdt.append(sdt_content)

        # -- Splice into the paragraph XML --------------------------------
        insert_idx = list(p_elem).index(first_r_elem)

        if before_text:
            p_elem.insert(
                insert_idx,
                _create_text_run_from(first_r_elem, before_text),
            )
            insert_idx += 1

        p_elem.insert(insert_idx, sdt)
        insert_idx += 1

        if after_text:
            p_elem.insert(
                insert_idx,
                _create_text_run_from(last_run._element, after_text),
            )

        # Remove original affected runs
        for run, _, _, _ in affected:
            r_elem = run._element
            if r_elem.getparent() is not None:
                r_elem.getparent().remove(r_elem)

        logger.info(
            "Created editable content control for %s with value: %s",
            tag_name, clean_value,
        )

    except Exception as e:
        logger.error(
            "Error creating editable content control in paragraph: %s",
            str(e), exc_info=True,
        )
        _replace_placeholder_in_paragraph_preserving_runs(
            paragraph, placeholder_text, value
        )


def enforce_locked_placeholders_in_docx(
    original_docx_path: Path,
    updated_docx_bytes: bytes,
    placeholder_config: Optional[Dict[str, Dict[str, bool]]],
    field_mappings: Optional[Dict[str, str]],
    site_profile: Optional[SiteProfile],
    agreement: Optional[Agreement],
    template_id: Optional[str] = None,
) -> bytes:
    """
    Enforce locked (non-editable) placeholders by comparing the original
    populated Version 1 document with the updated version and restoring
    locked values where they were modified.

    This works at a paragraph / table-cell granularity:
    - If a paragraph or cell in the original document contains a locked
      placeholder value and the corresponding paragraph/cell text was
      changed in the updated document, the original text is restored.
    - Editable fields are never touched.
    - Any failure in this logic will NOT block saving – the original
      updated_docx_bytes are returned.

    Args:
        original_docx_path: Path to the original Version 1 DOCX file.
        updated_docx_bytes: Bytes of the updated DOCX from ONLYOFFICE.
        placeholder_config: Template placeholder configuration
            {"PLACEHOLDER": {"editable": true/false}}.
        field_mappings: Dynamic field mappings
            {"PLACEHOLDER": "site_profile.site_name"}.
        site_profile: SiteProfile instance.
        agreement: Agreement instance.
        template_id: Optional template ID for logging.

    Returns:
        Potentially modified DOCX bytes with locked fields restored.
    """
    try:
        if not placeholder_config:
            logger.debug(
                "No placeholder_config provided for locked field enforcement; skipping."
            )
            return updated_docx_bytes

        # Determine locked placeholders (editable == False)
        locked_placeholders = [
            name
            for name, cfg in placeholder_config.items()
            if not cfg.get("editable", True)
        ]

        if not locked_placeholders:
            logger.debug(
                "No locked placeholders (editable=false) configured; skipping enforcement."
            )
            return updated_docx_bytes

        # Build canonical values for each locked placeholder
        locked_values: Dict[str, str] = {}
        for placeholder_name in locked_placeholders:
            value = get_placeholder_value_from_mapping(
                placeholder_name,
                field_mappings,
                site_profile=site_profile,
                agreement=agreement,
            )
            if value:
                locked_values[placeholder_name] = value
            else:
                logger.debug(
                    "Locked placeholder '%s' has no resolved value; it will be ignored.",
                    placeholder_name,
                )

        if not locked_values:
            logger.info(
                "Locked placeholders configured but none have resolvable values; "
                "skipping enforcement. template_id=%s",
                template_id,
            )
            return updated_docx_bytes

        logger.info(
            "Enforcing locked placeholders for template_id=%s. Locked fields: %s",
            template_id,
            ", ".join(
                f"{name}='{val}'" for name, val in locked_values.items()
            ),
        )

        # Load original and updated documents
        if not original_docx_path.exists():
            logger.warning(
                "Original Version 1 DOCX not found at %s; skipping locked field enforcement.",
                original_docx_path,
            )
            return updated_docx_bytes

        original_doc = Document(original_docx_path)
        updated_doc = Document(io.BytesIO(updated_docx_bytes))

        def _enforce_in_paragraphs(
            original_paragraphs,
            updated_paragraphs,
            scope: str,
        ) -> Tuple[int, int]:
            """Return (blocks_checked, blocks_reverted)."""
            checked = 0
            reverted = 0

            for idx, (orig_p, new_p) in enumerate(
                zip(original_paragraphs, updated_paragraphs)
            ):
                orig_text = orig_p.text or ""
                new_text = new_p.text or ""

                if not orig_text:
                    continue

                checked += 1

                if orig_text == new_text:
                    continue

                # Check if this block contains any locked value in the original
                for placeholder_name, locked_value in locked_values.items():
                    if locked_value and locked_value in orig_text:
                        if locked_value not in new_text or new_text != orig_text:
                            logger.warning(
                                "Locked field changed and will be restored. "
                                "template_id=%s placeholder=%s scope=%s index=%s",
                                template_id,
                                placeholder_name,
                                scope,
                                idx,
                            )
                            new_p.text = orig_text
                            reverted += 1
                            break

            return checked, reverted

        total_checked = 0
        total_reverted = 0

        # Enforce in top-level paragraphs
        c, r = _enforce_in_paragraphs(
            original_doc.paragraphs, updated_doc.paragraphs, scope="paragraph"
        )
        total_checked += c
        total_reverted += r

        # Enforce inside tables (cell by cell)
        for t_idx, (orig_table, new_table) in enumerate(
            zip(original_doc.tables, updated_doc.tables)
        ):
            for r_idx, (orig_row, new_row) in enumerate(
                zip(orig_table.rows, new_table.rows)
            ):
                for c_idx, (orig_cell, new_cell) in enumerate(
                    zip(orig_row.cells, new_row.cells)
                ):
                    scope = f"table[{t_idx}].row[{r_idx}].cell[{c_idx}]"
                    c, r = _enforce_in_paragraphs(
                        orig_cell.paragraphs,
                        new_cell.paragraphs,
                        scope=scope,
                    )
                    total_checked += c
                    total_reverted += r

        logger.info(
            "Locked placeholder enforcement complete. template_id=%s blocks_checked=%s blocks_reverted=%s",
            template_id,
            total_checked,
            total_reverted,
        )

        # Save updated document back to bytes
        output_buffer = io.BytesIO()
        updated_doc.save(output_buffer)
        return output_buffer.getvalue()

    except Exception as e:
        logger.error(
            "Error enforcing locked placeholders (template_id=%s): %s. "
            "Proceeding without reverting locked fields.",
            template_id,
            str(e),
            exc_info=True,
        )
        # In case of any error, do NOT block save – return original bytes
        return updated_docx_bytes


def _replace_multiline_placeholder_in_paragraph(
    paragraph, placeholder_text: str, replacement_text: str
) -> bool:
    """Replace a placeholder with multi-line text, converting ``\\n`` into
    ``<w:br/>`` elements so each line renders on its own row within the same
    paragraph.

    Used for signature blocks where line breaks are intentional and must NOT
    be stripped by ``_sanitize_replacement_value``.

    Returns ``True`` when a replacement was made.
    """
    from copy import deepcopy

    runs = list(paragraph.runs)
    if not runs:
        return False

    run_spans = []
    pos = 0
    for run in runs:
        t_text = _get_run_t_text(run._element)
        run_spans.append((run, t_text, pos, pos + len(t_text)))
        pos += len(t_text)

    full_text = "".join(info[1] for info in run_spans)
    ph_start = full_text.find(placeholder_text)
    if ph_start == -1:
        return False
    ph_end = ph_start + len(placeholder_text)

    affected = [
        (run, t_text, rs, re_)
        for run, t_text, rs, re_ in run_spans
        if re_ > ph_start and rs < ph_end
    ]
    if not affected:
        return False

    first_run = affected[0][0]
    first_r_elem = first_run._element
    first_text = affected[0][1]
    first_rs = affected[0][2]
    last_run = affected[-1][0]
    last_text = affected[-1][1]
    last_rs = affected[-1][2]

    before_text = first_text[: ph_start - first_rs]
    after_text = last_text[ph_end - last_rs :]

    r_pr = first_r_elem.find(qn('w:rPr'))
    p_elem = paragraph._element
    insert_idx = list(p_elem).index(first_r_elem)

    if before_text:
        p_elem.insert(insert_idx, _create_text_run_from(first_r_elem, before_text))
        insert_idx += 1

    clean = replacement_text.replace('\r', '')
    lines = clean.split('\n')
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped:
            p_elem.insert(insert_idx, _create_text_run_from(first_r_elem, stripped))
            insert_idx += 1
        if i < len(lines) - 1:
            br_run = OxmlElement('w:r')
            if r_pr is not None:
                br_run.append(deepcopy(r_pr))
            br_run.append(OxmlElement('w:br'))
            p_elem.insert(insert_idx, br_run)
            insert_idx += 1

    if after_text:
        src = last_run._element if len(affected) > 1 else first_r_elem
        p_elem.insert(insert_idx, _create_text_run_from(src, after_text))

    for run, _, _, _ in affected:
        r_elem = run._element
        if r_elem.getparent() is not None:
            r_elem.getparent().remove(r_elem)

    return True


def _replace_signature_blocks_in_docx(
    doc: Document,
    profile: SiteProfile,
    sponsor_signatory_name: Optional[str] = None,
    sponsor_signatory_email: Optional[str] = None,
    current_user_email: Optional[str] = None
):
    """
    Replace signature block placeholders in DOCX document.

    Signature and Date fields are placed via Zoho Sign API (no text tags in document).
    Recipient1 = Site, Recipient2 = Sponsor.
    
    Args:
        doc: python-docx Document object
        profile: SiteProfile instance
        sponsor_signatory_name: Optional sponsor signatory name
        sponsor_signatory_email: Optional sponsor signatory email
        current_user_email: Optional current user email for fallback
    """
    # Find and replace SITE_SIGNATURE_BLOCK
    site_sig_text = _generate_site_signature_block_text(profile)
    
    # Find and replace SPONSOR_SIGNATURE_BLOCK
    sponsor_sig_text = _generate_sponsor_signature_block_text(
        profile,
        sponsor_signatory_name,
        sponsor_signatory_email,
        current_user_email
    )
    
    for paragraph in doc.paragraphs:
        if "{{SITE_SIGNATURE_BLOCK}}" in paragraph.text:
            _replace_multiline_placeholder_in_paragraph(
                paragraph, "{{SITE_SIGNATURE_BLOCK}}", site_sig_text
            )
        if "{{SPONSOR_SIGNATURE_BLOCK}}" in paragraph.text:
            _replace_multiline_placeholder_in_paragraph(
                paragraph, "{{SPONSOR_SIGNATURE_BLOCK}}", sponsor_sig_text
            )
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "{{SITE_SIGNATURE_BLOCK}}" in paragraph.text:
                        _replace_multiline_placeholder_in_paragraph(
                            paragraph, "{{SITE_SIGNATURE_BLOCK}}", site_sig_text
                        )
                    if "{{SPONSOR_SIGNATURE_BLOCK}}" in paragraph.text:
                        _replace_multiline_placeholder_in_paragraph(
                            paragraph, "{{SPONSOR_SIGNATURE_BLOCK}}", sponsor_sig_text
                        )


def _generate_site_signature_block_text(profile: SiteProfile) -> str:
    """Generate site signature block text. Signature/Date fields are placed via Zoho Sign API (no tags)."""
    hospital_name = getattr(profile, "hospital_name", None) or ""
    signatory_name = getattr(profile, "authorized_signatory_name", None) or ""
    signatory_title = getattr(profile, "authorized_signatory_title", None) or ""
    
    return (
        f"Signed for and on behalf of: {hospital_name}\n"
        f"\n"
        f"Name: {signatory_name}\n"
        f"Title: {signatory_title}\n"
        f"\n"
        f"Signature:\n"
        f"\n"
        f"Date: "
    )


def _generate_sponsor_signature_block_text(
    profile: SiteProfile,
    sponsor_signatory_name: Optional[str] = None,
    sponsor_signatory_email: Optional[str] = None,
    current_user_email: Optional[str] = None
) -> str:
    """Generate sponsor signature block text."""
    contracting_entity = getattr(profile, "primary_contracting_entity", None) or ""
    
    # Get sponsor signatory name
    signatory_name = sponsor_signatory_name
    if not signatory_name:
        if current_user_email:
            signatory_name = current_user_email.split("@")[0]
        else:
            signatory_name = ""
    
    return (
        f"Signed for and on behalf of: {contracting_entity}\n"
        f"\n"
        f"Name: {signatory_name}\n"
        f"\n"
        f"Signature:\n"
        f"\n"
        f"Date: "
    )

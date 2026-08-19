"""
Generate a sample CTA template DOCX for the Template Library.

Contains every Site Profile + Study + IRB placeholder the backend's
`resolve_field_value` knows how to fill, plus a handful of "wild" tokens
(bracket-labels, ALL CAPS labels, underscore fill-ins) for testing the AI
placeholder detector.

Run inside the backend container:
  docker compose exec -e PYTHONPATH=/app backend python scripts/generate_sample_cta.py

The DOCX is written to /app/uploads/sample_cta_template.docx — copy it out
via `docker compose cp` or just upload it through the Template Library UI.
"""
from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as exc:  # pragma: no cover - dev-time hint
    raise SystemExit(
        "python-docx is not installed. Run `pip install python-docx` first."
    ) from exc


def build_doc() -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # ── Title ───────────────────────────────────────────────────────────────
    title = doc.add_heading("CLINICAL TRIAL AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("(Sample template — for Template Library testing)")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph()

    # ── Preamble ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.add_run(
        "This Clinical Trial Agreement (the “Agreement”) is entered into as of "
    )
    p.add_run("[INSERT EFFECTIVE DATE]").bold = True
    p.add_run(" by and between ")
    p.add_run("{{sponsor_name}}").bold = True
    p.add_run(" (the “Sponsor”) and ")
    p.add_run("{{primary_contracting_entity}}").bold = True
    p.add_run(" (the “Institution”), on behalf of investigator ")
    p.add_run("{{pi_name}}").bold = True
    p.add_run(", with respect to the clinical study described below.")

    doc.add_paragraph()

    # ── Article 1: Study ────────────────────────────────────────────────────
    doc.add_heading("1. The Study", level=1)
    doc.add_paragraph().add_run(
        "Sponsor wishes to engage the Institution to conduct the clinical "
        "investigation titled “{{study_name}}” (the “Study”), conducted under "
        "Protocol Number {{study_protocol_number}}, Protocol Version "
        "PROTOCOL VERSION, with a planned study start date of "
        "STUDY START DATE: ____________."
    )

    # ── Article 2: Site & Investigator ──────────────────────────────────────
    doc.add_heading("2. Site and Principal Investigator", level=1)
    p = doc.add_paragraph(
        "The Study shall be conducted at the following site under the "
        "supervision of the named Principal Investigator (“PI”)."
    )

    # Site / PI detail table
    t = doc.add_table(rows=10, cols=2)
    t.style = "Light Grid Accent 1"
    rows = [
        ("Site Name", "{{site_name}}"),
        ("Hospital / Institution", "{{hospital_name}}"),
        ("Site Address", "{{address_line_1}}, {{city}}, {{state}} {{postal_code}}, {{country}}"),
        ("Principal Investigator", "{{pi_name}}"),
        ("PI Designation", "{{pi_designation}}"),
        ("PI Department", "{{pi_department}}"),
        ("PI Email", "{{pi_email}}"),
        ("PI Phone", "{{pi_phone}}"),
        ("Site Coordinator", "{{site_coordinator_name}}"),
        ("Coordinator Contact", "{{site_coordinator_email}} / {{site_coordinator_phone}}"),
    ]
    for i, (label, value) in enumerate(rows):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ── Article 3: IRB / Ethics ─────────────────────────────────────────────
    doc.add_heading("3. Independent Review Board / Ethics Committee", level=1)
    doc.add_paragraph(
        "The Study has been reviewed and approved by the IRB/IEC of record. "
        "Identifiers and contact details are listed below."
    )

    t2 = doc.add_table(rows=4, cols=2)
    t2.style = "Light Grid Accent 1"
    irb_rows = [
        ("IRB / IEC Name", "{{irb_name}}"),
        ("IRB Registration Number", "{{irb_number}}"),
        ("IRB Country / Jurisdiction", "{{irb_country}}"),
        ("IRB Approval Date", "[INSERT IRB APPROVAL DATE]"),
    ]
    for i, (label, value) in enumerate(irb_rows):
        t2.rows[i].cells[0].text = label
        t2.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ── Article 4: Sponsor Authorized Signatory ─────────────────────────────
    doc.add_heading("4. Authorized Signatories", level=1)
    doc.add_paragraph(
        "Each Party represents that the individuals signing this Agreement "
        "on its behalf are duly authorized to do so."
    )
    doc.add_paragraph(
        "Institution Authorized Signatory: {{authorized_signatory_name}}, "
        "{{authorized_signatory_title}} ({{authorized_signatory_email}})."
    )
    doc.add_paragraph(
        "Designated Person for Notices: NAME OF DESIGNATED PERSON: ____________."
    )

    # ── Article 5: Budget ───────────────────────────────────────────────────
    doc.add_heading("5. Budget and Payments", level=1)
    doc.add_paragraph(
        "Sponsor shall pay the Institution in accordance with the Site Budget "
        "appended hereto. The Site Budget — including per-visit costs and "
        "study-wide totals — is appended to this Agreement and forms an "
        "integral part hereof."
    )

    # ── Article 6: Term ─────────────────────────────────────────────────────
    doc.add_heading("6. Term", level=1)
    doc.add_paragraph(
        "This Agreement shall commence on the effective date set out above and "
        "shall continue until the Study is completed, unless terminated "
        "earlier in accordance with the termination provisions herein."
    )

    # ── Article 7: Amendments ───────────────────────────────────────────────
    doc.add_heading("7. Amendments", level=1)
    doc.add_paragraph(
        "Any amendment to this Agreement shall be made in writing and signed "
        "by both Parties. Reference CTA AMENDMENT NUMBER for tracking."
    )

    # ── Signature Block ─────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Signatures", level=1)

    sig = doc.add_table(rows=4, cols=2)
    sig.style = "Light Grid Accent 1"
    sig.rows[0].cells[0].text = "Sponsor"
    sig.rows[0].cells[1].text = "Institution"
    sig.rows[1].cells[0].text = "Signed: ______________________"
    sig.rows[1].cells[1].text = "Signed: ______________________"
    sig.rows[2].cells[0].text = "Name: NAME OF SPONSOR"
    sig.rows[2].cells[1].text = "Name: {{authorized_signatory_name}}"
    sig.rows[3].cells[0].text = "Title: TITLE OF SPONSOR SIGNATORY"
    sig.rows[3].cells[1].text = "Title: {{authorized_signatory_title}}"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Principal Investigator Acknowledgement").bold = True
    doc.add_paragraph("Signed: ______________________")
    doc.add_paragraph("Name: {{pi_name}}")
    doc.add_paragraph("Date: [INSERT PI SIGNATURE DATE]")

    return doc


if __name__ == "__main__":
    out_dir = Path("/app/uploads")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "sample_cta_template.docx"

    document = build_doc()
    document.save(str(out_path))
    print(f"Sample CTA template written to: {out_path}")
    print(
        "\nTo copy out of the container:\n"
        "  docker compose cp backend:/app/uploads/sample_cta_template.docx ./sample_cta_template.docx"
    )

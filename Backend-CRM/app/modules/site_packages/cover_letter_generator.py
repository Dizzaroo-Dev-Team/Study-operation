"""
Standard IRB / Ethics Committee Cover Letter generator (pure Python).

Produces the same layout/intent as Frontend-CRM/scripts/irb-cover-letter/
(coverLetterGenerator.js) but using python-docx so it runs inside the
backend without requiring Node.js at runtime.
"""
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

# ---- Color palette (mirrors coverLetterGenerator.js) -------------------------
NAVY    = RGBColor(0x1B, 0x3A, 0x6B)
STEEL   = RGBColor(0x2E, 0x6E, 0xA6)
SILVER  = "EEF3FA"
MIDGRAY_HEX = "C5D4E8"
MIDGRAY = RGBColor(0xC5, 0xD4, 0xE8)
DKGRAY  = RGBColor(0x44, 0x44, 0x44)
LTGRAY  = RGBColor(0x77, 0x77, 0x77)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
GOLD    = RGBColor(0xC8, 0xA8, 0x4B)

# ---- Regulatory lookup (mirrors regulatoryLookup.js) -------------------------
REGULATORY_MAP: Dict[str, Dict[str, str]] = {
    "India":          {"authority": "CDSCO",          "local_law": "ICMR National Ethical Guidelines for Biomedical Research (2017)"},
    "United States":  {"authority": "FDA",             "local_law": "45 CFR Part 46 (Common Rule) and 21 CFR Parts 50 & 56"},
    "China":          {"authority": "NMPA",            "local_law": "Measures for the Ethical Review of Life Science and Medical Research Involving Humans (2023)"},
    "European Union": {"authority": "EMA",             "local_law": "EU Clinical Trials Regulation 536/2014"},
    "Germany":        {"authority": "BfArM / PEI",     "local_law": "AMG (Arzneimittelgesetz) and EU Regulation 536/2014"},
    "France":         {"authority": "ANSM",            "local_law": "Loi Jarde (Law No. 2012-300) and EU Regulation 536/2014"},
    "United Kingdom": {"authority": "MHRA",            "local_law": "UK Policy Framework for Health and Social Care Research (2017)"},
    "Australia":      {"authority": "TGA",             "local_law": "National Statement on Ethical Conduct in Human Research (2007, updated 2018)"},
    "Canada":         {"authority": "Health Canada",   "local_law": "Tri-Council Policy Statement: Ethical Conduct for Research Involving Humans (TCPS2 2022)"},
    "Japan":          {"authority": "PMDA",            "local_law": "Act on the Protection of Personal Information (APPI) and MHLW GCP Ordinance"},
    "South Korea":    {"authority": "MFDS",            "local_law": "Bioethics and Safety Act (2013, amended 2021)"},
    "Brazil":         {"authority": "ANVISA",          "local_law": "Resolution RDC No. 204/2017 and CNS Resolution No. 466/2012"},
    "Mexico":         {"authority": "COFEPRIS",        "local_law": "NOM-012-SSA3-2012 (Clinical Research Guidelines)"},
    "South Africa":   {"authority": "SAHPRA",          "local_law": "South African Good Clinical Practice Guidelines (3rd ed.)"},
    "Saudi Arabia":   {"authority": "NCBE / SFDA",     "local_law": "Saudi Good Clinical Practice Guidelines and NCBE Regulations"},
    "UAE":            {"authority": "DOH / MOHAP",     "local_law": "UAE Federal Law No. 4 of 2020 on Medical Liability"},
    "Singapore":      {"authority": "HSA",             "local_law": "Medicines Act and HSA CT Regulatory Guidelines"},
    "Israel":         {"authority": "MOH Israel",      "local_law": "Public Health Regulations (Clinical Trials in Humans) 1980"},
    "Switzerland":    {"authority": "Swissmedic",      "local_law": "Human Research Act (HRA/HFG) and ClinO"},
    "Netherlands":    {"authority": "CCMO",            "local_law": "Medical Research Involving Human Subjects Act (WMO)"},
    "Spain":          {"authority": "AEMPS",           "local_law": "RD 1090/2015 and EU Regulation 536/2014"},
    "Italy":          {"authority": "AIFA",            "local_law": "D.Lgs. 211/2003 and EU Regulation 536/2014"},
    "Russia":         {"authority": "Ministry of Health Russia", "local_law": "Federal Law No. 61-FZ on Circulation of Medicines"},
    "Turkey":         {"authority": "TITCK",           "local_law": "Clinical Research Regulation (2017)"},
    "Argentina":      {"authority": "ANMAT",           "local_law": "Disposition 6677/2010 (GCP Guidelines)"},
    "Taiwan":         {"authority": "TFDA",            "local_law": "Pharmaceutical Affairs Act and TFDA CT Guidelines"},
    "Thailand":       {"authority": "FDA Thailand",    "local_law": "ICH E6(R2) as adopted by Thai FDA"},
    "Malaysia":       {"authority": "NPRA",            "local_law": "Malaysian Guidelines for Good Clinical Practice (4th ed.)"},
    "Indonesia":      {"authority": "BPOM",            "local_law": "BPOM Regulation No. 21 of 2021 on Clinical Trials"},
    "Nigeria":        {"authority": "NAFDAC",          "local_law": "NAFDAC Clinical Trial Regulations (2021)"},
    "Egypt":          {"authority": "EDA",             "local_law": "Egyptian Drug Authority Clinical Research Guidelines"},
    "Kenya":          {"authority": "PPB / NCST",      "local_law": "Science, Technology and Innovation Act (2013)"},
    "Pakistan":       {"authority": "DRAP",            "local_law": "DRAP Clinical Trial Regulations (2017)"},
    "Bangladesh":     {"authority": "DGDA",            "local_law": "Drug Control Ordinance 1982 (amended) and DGDA Guidelines"},
}


COUNTRY_ALIASES: Dict[str, str] = {
    "republic of korea": "South Korea",
    "korea, republic of": "South Korea",
    "korea republic": "South Korea",
    "korea (south)": "South Korea",
    "korea": "South Korea",
}


def _canonical_country_name(country: str) -> str:
    cleaned = (country or "").strip()
    if not cleaned:
        return ""
    return COUNTRY_ALIASES.get(cleaned.lower(), cleaned)


def _resolve_regulatory_defaults(data: Dict[str, Any]) -> Dict[str, Any]:
    country = _canonical_country_name(str(data.get("country") or ""))
    defaults = REGULATORY_MAP.get(country) or {
        "authority": data.get("local_regulatory_authority") or "[REGULATORY AUTHORITY - NOT FOUND FOR COUNTRY]",
        "local_law": data.get("additional_local_regulation") or "[LOCAL REGULATION - ADD MANUALLY FOR THIS COUNTRY]",
    }
    merged = dict(data)
    merged["local_regulatory_authority"] = data.get("local_regulatory_authority") or defaults["authority"]
    merged["additional_local_regulation"] = data.get("additional_local_regulation") or defaults["local_law"]
    merged["doi_version"] = data.get("doi_version") or "2013 Fortaleza revision"
    return merged


def _f(data: Dict[str, Any], key: str) -> str:
    v = data.get(key)
    if v is None or v == "":
        return f"[MISSING: {key}]"
    return str(v)


# ---- Low-level docx helpers --------------------------------------------------
def _set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, *, color_hex: str = MIDGRAY_HEX, size: int = 4, style: str = "single") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), style)
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color_hex)
        tc_borders.append(node)
    tc_pr.append(tc_borders)


def _set_no_cell_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        tc_borders.append(node)
    tc_pr.append(tc_borders)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def _add_paragraph_bottom_border(paragraph, color_hex: str, size: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
             color: Optional[RGBColor] = None, size_pt: float = 11.0, font: str = "Arial",
             all_caps: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if all_caps:
        rpr = run._element.get_or_add_rPr()
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "true")
        rpr.append(caps)


def _spacer(doc, *, pt_before: float = 6.0, pt_after: float = 6.0) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(pt_before)
    p.paragraph_format.space_after = Pt(pt_after)


def _divider(doc, color_hex: str = MIDGRAY_HEX) -> None:
    p = doc.add_paragraph()
    _add_paragraph_bottom_border(p, color_hex)
    p.paragraph_format.space_before = Pt(3.0)
    p.paragraph_format.space_after = Pt(3.0)


# ---- Section builders --------------------------------------------------------
def _build_header_banner(doc, data: Dict[str, Any]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    # Approx widths: 6000 / 3360 DXA
    left, right = table.rows[0].cells
    left.width = Twips(6000)
    right.width = Twips(3360)

    # Left cell: Institution block (NAVY background)
    _set_cell_shading(left, "1B3A6B")
    _set_no_cell_border(left)
    _set_cell_margins(left, top=200, bottom=200, left=240, right=120)
    left.paragraphs[0].text = ""
    p1 = left.paragraphs[0]
    _add_run(p1, "CLINICAL RESEARCH OFFICE", bold=True, color=GOLD, size_pt=8, all_caps=True)
    p2 = left.add_paragraph()
    _add_run(p2, _f(data, "sponsor_institution_name"), bold=True, color=WHITE, size_pt=12)
    p3 = left.add_paragraph()
    _add_run(p3, "IRB / Ethics Committee Submission", italic=True, color=MIDGRAY, size_pt=9)

    # Right cell: Date + ref block (STEEL background)
    _set_cell_shading(right, "2E6EA6")
    _set_no_cell_border(right)
    _set_cell_margins(right, top=200, bottom=200, left=200, right=200)
    right.paragraphs[0].text = ""
    r1 = right.paragraphs[0]
    r1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(r1, "DATE OF SUBMISSION", color=MIDGRAY, size_pt=7, all_caps=True)
    r2 = right.add_paragraph()
    r2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(r2, _f(data, "submission_date"), bold=True, color=WHITE, size_pt=11)
    r3 = right.add_paragraph()
    r3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(r3, "PROTOCOL REF.", color=MIDGRAY, size_pt=7, all_caps=True)
    r4 = right.add_paragraph()
    r4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(r4, _f(data, "protocol_number"), bold=True, color=WHITE, size_pt=11)


def _build_address_block(doc, data: Dict[str, Any]) -> None:
    _spacer(doc, pt_before=10.0, pt_after=3.0)

    p = doc.add_paragraph()
    _add_run(p, "TO:", bold=True, color=STEEL)

    p = doc.add_paragraph()
    _add_run(p, _f(data, "irb_chair_name"), color=DKGRAY)
    _add_run(p, "  . Chairperson / Authorized Representative", italic=True, color=LTGRAY, size_pt=10)

    for field in ("irb_committee_name", "institution_name", "institution_address_line1"):
        p = doc.add_paragraph()
        _add_run(p, _f(data, field), color=DKGRAY)

    if data.get("institution_address_line2"):
        p = doc.add_paragraph()
        _add_run(p, str(data["institution_address_line2"]), color=DKGRAY)

    for field in ("city_state_zip", "country"):
        p = doc.add_paragraph()
        _add_run(p, _f(data, field), color=DKGRAY)

    _spacer(doc, pt_before=4.0, pt_after=4.0)
    _divider(doc)


def _build_subject_line(doc, data: Dict[str, Any]) -> None:
    _spacer(doc, pt_before=4.0, pt_after=3.0)
    p = doc.add_paragraph()
    _add_run(p, "RE: ", bold=True, color=DKGRAY)
    _add_run(p, _f(data, "submission_type"), color=DKGRAY)
    _add_run(p, " - Protocol No. ", color=DKGRAY)
    _add_run(p, _f(data, "protocol_number"), color=DKGRAY)

    p = doc.add_paragraph()
    _add_run(p, "      ", color=DKGRAY)
    _add_run(p, _f(data, "study_title"), italic=True, color=DKGRAY)

    _spacer(doc, pt_before=4.0, pt_after=2.0)
    _divider(doc, "2E6EA6")
    _spacer(doc, pt_before=2.0, pt_after=4.0)


def _build_salutation(doc, data: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    _add_run(p, "Dear ", color=DKGRAY)
    _add_run(p, _f(data, "salutation"), color=DKGRAY)
    _add_run(p, ",", color=DKGRAY)
    _spacer(doc, pt_before=5.0, pt_after=2.0)


def _section_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    _add_run(p, text, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(4.0)


def _build_purpose_paragraph(doc, data: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8.0)
    _add_run(p, "On behalf of ", color=DKGRAY)
    _add_run(p, _f(data, "sponsor_institution_name"), bold=True, color=DKGRAY)
    _add_run(p, ", we hereby submit the ", color=DKGRAY)
    _add_run(p, _f(data, "submission_type"), bold=True, color=DKGRAY)
    _add_run(
        p,
        " for the above-referenced clinical investigation to your committee for review and approval. "
        "This submission is made in accordance with applicable local regulations and international "
        "ethical standards governing human subjects research.",
        color=DKGRAY,
    )


def _build_study_overview_paragraph(doc, data: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8.0)
    _add_run(p, "The proposed study, \"", color=DKGRAY)
    _add_run(p, _f(data, "study_title"), italic=True, color=DKGRAY)
    _add_run(p, "\", is briefly described as follows: ", color=DKGRAY)
    _add_run(p, _f(data, "study_description"), color=DKGRAY)
    _add_run(p, " The investigation will be conducted under the principal investigatorship of ", color=DKGRAY)
    _add_run(p, _f(data, "pi_name"), bold=True, color=DKGRAY)
    _add_run(p, ", ", color=DKGRAY)
    _add_run(p, _f(data, "pi_title"), color=DKGRAY)
    _add_run(p, ", at ", color=DKGRAY)
    _add_run(p, _f(data, "study_site_name"), bold=True, color=DKGRAY)
    _add_run(p, ".", color=DKGRAY)


def _build_regulatory_paragraph(doc, data: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    _add_run(p, "This study will be conducted in full compliance with:", color=DKGRAY)

    def _bullet(runs_fn):
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(4.0)
        runs_fn(bp)

    _bullet(lambda bp: (
        _add_run(bp, "The current guidelines of ", color=DKGRAY),
        _add_run(bp, _f(data, "local_regulatory_authority"), bold=True, color=DKGRAY),
        _add_run(bp, ";", color=DKGRAY),
    ))
    _bullet(lambda bp: _add_run(bp, "ICH E6(R2) Good Clinical Practice (GCP) guidelines;", color=DKGRAY))
    _bullet(lambda bp: (
        _add_run(bp, "The Declaration of Helsinki (", color=DKGRAY),
        _add_run(bp, _f(data, "doi_version"), color=DKGRAY),
        _add_run(bp, ") and its applicable amendments;", color=DKGRAY),
    ))
    _bullet(lambda bp: (
        _add_run(bp, _f(data, "additional_local_regulation"), bold=True, color=DKGRAY),
        _add_run(bp, ".", color=DKGRAY),
    ))


def _build_enclosed_docs_table(doc, data: Dict[str, Any]) -> None:
    docs: List[Dict[str, Any]] = list(data.get("enclosed_document_list") or [])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4.0)
    _add_run(p, "The following documents are enclosed with this submission (", color=DKGRAY)
    _add_run(p, _f(data, "submission_version"), bold=True, color=DKGRAY)
    _add_run(p, "):", color=DKGRAY)

    table = doc.add_table(rows=1 + len(docs), cols=4)
    table.autofit = False
    widths = [Twips(600), Twips(5760), Twips(1800), Twips(1200)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    # Header row
    header_titles = ("#", "Document Title", "Version / Date", "File Format")
    header_alignments = (
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    for cell, title, align in zip(table.rows[0].cells, header_titles, header_alignments):
        _set_cell_shading(cell, "1B3A6B")
        _set_cell_border(cell)
        _set_cell_margins(cell)
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].alignment = align
        _add_run(cell.paragraphs[0], title, bold=True, color=WHITE, size_pt=9)

    # Data rows
    for i, d in enumerate(docs):
        row_cells = table.rows[i + 1].cells
        zebra_even = (i % 2 == 0)

        num_cell, title_cell, ver_cell, fmt_cell = row_cells

        for cell in row_cells:
            if zebra_even:
                _set_cell_shading(cell, SILVER)
            _set_cell_border(cell)
            _set_cell_margins(cell)
            cell.paragraphs[0].text = ""

        num_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(num_cell.paragraphs[0], str(i + 1), color=DKGRAY)

        _add_run(title_cell.paragraphs[0], str(d.get("title") or ""), color=DKGRAY)

        ver_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(ver_cell.paragraphs[0], str(d.get("version") or ""), color=DKGRAY)

        fmt_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(fmt_cell.paragraphs[0], str(d.get("format") or ""), color=DKGRAY)

    _spacer(doc, pt_before=4.0, pt_after=4.0)
    total = doc.add_paragraph()
    _add_run(total, "Total documents submitted: ", italic=True, color=LTGRAY, size_pt=10)
    _add_run(total, str(len(docs)), bold=True, color=DKGRAY)
    _add_run(total, ".", italic=True, color=LTGRAY, size_pt=10)
    _spacer(doc, pt_before=4.0, pt_after=6.0)


def _build_closing_paragraph(doc, data: Dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8.0)
    _add_run(p, "We respectfully request that your committee undertake a ", color=DKGRAY)
    _add_run(p, _f(data, "review_type"), bold=True, color=DKGRAY)
    _add_run(
        p,
        " of the above documentation. Please direct any queries or requests for additional "
        "information to the contact below. We look forward to your written determination within "
        "the timeframe specified by your committee's Standard Operating Procedures.",
        color=DKGRAY,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4.0)
    _add_run(p, "Primary Contact: ", color=DKGRAY)
    _add_run(p, _f(data, "primary_contact_name"), bold=True, color=DKGRAY)
    _add_run(p, "  |  ", color=DKGRAY)
    _add_run(p, _f(data, "primary_contact_email"), color=DKGRAY)
    _add_run(p, "  |  ", color=DKGRAY)
    _add_run(p, _f(data, "primary_contact_phone"), color=DKGRAY)
    _spacer(doc, pt_before=4.0, pt_after=6.0)


def _build_signature_block(doc, data: Dict[str, Any]) -> None:
    _divider(doc)
    _spacer(doc, pt_before=4.0, pt_after=4.0)
    p = doc.add_paragraph()
    _add_run(p, "Sincerely,", bold=True, color=NAVY)
    _spacer(doc, pt_before=2.0, pt_after=2.0)

    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    sig_cell, spacer_cell, pi_cell = table.rows[0].cells
    sig_cell.width = Twips(2400)
    spacer_cell.width = Twips(300)
    pi_cell.width = Twips(6660)

    for cell in (sig_cell, pi_cell):
        _set_no_cell_border(cell)
        _set_cell_margins(cell, left=0, right=200)
    _set_no_cell_border(spacer_cell)

    # Signature stamp cell
    sig_cell.paragraphs[0].text = ""
    _add_run(sig_cell.paragraphs[0], "[ ", color=LTGRAY, size_pt=10)
    _add_run(sig_cell.paragraphs[0], _f(data, "signature_stamp"), italic=True, color=DKGRAY)
    _add_run(sig_cell.paragraphs[0], " ]", color=LTGRAY, size_pt=10)
    sp = sig_cell.add_paragraph()
    _add_run(sp, "Wet ink / e-Signature / Stamp", italic=True, color=LTGRAY, size_pt=10)

    # PI block cell
    pi_cell.paragraphs[0].text = ""
    _add_run(pi_cell.paragraphs[0], _f(data, "pi_name"), bold=True, color=DKGRAY)
    for field in ("pi_title", "pi_department", "sponsor_institution_name"):
        p = pi_cell.add_paragraph()
        _add_run(p, _f(data, field), color=DKGRAY)
    if data.get("co_pi_name") or data.get("co_pi_title"):
        gap = pi_cell.add_paragraph()
        gap.paragraph_format.space_before = Pt(3.0)
        co = pi_cell.add_paragraph()
        _add_run(co, "Co-Signatory (if applicable):", italic=True, color=LTGRAY, size_pt=10)
        _add_run(co, "  ", color=DKGRAY)
        _add_run(co, str(data.get("co_pi_name") or ""), color=DKGRAY)
        co2 = pi_cell.add_paragraph()
        _add_run(co2, str(data.get("co_pi_title") or ""), color=DKGRAY)

    _spacer(doc, pt_before=4.0, pt_after=4.0)
    _divider(doc, "C8A84B")
    _spacer(doc, pt_before=2.0, pt_after=2.0)

    meta = doc.add_paragraph()
    _add_run(meta, "Submission ID (system-generated): ", italic=True, color=LTGRAY, size_pt=10)
    _add_run(meta, str(data.get("system_submission_id") or ""), color=DKGRAY)
    _add_run(meta, "   |   IRB Reference No. (if assigned): ", italic=True, color=LTGRAY, size_pt=10)
    _add_run(meta, str(data.get("irb_assigned_ref") or ""), color=DKGRAY)


# ---- Public API --------------------------------------------------------------
def generate_cover_letter_docx_bytes(data: Dict[str, Any]) -> bytes:
    """
    Build the standardized IRB cover letter as a .docx file and return its bytes.
    Accepts the same payload shape as coverLetterGenerator.js.
    """
    resolved = _resolve_regulatory_defaults(data or {})

    doc = Document()

    # Default paragraph font/color like the JS generator's document default.
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = DKGRAY

    # Page size / margins (US Letter, DXA values approximated via EMU here).
    section = doc.sections[0]
    section.left_margin = Twips(1080)
    section.right_margin = Twips(1080)
    section.top_margin = Twips(720)
    section.bottom_margin = Twips(1080)

    _build_header_banner(doc, resolved)
    _build_address_block(doc, resolved)
    _build_subject_line(doc, resolved)
    _build_salutation(doc, resolved)

    _section_heading(doc, "1.  PURPOSE OF SUBMISSION")
    _build_purpose_paragraph(doc, resolved)
    _spacer(doc, pt_before=3.0, pt_after=3.0)

    _section_heading(doc, "2.  STUDY OVERVIEW")
    _build_study_overview_paragraph(doc, resolved)
    _spacer(doc, pt_before=3.0, pt_after=3.0)

    _section_heading(doc, "3.  REGULATORY & ETHICAL COMPLIANCE")
    _build_regulatory_paragraph(doc, resolved)
    _spacer(doc, pt_before=3.0, pt_after=3.0)

    _section_heading(doc, "4.  ENCLOSED DOCUMENTS")
    _build_enclosed_docs_table(doc, resolved)

    _section_heading(doc, "5.  REVIEW REQUEST & CLOSING")
    _build_closing_paragraph(doc, resolved)
    _build_signature_block(doc, resolved)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()

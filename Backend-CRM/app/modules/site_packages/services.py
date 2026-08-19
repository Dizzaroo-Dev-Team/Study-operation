from __future__ import annotations

import asyncio
from copy import deepcopy
import logging
import os
import re
import time
from urllib.parse import urlparse, urlunparse
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple
from uuid import UUID

import httpx
from fastapi import HTTPException, UploadFile
from sqlalchemy import and_, func, or_, select
from sqlalchemy.ext.asyncio import AsyncSession
from docx import Document
from docx.text.paragraph import Paragraph

from app.modules.sites.routes.irb_requirements import resolve_requirements_for_irb
from app.config import settings
from app.models import IRB, Site, SiteIRBMapping, SiteProfile, Study, User, IRBAdministrativeInfo
from app.utils.docx_placeholder_replace import remove_document_protection, replace_placeholders_in_docx
from app.utils.log_sanitize import sfmt
from app.utils.docx_to_pdf import docx_to_pdf

from .models import SitePackage

logger = logging.getLogger(__name__)


def _make_named_tempfile(suffix: str, data: Optional[bytes] = None) -> Path:
    """Create a NamedTemporaryFile(delete=False) with ``suffix``, optionally
    writing ``data`` into it, and return its path. Sync on purpose — call via
    ``asyncio.to_thread`` from async code so the blocking file I/O stays off
    the event loop."""
    from tempfile import NamedTemporaryFile

    with NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        if data is not None:
            tmp.write(data)
        return Path(tmp.name)


_FALLBACK_TMC_REQUIRED_DOCUMENTS: List[str] = [
    "Cover Letter (This document)",
    "IEC Application Form (Duly filled and signed)",
    "Principal Investigator Undertaking",
    "Conflict of Interest (COI) Declaration",
    "Main Clinical Trial Protocol",
    "Investigator's Brochure (IB)",
    "Master Informed Consent Document (ICD) - English",
    "Translated ICDs",
    "Translation Certificates / Back Translations",
    "Patient Diaries / Questionnaires / PROs",
    "Patient Recruitment Materials",
    "PI & Co-I CVs and Medical Registrations",
    "Good Clinical Practice (GCP) Certificates",
    "CDSCO / DCGI Approval (NOC)",
    "CTRI Registration Proof",
    "Clinical Trial Agreement (CTA) / Draft Budget",
    "Insurance Policy / Certificate",
    "Patient Compensation Formula / Undertaking",
    "HMSC Approval / Proof of Submission",
]

_MD_ANDERSON_IRB_NAME = "md anderson irb"
_MD_ANDERSON_IRB_CODE = "md_anderson_irb"
_HEIDELBERG_EC_NAME = "heidelberg university ethics committee (germany)"
_HEIDELBERG_EC_LEGACY_NAME = "ethics committee of the medical faculty of heidelberg university"
_HEIDELBERG_EC_CODE = "heidelberg_medical_faculty_ec"
_PETER_MAC_HREC_NAME = "peter maccallum cancer centre human research ethics committee (peter mac hrec), australia"
_PETER_MAC_HREC_CODE = "peter_mac_hrec"
_NCC_IRB_JAPAN_NAME = "national cancer center institutional review board (ncc irb), japan"
_NCC_IRB_JAPAN_CODE = "ncc_irb_japan"
_CHCAMS_EC_NAME = "national cancer center / chcams institutional ethics committee"
_CHCAMS_EC_LEGACY_NAME = (
    "ethics committee of national cancer center / cancer hospital, chinese academy of medical "
    "sciences and peking union medical college (chcams ethics committee)"
)
_CHCAMS_EC_CODE = "chcams_ethics_committee"
_HRA_OXFORD_REC_UK_NAME = "the oxford cancer & heamatology centre , united kingdom"
_HRA_OXFORD_REC_UK_CODE = "hra_oxford_rec_uk"
_SMC_IRB_NAME = "samsung medical center institutional review board (smc irb)"
_SMC_IRB_CODE = "smc_irb"

_COVER_LETTER_TYPES: Set[str] = {
    "cover_letter",
    "tmc_cover_letter",
    "cover_submission_letter",
    "notification_cover_letter",
    "protocol_deviation_cover_letter",
}

_AMENDMENT_DOC_MAX_BY_TYPE: Dict[str, int] = {
    "summary_of_changes": 1,
    "amended_protocol_clean_tracked": 10_000,
    "updated_icf": 10_000,
    "revised_recruitment_materials": 10_000,
    "updated_ib_safety_efficacy": 1,
    "impact_assessment_risk_benefit": 1,
    "updated_questionnaires_crfs": 10_000,
}

_AMENDMENT_DOC_TYPE_BY_NAME_LOWER: Dict[str, str] = {
    "summary of changes": "summary_of_changes",
    "amended protocol (clean + tracked/redline)": "amended_protocol_clean_tracked",
    "updated icf (if applicable)": "updated_icf",
    "revised recruitment materials": "revised_recruitment_materials",
    "updated ib (if safety/efficacy changes)": "updated_ib_safety_efficacy",
    "impact assessment (risk/benefit)": "impact_assessment_risk_benefit",
    "updated questionnaires or crfs": "updated_questionnaires_crfs",
}

_AMENDMENT_DOC_MAX_BY_NAME_LOWER: Dict[str, int] = {
    name: _AMENDMENT_DOC_MAX_BY_TYPE[type_key]
    for name, type_key in _AMENDMENT_DOC_TYPE_BY_NAME_LOWER.items()
}

_CONTINUING_REVIEW_DOC_MAX_BY_TYPE: Dict[str, int] = {
    "progress_report_status_summary": 1,
    "site_enrollment_numbers": 1,
    "safety_summary_ae_sae_overview": 1,
    "protocol_deviation_violation_log_summary": 1,
    "updated_risk_benefit_assessment": 1,
    "current_icf_latest_approved_version": 10_000,
    "current_irb_approval_letter": 1,
    "publication_summary_if_any": 10_000,
}

_CONTINUING_REVIEW_DOC_TYPE_BY_NAME_LOWER: Dict[str, str] = {
    "progress report / status summary": "progress_report_status_summary",
    "site enrollment numbers": "site_enrollment_numbers",
    "safety summary (ae/sae overview)": "safety_summary_ae_sae_overview",
    "protocol deviation / violation log summary": "protocol_deviation_violation_log_summary",
    "updated risk-benefit assessment": "updated_risk_benefit_assessment",
    "current icf (latest approved version)": "current_icf_latest_approved_version",
    "current irb approval letter": "current_irb_approval_letter",
    "publication summary (if any)": "publication_summary_if_any",
}

_CONTINUING_REVIEW_DOC_MAX_BY_NAME_LOWER: Dict[str, int] = {
    name: _CONTINUING_REVIEW_DOC_MAX_BY_TYPE[type_key]
    for name, type_key in _CONTINUING_REVIEW_DOC_TYPE_BY_NAME_LOWER.items()
}

_SAFETY_REPORT_DOC_MAX_BY_TYPE: Dict[str, int] = {
    "sae_report_susar_report": 10_000,
    "narrative_report": 1,
    "causality_assessment": 1,
    "line_listing": 1,
    "dsmb_recommendations": 1,
    "updated_ib": 1,
}

_SAFETY_REPORT_DOC_TYPE_BY_NAME_LOWER: Dict[str, str] = {
    "sae report / susar report": "sae_report_susar_report",
    "narrative report": "narrative_report",
    "causality assessment": "causality_assessment",
    "line listing": "line_listing",
    "dsmb recommendations": "dsmb_recommendations",
    "updated ib": "updated_ib",
}

_SAFETY_REPORT_DOC_MAX_BY_NAME_LOWER: Dict[str, int] = {
    name: _SAFETY_REPORT_DOC_MAX_BY_TYPE[type_key]
    for name, type_key in _SAFETY_REPORT_DOC_TYPE_BY_NAME_LOWER.items()
}

_STUDY_CLOSURE_DOC_MAX_BY_TYPE: Dict[str, int] = {
    "final_study_report_close_out_report": 1,
    "enrollment_summary_final": 1,
    "safety_summary_cumulative": 1,
    "deviations_summary": 1,
    "publication_status": 10_000,
    "irb_closure_acknowledgment_letter": 1,
    "drug_device_accountability_records": 10_000,
    "archival_plan_statement": 1,
    "final_study_financial_disclosure_log": 10_000,
}

_STUDY_CLOSURE_DOC_TYPE_BY_NAME_LOWER: Dict[str, str] = {
    "final study report / close-out report": "final_study_report_close_out_report",
    "enrollment summary (final)": "enrollment_summary_final",
    "safety summary (cumulative)": "safety_summary_cumulative",
    "deviations summary": "deviations_summary",
    "publication status": "publication_status",
    "irb closure acknowledgment letter": "irb_closure_acknowledgment_letter",
    "drug/device accountability records": "drug_device_accountability_records",
    "archival plan statement": "archival_plan_statement",
    "final study financial disclosure log": "final_study_financial_disclosure_log",
}

_STUDY_CLOSURE_DOC_MAX_BY_NAME_LOWER: Dict[str, int] = {
    name: _STUDY_CLOSURE_DOC_MAX_BY_TYPE[type_key]
    for name, type_key in _STUDY_CLOSURE_DOC_TYPE_BY_NAME_LOWER.items()
}

_PROTOCOL_DEVIATION_DOC_MAX_BY_TYPE: Dict[str, int] = {
    "protocol_deviation_cover_letter": 1,
    "deviation_report_form": 1,
    "root_cause_analysis": 1,
    "capa_corrective_preventive_actions": 1,
    "impact_subject_safety_data_integrity": 1,
    "pi_acknowledgement_memo_to_file_mtf": 10_000,
}

_PROTOCOL_DEVIATION_DOC_TYPE_BY_NAME_LOWER: Dict[str, str] = {
    "deviation report form": "deviation_report_form",
    "root cause analysis": "root_cause_analysis",
    "capa (corrective and preventive actions)": "capa_corrective_preventive_actions",
    "impact on subject safety/data integrity": "impact_subject_safety_data_integrity",
    "pi acknowledgement / memo to file (mtf)": "pi_acknowledgement_memo_to_file_mtf",
}

_PROTOCOL_DEVIATION_DOC_MAX_BY_NAME_LOWER: Dict[str, int] = {
    name: _PROTOCOL_DEVIATION_DOC_MAX_BY_TYPE[type_key]
    for name, type_key in _PROTOCOL_DEVIATION_DOC_TYPE_BY_NAME_LOWER.items()
}

_NOTIFICATION_DOC_MAX_BY_TYPE: Dict[str, int] = {
    "notification_cover_letter": 1,
    "updated_fda_1572_regulatory_statement": 1,
    "incoming_pi_cv": 1,
    "incoming_pi_medical_license": 1,
    "incoming_pi_gcp_certificate": 1,
    "incoming_pi_financial_disclosure": 1,
    "pi_change_irb_submission_receipt": 1,
}

_NOTIFICATION_DOC_TYPE_BY_NAME_LOWER: Dict[str, str] = {
    "updated form fda 1572 / regulatory statement of investigator": "updated_fda_1572_regulatory_statement",
    "current curriculum vitae (cv) of incoming pi": "incoming_pi_cv",
    "active medical license": "incoming_pi_medical_license",
    "good clinical practice (gcp) training certificate": "incoming_pi_gcp_certificate",
    "executed financial disclosure form (fdf)": "incoming_pi_financial_disclosure",
    "irb/iec submission receipt for pi change": "pi_change_irb_submission_receipt",
}

_NOTIFICATION_DOC_MAX_BY_NAME_LOWER: Dict[str, int] = {
    name: _NOTIFICATION_DOC_MAX_BY_TYPE[type_key]
    for name, type_key in _NOTIFICATION_DOC_TYPE_BY_NAME_LOWER.items()
}

# Per-submission-type required document sets now live in MongoDB (see
# notification_templates.py — collection `submission_type_documents`, cached
# lookup `workflow_requirements_for`). "Notification" intentionally has no
# entry there: notification packages are built by the Notification wizard with
# per-type document sets, so progress is computed from the package's own
# documents (see _progress_for_site_package).
from .notification_templates import workflow_requirements_for as _workflow_requirements_for_submission_type_async  # noqa: E402


def _now_utc() -> datetime:
    return datetime.now(timezone.utc)


def _normalize_irb_name(name: Optional[str]) -> str:
    if not name:
        return ""
    return " ".join(str(name).strip().lower().split())


def _normalize_unique_code(code: Optional[str]) -> str:
    if not code:
        return ""
    return "_".join(str(code).strip().lower().split())


def _resolve_onlyoffice_internal_base(raw_base: str) -> str:
    """Return a backend URL that ONLYOFFICE can call from Docker network."""
    parsed = urlparse(raw_base or "")
    if parsed.hostname in {"localhost", "127.0.0.1"} and parsed.port in {3000, 5173}:
        return "http://host.docker.internal:8000"
    return (raw_base or "").rstrip("/")


def _is_cover_letter_document_entry(doc: Dict[str, Any]) -> bool:
    raw_type = str(doc.get("type") or "").strip().lower()
    if raw_type in _COVER_LETTER_TYPES:
        return True
    raw_name = str(doc.get("name") or "").strip().lower()
    return "cover letter" in raw_name


def _normalize_cover_doc_name(value: Optional[str]) -> str:
    raw = str(value or "").strip().lower()
    if not raw:
        return ""
    raw = raw.replace("_", " ").replace("-", " ")
    # Compare by semantic base name instead of extension/format details.
    raw = re.sub(r"\.(docx|pdf)$", "", raw)
    raw = re.sub(r"\s+", " ", raw).strip()
    return raw


def _resolve_cover_docx_url(doc: Dict[str, Any]) -> str:
    docx_url = str(doc.get("docxFileUrl") or "").strip()
    if docx_url:
        return docx_url
    fallback_url = str(doc.get("fileUrl") or "").strip()
    fallback_mime = str(doc.get("mimeType") or "").lower()
    fallback_name = str(doc.get("name") or "").lower()
    if fallback_url and (
        "wordprocessingml" in fallback_mime
        or fallback_name.endswith(".docx")
    ):
        doc["docxFileUrl"] = fallback_url
        return fallback_url
    return ""


def _select_cover_letter_document(
    *,
    documents: List[Any],
    requested_name: str,
    require_editable_docx: bool,
) -> Tuple[int, Optional[Dict[str, Any]]]:
    cover_docs: List[Tuple[int, Dict[str, Any]]] = []
    for idx, doc in enumerate(documents):
        if not isinstance(doc, dict):
            continue
        if not _is_cover_letter_document_entry(doc):
            continue
        cover_docs.append((idx, doc))
    if not cover_docs:
        return -1, None

    requested_raw = str(requested_name or "").strip().lower()
    requested_norm = _normalize_cover_doc_name(requested_raw)

    def _acceptable(doc: Dict[str, Any]) -> bool:
        if not require_editable_docx:
            return True
        return bool(_resolve_cover_docx_url(doc))

    # 1) Exact raw name match (legacy behavior).
    if requested_raw:
        for idx, doc in cover_docs:
            if str(doc.get("name") or "").strip().lower() != requested_raw:
                continue
            if _acceptable(doc):
                return idx, doc

    # 2) Normalized match (ignores extension/underscore/hyphen differences).
    if requested_norm:
        for idx, doc in cover_docs:
            doc_norm = _normalize_cover_doc_name(doc.get("name"))
            if doc_norm != requested_norm:
                continue
            if _acceptable(doc):
                return idx, doc

    # 3) Type-aware match for submission-cover variants across IRBs.
    if requested_norm:
        wants_submission = "submission" in requested_norm
        for idx, doc in cover_docs:
            doc_type = str(doc.get("type") or "").strip().lower()
            doc_name_norm = _normalize_cover_doc_name(doc.get("name"))
            is_submission = (
                doc_type == "cover_submission_letter"
                or "submission" in doc_name_norm
            )
            if wants_submission != is_submission:
                continue
            if _acceptable(doc):
                return idx, doc

    # 4) Final fallback: first editable cover doc (or first cover doc).
    for idx, doc in cover_docs:
        if _acceptable(doc):
            return idx, doc
    return -1, None


def _rewrite_onlyoffice_callback_url_if_needed(url: str) -> str:
    """Rewrite localhost callback URL to internal ONLYOFFICE host when needed."""
    raw = (url or "").strip()
    if not raw:
        return raw
    try:
        parsed = urlparse(raw)
        if parsed.hostname not in {"localhost", "127.0.0.1"}:
            return raw
        internal = urlparse(settings.onlyoffice_url or "")
        new_netloc = internal.netloc or "onlyoffice"
        new_scheme = internal.scheme or parsed.scheme or "http"
        return urlunparse((new_scheme, new_netloc, parsed.path, parsed.params, parsed.query, parsed.fragment))
    except Exception:
        return raw


def _remove_docx_protection_bytes(data: bytes) -> bytes:
    """Best-effort removal of Word forms protection for editor compatibility."""
    from tempfile import NamedTemporaryFile

    with NamedTemporaryFile(suffix=".docx", delete=False) as src:
        src_path = Path(src.name)
        src.write(data)
    with NamedTemporaryFile(suffix=".docx", delete=False) as dst:
        dst_path = Path(dst.name)

    try:
        remove_document_protection(src_path, dst_path)
        return dst_path.read_bytes()
    except Exception:
        return data
    finally:
        try:
            src_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            dst_path.unlink(missing_ok=True)
        except Exception:
            pass


def _is_md_anderson_irb(irb: Optional[IRB]) -> bool:
    if not irb:
        return False
    uc = _normalize_unique_code(getattr(irb, "unique_code", None))
    if uc and uc == _MD_ANDERSON_IRB_CODE:
        return True
    name = _normalize_irb_name(irb.name)
    if name == _MD_ANDERSON_IRB_NAME:
        return True
    return "md anderson irb" in name or (
        "md anderson cancer center" in name and "institutional review board" in name
    )


def _is_oxford_rec_irb(irb: Optional[IRB]) -> bool:
    if not irb:
        return False
    uc = _normalize_unique_code(getattr(irb, "unique_code", None))
    if uc and uc == _HRA_OXFORD_REC_UK_CODE:
        return True
    name = _normalize_irb_name(irb.name)
    if name == _HRA_OXFORD_REC_UK_NAME:
        return True
    return (
        "south central" in name
        and "oxford" in name
        and "research ethics committee" in name
    ) or (
        "oxford" in name
        and "cancer" in name
        and ("heamatology" in name or "hematology" in name)
        and "united kingdom" in name
    ) or (
        "health research authority" in name
        and "oxford rec" in name
        and "uk" in name
    ) or (
        "hra" in name
        and "south central" in name
        and "oxford" in name
        and "rec" in name
    )


def _is_heidelberg_ec_irb(irb: Optional[IRB]) -> bool:
    if not irb:
        return False
    uc = _normalize_unique_code(getattr(irb, "unique_code", None))
    if uc and uc == _HEIDELBERG_EC_CODE:
        return True
    name = _normalize_irb_name(irb.name)
    if name in {_HEIDELBERG_EC_NAME, _HEIDELBERG_EC_LEGACY_NAME}:
        return True
    return (
        "heidelberg" in name
        and "ethics committee" in name
        and (
            "medical faculty" in name
            or "medizinischen fakultat" in name
            or "medizinische fakultat" in name
            or "ethikkommission" in name
        )
    )


def _is_smc_irb(irb: Optional[IRB]) -> bool:
    if not irb:
        return False
    uc = _normalize_unique_code(getattr(irb, "unique_code", None))
    if uc and uc == _SMC_IRB_CODE:
        return True
    name = _normalize_irb_name(irb.name)
    if name == _SMC_IRB_NAME:
        return True
    return "smc irb" in name or (
        "samsung medical center" in name and "institutional review board" in name
    )


def _is_chcams_ec_irb(irb: Optional[IRB]) -> bool:
    if not irb:
        return False
    uc = _normalize_unique_code(getattr(irb, "unique_code", None))
    if uc and uc == _CHCAMS_EC_CODE:
        return True
    name = _normalize_irb_name(irb.name)
    if name in {_CHCAMS_EC_NAME, _CHCAMS_EC_LEGACY_NAME}:
        return True
    return (
        ("chcams ethics committee" in name)
        or (
            "national cancer center" in name
            and "chcams" in name
            and "ethics committee" in name
        )
    )


def _is_ncc_irb_japan(irb: Optional[IRB]) -> bool:
    if not irb:
        return False
    uc = _normalize_unique_code(getattr(irb, "unique_code", None))
    if uc and uc == _NCC_IRB_JAPAN_CODE:
        return True
    name = _normalize_irb_name(irb.name)
    if name == _NCC_IRB_JAPAN_NAME:
        return True
    if "ncc irb" in name and "japan" in name:
        return True
    return (
        "national cancer center" in name
        and "institutional review board" in name
        and "japan" in name
        and "chcams" not in name
    )


def _is_peter_mac_hrec_irb(irb: Optional[IRB]) -> bool:
    if not irb:
        return False
    uc = _normalize_unique_code(getattr(irb, "unique_code", None))
    if uc and uc == _PETER_MAC_HREC_CODE:
        return True
    name = _normalize_irb_name(irb.name)
    if name == _PETER_MAC_HREC_NAME:
        return True
    return (
        "peter maccallum" in name
        and "peter mac hrec" in name
        and "human research ethics committee" in name
    )


def _resolve_tmc_cover_letter_template_path() -> Path:
    """
    Resolve TMC IEC cover letter template path from known locations.
    """
    env_path = (os.environ.get("TMC_IEC_COVER_LETTER_TEMPLATE") or "").strip()
    module_path = Path(__file__).resolve()
    cwd = Path.cwd()
    filename = "TMC_IEC_Cover_Letter.docx"

    candidates = [
        Path(env_path) if env_path else None,
        # Preferred in-repo location for site package document templates
        module_path.parent / "templates" / filename,  # Backend-CRM/app/site_packages/templates/TMC_IEC_Cover_Letter.docx
        # Relative to module tree
        module_path.parents[2] / filename,  # Backend-CRM/TMC_IEC_Cover_Letter.docx
        module_path.parents[3] / filename,  # CRM_/TMC_IEC_Cover_Letter.docx
        module_path.parents[2] / "templates" / filename,
        # Relative to runtime cwd (backend may be started from different location)
        cwd / filename,
        cwd.parent / filename,
        cwd / "templates" / filename,
        cwd.parent / "templates" / filename,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue

    # Last fallback: scan a few likely roots non-recursively/limited recursively.
    home = Path.home()
    scan_roots = [
        cwd,
        cwd.parent,
        module_path.parents[2],
        module_path.parents[3],
        home / "Desktop",
        home / "OneDrive" / "Desktop",
        home / "Documents",
    ]
    for root in scan_roots:
        try:
            direct = root / filename
            if direct.exists():
                return direct
            # Case-insensitive check in root directory first.
            for child in root.iterdir():
                if child.is_file() and child.name.lower() == filename.lower():
                    return child
            # Recursive fallback.
            for hit in root.glob("**/*"):
                if hit.is_file() and hit.name.lower() == filename.lower():
                    return hit
        except Exception:
            continue

    raise HTTPException(
        status_code=500,
        detail=(
            "TMC IEC cover letter template not found. Expected TMC_IEC_Cover_Letter.docx. "
            "Set TMC_IEC_COVER_LETTER_TEMPLATE to an absolute file path if needed."
        ),
    )


def _resolve_md_anderson_cover_letter_template_path() -> Path:
    """
    Resolve MD Anderson IRB cover letter template path from known locations.
    """
    env_path = (os.environ.get("MD_ANDERSON_COVER_LETTER_TEMPLATE") or "").strip()
    module_path = Path(__file__).resolve()
    cwd = Path.cwd()
    filename = "CoverLetter_MDAnderson_IRB.docx"

    candidates = [
        Path(env_path) if env_path else None,
        module_path.parent / "templates" / filename,
        module_path.parents[2] / filename,
        module_path.parents[3] / filename,
        module_path.parents[2] / "templates" / filename,
        cwd / filename,
        cwd.parent / filename,
        cwd / "templates" / filename,
        cwd.parent / "templates" / filename,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue

    raise HTTPException(
        status_code=500,
        detail=(
            "MD Anderson cover letter template not found. Expected CoverLetter_MDAnderson_IRB.docx. "
            "Set MD_ANDERSON_COVER_LETTER_TEMPLATE to an absolute file path if needed."
        ),
    )


def _resolve_oxford_cover_letter_template_path() -> Path:
    """
    Resolve Oxford REC cover letter template path from known locations.
    """
    env_path = (os.environ.get("OXFORD_REC_COVER_LETTER_TEMPLATE") or "").strip()
    module_path = Path(__file__).resolve()
    cwd = Path.cwd()
    filename = "CoverLetter_Oxford_REC.docx"

    candidates = [
        Path(env_path) if env_path else None,
        module_path.parent / "templates" / filename,
        module_path.parents[2] / filename,
        module_path.parents[3] / filename,
        module_path.parents[2] / "templates" / filename,
        cwd / filename,
        cwd.parent / filename,
        cwd / "templates" / filename,
        cwd.parent / "templates" / filename,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue

    raise HTTPException(
        status_code=500,
        detail=(
            "Oxford REC cover letter template not found. Expected CoverLetter_Oxford_REC.docx. "
            "Set OXFORD_REC_COVER_LETTER_TEMPLATE to an absolute file path if needed."
        ),
    )


def _resolve_heidelberg_cover_letter_template_path() -> Path:
    """
    Resolve Heidelberg Ethics Committee cover letter template path from known locations.
    """
    env_path = (os.environ.get("HEIDELBERG_EC_COVER_LETTER_TEMPLATE") or "").strip()
    module_path = Path(__file__).resolve()
    cwd = Path.cwd()
    filename = "CoverLetter_Heidelberg_EC.docx"

    candidates = [
        Path(env_path) if env_path else None,
        module_path.parent / "templates" / filename,
        module_path.parents[2] / filename,
        module_path.parents[3] / filename,
        module_path.parents[2] / "templates" / filename,
        cwd / filename,
        cwd.parent / filename,
        cwd / "templates" / filename,
        cwd.parent / "templates" / filename,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue

    raise HTTPException(
        status_code=500,
        detail=(
            "Heidelberg EC cover letter template not found. Expected CoverLetter_Heidelberg_EC.docx. "
            "Set HEIDELBERG_EC_COVER_LETTER_TEMPLATE to an absolute file path if needed."
        ),
    )


def _resolve_smc_cover_letter_template_path() -> Path:
    """
    Resolve Samsung Medical Center IRB cover letter template path from known locations.
    """
    env_path = (os.environ.get("SMC_IRB_COVER_LETTER_TEMPLATE") or "").strip()
    module_path = Path(__file__).resolve()
    cwd = Path.cwd()
    filename = "CoverLetter_SamsungMC_IRB.docx"

    candidates = [
        Path(env_path) if env_path else None,
        module_path.parent / "templates" / filename,
        module_path.parents[2] / filename,
        module_path.parents[3] / filename,
        module_path.parents[2] / "templates" / filename,
        cwd / filename,
        cwd.parent / filename,
        cwd / "templates" / filename,
        cwd.parent / "templates" / filename,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue

    raise HTTPException(
        status_code=500,
        detail=(
            "Samsung Medical Center IRB cover letter template not found. Expected CoverLetter_SamsungMC_IRB.docx. "
            "Set SMC_IRB_COVER_LETTER_TEMPLATE to an absolute file path if needed."
        ),
    )


def _resolve_chcams_cover_letter_template_path() -> Path:
    """
    Resolve CHCAMS / NCC China ethics committee cover letter template path from known locations.
    """
    env_path = (os.environ.get("CHCAMS_EC_COVER_LETTER_TEMPLATE") or "").strip()
    module_path = Path(__file__).resolve()
    cwd = Path.cwd()
    filename = "CoverLetter_NCC_China_IEC.docx"

    candidates = [
        Path(env_path) if env_path else None,
        module_path.parent / "templates" / filename,
        module_path.parents[2] / filename,
        module_path.parents[3] / filename,
        module_path.parents[2] / "templates" / filename,
        cwd / filename,
        cwd.parent / filename,
        cwd / "templates" / filename,
        cwd.parent / "templates" / filename,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue

    raise HTTPException(
        status_code=500,
        detail=(
            "CHCAMS China IEC cover letter template not found. Expected CoverLetter_NCC_China_IEC.docx. "
            "Set CHCAMS_EC_COVER_LETTER_TEMPLATE to an absolute file path if needed."
        ),
    )


def _resolve_ncc_japan_cover_letter_template_path() -> Path:
    """
    Resolve NCC Japan IRB cover letter template path from known locations.
    """
    env_path = (os.environ.get("NCC_JAPAN_IRB_COVER_LETTER_TEMPLATE") or "").strip()
    module_path = Path(__file__).resolve()
    cwd = Path.cwd()
    filename = "CoverLetter_NCC_Japan_IRB.docx"

    candidates = [
        Path(env_path) if env_path else None,
        module_path.parent / "templates" / filename,
        module_path.parents[2] / filename,
        module_path.parents[3] / filename,
        module_path.parents[2] / "templates" / filename,
        cwd / filename,
        cwd.parent / filename,
        cwd / "templates" / filename,
        cwd.parent / "templates" / filename,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue

    raise HTTPException(
        status_code=500,
        detail=(
            "NCC Japan IRB cover letter template not found. Expected CoverLetter_NCC_Japan_IRB.docx. "
            "Set NCC_JAPAN_IRB_COVER_LETTER_TEMPLATE to an absolute file path if needed."
        ),
    )


def _resolve_peter_mac_cover_letter_template_path() -> Path:
    """
    Resolve Peter Mac HREC cover letter template path from known locations.
    """
    env_path = (os.environ.get("PETER_MAC_HREC_COVER_LETTER_TEMPLATE") or "").strip()
    module_path = Path(__file__).resolve()
    cwd = Path.cwd()
    filename = "CoverLetter_PeterMac_HREC.docx"

    candidates = [
        Path(env_path) if env_path else None,
        module_path.parent / "templates" / filename,
        module_path.parents[2] / filename,
        module_path.parents[3] / filename,
        module_path.parents[2] / "templates" / filename,
        cwd / filename,
        cwd.parent / filename,
        cwd / "templates" / filename,
        cwd.parent / "templates" / filename,
    ]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            if candidate.exists():
                return candidate
        except Exception:
            continue

    raise HTTPException(
        status_code=500,
        detail=(
            "Peter Mac HREC cover letter template not found. Expected CoverLetter_PeterMac_HREC.docx. "
            "Set PETER_MAC_HREC_COVER_LETTER_TEMPLATE to an absolute file path if needed."
        ),
    )


def _build_tmc_cover_letter_field_mappings() -> Dict[str, str]:
    """
    Placeholder mapping source for TMC IEC cover letter.
    Only site_profile.* mappings are used as required.
    """
    return {
        "IEC_REF_NO": "site_profile.iec_ref_no",
        "DATE": "site_profile.generated_date",
        "HOSPITAL_NAME": "site_profile.hospital_name",
        "CITY": "site_profile.city",
        "STATE": "site_profile.state",
        "POSTAL_CODE": "site_profile.postal_code",
        "PI_NAME": "site_profile.pi_name",
        "PI_DESIGNATION": "site_profile.pi_designation",
        "DEPARTMENT": "site_profile.pi_department",
        "PI_PHONE": "site_profile.pi_phone",
        "PI_EMAIL": "site_profile.pi_email",
        "STUDY_TITLE": "site_profile.study_title",
        "STUDY_TYPE": "site_profile.study_type",
        "PRIMARY_OBJECTIVE": "site_profile.primary_objective",
        "NUM_HARD_COPIES": "site_profile.num_hard_copies",
        "SUBMISSION_MODE": "site_profile.submission_mode",
    }


def _build_md_anderson_cover_letter_field_mappings() -> Dict[str, str]:
    """
    Placeholder mapping source for MD Anderson IRB cover letter.
    """
    return {
        "FULL STUDY TITLE": "site_profile.study_title",
        "PROTOCOL #": "site_profile.protocol_number",
        "PROTOCOL NO. [PROTOCOL #]": "site_profile.protocol_number",
        "DATE": "site_profile.generated_date",
        "DEPARTMENT": "site_profile.pi_department",
        "PRINCIPAL INVESTIGATOR NAME": "site_profile.pi_name",
        "PI_PHONE_NUMBER": "site_profile.pi_phone",
        "PI_EMAIL": "site_profile.pi_email",
        "COORDINATOR NAME": "site_profile.site_coordinator_name",
        "COORDINATOR EMAIL": "site_profile.site_coordinator_email",
        "COORDINATOR_NAME": "site_profile.site_coordinator_name",
        "COORDINATOR_EMAIL": "site_profile.site_coordinator_email",
    }


def _build_oxford_cover_letter_field_mappings() -> Dict[str, str]:
    """
    Placeholder mapping source for South Central - Oxford REC cover letter.
    """
    return {
        "DATE": "site_profile.generated_date",
        "FULL STUDY TITLE": "site_profile.study_title",
        "IRAS ID": "site_profile.iras_id",
        "IRAS_ID": "site_profile.iras_id",
        "REC REF": "site_profile.rec_ref",
        "REC_REF": "site_profile.rec_ref",
        "SPONSOR NAME": "site_profile.primary_contracting_entity",
        "SPONSOR / CHIEF INVESTIGATOR INSTITUTION": "site_profile.primary_contracting_entity",
        "SPONSOR_CHIEF_INVESTIGATOR_INSTITUTION": "site_profile.primary_contracting_entity",
        "BRIEF OBJECTIVE": "site_profile.primary_objective",
        "STUDY DESIGN": "site_profile.study_type",
        "HOSPITAL_NAME": "site_profile.hospital_name",
        "DEPARTMENT": "site_profile.pi_department",
        "ADDRESS LINE 1": "site_profile.address_line_1",
        "ADDRESS_LINE_1": "site_profile.address_line_1",
        "CITY": "site_profile.city",
        "STATE": "site_profile.state",
        "POSTAL CODE": "site_profile.postal_code",
        "POSTAL_CODE": "site_profile.postal_code",
        "COUNTRY": "site_profile.country",
        "COORDINATOR_NAME": "site_profile.site_coordinator_name",
        "CONTACT_EMAIL": "site_profile.contact_email",
        "PRINCIPAL INVESTIGATOR NAME": "site_profile.pi_name",
        "PI_PHONE_NUMBER": "site_profile.pi_phone",
        "PI_EMAIL_ADDRESS": "site_profile.pi_email",
    }


def _build_heidelberg_cover_letter_field_mappings() -> Dict[str, str]:
    """
    Placeholder mapping source for Heidelberg University Ethics Committee cover letter.
    """
    return {
        "DATE": "site_profile.generated_date",
        "PRINCIPAL INVESTIGATOR NAME": "site_profile.pi_name",
        "HOSPITAL NAME": "site_profile.hospital_name",
        "HOSPITAL_NAME": "site_profile.hospital_name",
        "ADDRESS LINE 1": "site_profile.address_line_1",
        "ADDRESS_LINE_1": "site_profile.address_line_1",
        "CITY": "site_profile.city",
        "STATE": "site_profile.state",
        "PROTOCOL NUMBER": "site_profile.protocol_number",
        "INSERT NUMBER": "site_profile.protocol_number",
        "FULL PROTOCOL TITLE": "site_profile.study_title",
        "COORDINATOR EMAIL": "site_profile.site_coordinator_email",
        "COORDINATOR_EMAIL": "site_profile.site_coordinator_email",
        "SIGNATURE": "site_profile.pi_name",
        "TITLE": "site_profile.pi_designation",
        "DEPARTMENT": "site_profile.pi_department",
    }


def _build_chcams_cover_letter_field_mappings() -> Dict[str, str]:
    """
    Placeholder mapping source for NCC / CHCAMS China IEC cover letter.
    Supports {{...}} style placeholders used by the US-style generation flow.
    """
    return {
        "DATE": "site_profile.generated_date",
        "PRINCIPAL INVESTIGATOR NAME": "site_profile.pi_name",
        "DEPARTMENT": "site_profile.pi_department",
        "PROTOCOL NUMBER": "site_profile.protocol_number",
        "FULL PROTOCOL TITLE": "site_profile.study_title",
        "BRIEF OBJECTIVE": "site_profile.primary_objective",
        "TITLE": "site_profile.pi_designation",
    }


def _build_ncc_japan_cover_letter_field_mappings() -> Dict[str, str]:
    """
    Placeholder mapping source for National Cancer Center IRB (Japan) cover letter.
    """
    return {
        "DATE": "site_profile.generated_date",
        "PRINCIPAL INVESTIGATOR NAME": "site_profile.pi_name",
        "DEPARTMENT": "site_profile.pi_department",
        "PROTOCOL NUMBER": "site_profile.protocol_number",
        "FULL PROTOCOL TITLE": "site_profile.study_title",
        "COORDINATOR EMAIL": "site_profile.site_coordinator_email",
        "COORDINATOR_EMAIL": "site_profile.site_coordinator_email",
        "SIGNATURE": "site_profile.pi_name",
        "TITLE": "site_profile.pi_designation",
    }


def _build_peter_mac_cover_letter_field_mappings() -> Dict[str, str]:
    """
    Placeholder mapping source for Peter Mac HREC (Australia) cover letter.
    """
    return {
        "DATE": "site_profile.generated_date",
        "PRINCIPAL INVESTIGATOR NAME": "site_profile.pi_name",
        "ADDRESS LINE 1": "site_profile.address_line_1",
        "ADDRESS_LINE_1": "site_profile.address_line_1",
        "CITY": "site_profile.city",
        "STATE": "site_profile.state",
        "PROTOCOL NUMBER": "site_profile.protocol_number",
        "INSERT PROTOCOL NUMBER": "site_profile.protocol_number",
        "INSERT FULL PROTOCOL TITLE": "site_profile.study_title",
        "TITLE": "site_profile.pi_designation",
    }


def _build_smc_cover_letter_field_mappings() -> Dict[str, str]:
    """
    Placeholder mapping source for Samsung Medical Center IRB cover letter.
    """
    return {
        "HOSPITAL_NAME": "site_profile.hospital_name",
        "DEPARTMENT": "site_profile.pi_department",
        "ADDRESS_LINE_1": "site_profile.address_line_1",
        "CITY": "site_profile.city",
        "STATE": "site_profile.state",
        "POSTAL_CODE": "site_profile.postal_code",
        "COUNTRY": "site_profile.country",
        "DATE": "site_profile.generated_date",
        "FULL_STUDY_TITLE": "site_profile.study_title",
        "PROTOCOL_NO": "site_profile.protocol_number",
        "BRIEF_OBJECTIVE": "site_profile.primary_objective",
        "STUDY_DESIGN": "site_profile.study_type",
        "N": "site_profile.target_enrollment",
        "COORDINATOR_NAME": "site_profile.site_coordinator_name",
        "COORDINATOR_EMAIL": "site_profile.site_coordinator_email",
        "PRINCIPAL_INVESTIGATOR_NAME": "site_profile.pi_name",
        "PI_PHONE_NUMBER": "site_profile.pi_phone",
        "PI_EMAIL_ADDRESS": "site_profile.pi_email",
    }


def _placeholder_config_from_mapping(field_mappings: Dict[str, str]) -> Dict[str, Dict[str, bool]]:
    """
    Force plain text replacement (no content controls) to preserve original letter format.
    """
    return {key: {"editable": False} for key in field_mappings.keys()}


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_cover_letter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _insert_numbered_paragraph_after(paragraph, text: str):
    """Clone a numbered paragraph and insert it immediately after the anchor paragraph."""
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.text = text
    return new_para


def _remove_paragraph(paragraph) -> None:
    """Remove a paragraph element from the document tree."""
    p = paragraph._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _inject_required_documents_list(docx_path: Path, requirement_names: List[str]) -> None:
    """Replace numbered list entries in TMC cover letter with dynamic requirements."""
    if not requirement_names:
        return

    doc = Document(docx_path)

    label_paragraph = None
    numbered_after_label = []
    for p in _iter_cover_letter_paragraphs(doc):
        text = (p.text or "").strip()
        if label_paragraph is None and text.lower().startswith("list of enclosures"):
            label_paragraph = p
            continue
        if label_paragraph is not None:
            if re.match(r"^\d+\.\s+", text):
                numbered_after_label.append(p)
                continue
            # Stop once we pass the list section.
            if numbered_after_label:
                break

    if label_paragraph is None:
        # Fallback to old behavior if label not found.
        numbered = []
        for p in _iter_cover_letter_paragraphs(doc):
            text = (p.text or "").strip()
            if re.match(r"^\d+\.\s+", text):
                numbered.append(p)

        if not numbered:
            logger.info("No numbered paragraphs found in cover letter template; skipped dynamic list injection")
            return

        template_slots = len(numbered)
        limit = min(template_slots, len(requirement_names))
        for i in range(limit):
            numbered[i].text = f"{i + 1}. {requirement_names[i]}"

        for i in range(limit, template_slots):
            numbered[i].text = ""

        if len(requirement_names) > template_slots:
            anchor = numbered[-1]
            for i in range(template_slots, len(requirement_names)):
                anchor = _insert_numbered_paragraph_after(anchor, f"{i + 1}. {requirement_names[i]}")
    else:
        template_anchor = numbered_after_label[0] if numbered_after_label else label_paragraph
        for p in reversed(numbered_after_label):
            _remove_paragraph(p)

        anchor = label_paragraph
        for i, name in enumerate(requirement_names, start=1):
            # Same insertion whether or not a numbered template paragraph was
            # available — _insert_numbered_paragraph_after handles both.
            anchor = _insert_numbered_paragraph_after(anchor, f"{i}. {name}")

    doc.save(docx_path)


def _inject_required_documents_after_label(
    docx_path: Path,
    requirement_names: List[str],
    label_text: str,
) -> None:
    if not requirement_names:
        return

    doc = Document(docx_path)
    label_paragraph = None
    for p in _iter_cover_letter_paragraphs(doc):
        text = (p.text or "").strip().lower()
        if text.startswith(label_text.lower()):
            label_paragraph = p
            break

    if label_paragraph is None:
        logger.info("Cover letter label paragraph not found; skipped list injection")
        return

    anchor = label_paragraph
    for idx, name in enumerate(requirement_names, start=1):
        anchor = _insert_numbered_paragraph_after(anchor, f"{idx}. {name}")

    doc.save(docx_path)
    logger.info(
        "Injected %s dynamic required-document entries into cover letter",
        len(requirement_names),
    )


def _replace_square_bracket_tokens_in_docx(docx_path: Path, replacements: Dict[str, str]) -> None:
    """Replace [token] style placeholders used by some legacy cover letter templates."""
    if not replacements:
        return

    doc = Document(docx_path)

    for paragraph in _iter_cover_letter_paragraphs(doc):
        if not paragraph.text:
            continue
        for token, value in replacements.items():
            if value is None or not str(value).strip():
                continue
            if not token or token not in paragraph.text:
                continue
            replaced_in_runs = False
            for run in paragraph.runs:
                if token in run.text:
                    run.text = run.text.replace(token, value)
                    replaced_in_runs = True
            # Fallback for placeholders split across runs.
            if not replaced_in_runs and token in paragraph.text:
                paragraph.text = paragraph.text.replace(token, value)

    doc.save(docx_path)


async def _resolve_tmc_required_documents(
    db: AsyncSession,
    site: Site,
    preferred_names: Optional[List[str]] = None,
) -> Tuple[List[str], str]:
    """Return requirement names and source label (ui|irb|fallback)."""
    if preferred_names:
        cleaned = []
        for name in preferred_names:
            val = str(name or "").strip()
            if not val:
                continue
            cleaned.append(val)
        if cleaned:
            return cleaned, "ui"

    try:
        mapping_result = await db.execute(
            select(SiteIRBMapping).where(SiteIRBMapping.site_id == site.id)
        )
        mapping = mapping_result.scalar_one_or_none()
        if mapping:
            irb_result = await db.execute(select(IRB).where(IRB.id == int(mapping.irb_id)))
            irb = irb_result.scalar_one_or_none()
            if irb:
                reqs = await resolve_requirements_for_irb(db, irb)
                names = [
                    str(r.get("name", "")).strip()
                    for r in reqs
                    if isinstance(r, dict) and str(r.get("name", "")).strip()
                ]
                if names:
                    return names, "irb"
    except Exception:
        logger.warning("Failed to resolve dynamic TMC required-document list; using defaults", exc_info=True)

    return list(_FALLBACK_TMC_REQUIRED_DOCUMENTS), "fallback"


def _get_azure_connection_string() -> str:
    # Current backend settings use `azure_storage_connection_string`.
    # The legacy ISF helper uses `AZURE_STORAGE_CONNECTION_STRING` env var.
    return (
        getattr(settings, "azure_storage_connection_string", None)
        or os.environ.get("AZURE_STORAGE_CONNECTION_STRING", "")
    )


def _get_azure_sitepackages_container() -> str:
    return (
        os.environ.get("AZURE_SITEPACKAGES_CONTAINER")
        or getattr(settings, "azure_templates_container_name", None)
        or "site-packages"
    )


def _parse_connection_string(conn_str: str) -> Dict[str, str]:
    parts: Dict[str, str] = {}
    for part in conn_str.split(";"):
        if "=" in part:
            k, _, v = part.partition("=")
            parts[k.strip()] = v.strip()
    return parts


def _is_uuid_like(value: str) -> bool:
    try:
        UUID(str(value))
        return True
    except Exception:
        return False


async def _resolve_study_uuid(db: AsyncSession, study_value: str) -> UUID:
    # Accept either:
    # - UUID primary key (Study.id)
    # - external study_id string (Study.study_id)
    if _is_uuid_like(study_value):
        study_uuid = UUID(str(study_value))
        result = await db.execute(select(Study).where(Study.id == study_uuid))
        study = result.scalar_one_or_none()
        if not study:
            raise HTTPException(status_code=404, detail="Study not found")
        return study.id

    result = await db.execute(select(Study).where(Study.study_id == study_value))
    study = result.scalar_one_or_none()
    if not study:
        raise HTTPException(status_code=404, detail="Study not found")
    return study.id


async def _resolve_site_uuid(db: AsyncSession, site_value: str) -> UUID:
    if _is_uuid_like(site_value):
        site_uuid = UUID(str(site_value))
        result = await db.execute(select(Site).where(Site.id == site_uuid))
        site = result.scalar_one_or_none()
        if not site:
            raise HTTPException(status_code=404, detail="Site not found")
        return site.id

    result = await db.execute(select(Site).where(Site.site_id == site_value))
    site = result.scalar_one_or_none()
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")
    return site.id


def _serialize_contact_person(value: Any) -> Any:
    # Keep it JSON-serializable; allow empty object.
    if value is None:
        return {}
    return value


def _uploaded_requirement_names_from_package_documents(documents: Any) -> Set[str]:
    """Requirement `name` values that have at least one stored file URL."""
    names: Set[str] = set()
    if not isinstance(documents, list):
        return names
    for d in documents:
        if not isinstance(d, dict):
            continue
        if not (d.get("fileUrl") or d.get("file_url") or d.get("sourceUrl") or d.get("source_url")):
            continue
        name = d.get("name")
        if name is not None and str(name).strip():
            names.add(str(name).strip())
    return names


def _compute_site_package_progress_percent(requirements: List[Dict[str, Any]], documents: Any) -> int:
    """
    Progress = mandatory requirement slots satisfied / total mandatory slots (0–100).
    Matches the wizard rule: all mandatory documents must be present before proceeding.
    """
    uploaded = _uploaded_requirement_names_from_package_documents(documents)
    uploaded_lower = {n.lower() for n in uploaded}

    def _has_slot(name: str) -> bool:
        n = str(name).strip()
        return n in uploaded or n.lower() in uploaded_lower

    mandatory_names = [
        str(r.get("name", "")).strip()
        for r in requirements
        if str(r.get("requirement", "")).strip().lower() == "mandatory" and str(r.get("name", "")).strip()
    ]
    if mandatory_names:
        total = len(mandatory_names)
        satisfied = sum(1 for n in mandatory_names if _has_slot(n))
        return min(100, int(round((satisfied / total) * 100))) if total else 0
    # No explicit mandatory rows: use every listed requirement as the bar.
    all_names = [str(r.get("name", "")).strip() for r in requirements if str(r.get("name", "")).strip()]
    if not all_names:
        return 100 if uploaded else 0
    total = len(all_names)
    satisfied = sum(1 for n in all_names if _has_slot(n))
    return min(100, int(round((satisfied / total) * 100))) if total else 0


def _enforce_document_limits(requirements: List[Dict[str, Any]], documents: Any) -> None:
    """Validate per-requirement max file limits (if configured)."""
    if not isinstance(documents, list) or not requirements:
        # Still enforce known amendment/modification document constraints even
        # when IRB requirements are not available.
        if not isinstance(documents, list):
            return

    max_by_name: Dict[str, int] = {}
    max_by_name_lower: Dict[str, int] = {}
    max_by_type: Dict[str, int] = {}
    for r in requirements:
        if not isinstance(r, dict):
            continue
        name = str(r.get("name", "")).strip()
        type_key = str(r.get("type", "")).strip().lower()
        max_files = r.get("max_files")
        if (not name and not type_key) or max_files is None:
            continue
        try:
            max_int = int(max_files)
        except (TypeError, ValueError):
            continue
        if max_int < 1:
            continue
        if name:
            max_by_name[name] = max_int
            max_by_name_lower[name.lower()] = max_int
        if type_key:
            max_by_type[type_key] = max_int

    # Enforce amendment/modification limits even if IRB metadata does not
    # include these workflow-specific document definitions.
    for name_lower, max_int in _AMENDMENT_DOC_MAX_BY_NAME_LOWER.items():
        max_by_name_lower.setdefault(name_lower, max_int)
    for type_key, max_int in _AMENDMENT_DOC_MAX_BY_TYPE.items():
        max_by_type.setdefault(type_key, max_int)
    for name_lower, max_int in _CONTINUING_REVIEW_DOC_MAX_BY_NAME_LOWER.items():
        max_by_name_lower.setdefault(name_lower, max_int)
    for type_key, max_int in _CONTINUING_REVIEW_DOC_MAX_BY_TYPE.items():
        max_by_type.setdefault(type_key, max_int)
    for name_lower, max_int in _SAFETY_REPORT_DOC_MAX_BY_NAME_LOWER.items():
        max_by_name_lower.setdefault(name_lower, max_int)
    for type_key, max_int in _SAFETY_REPORT_DOC_MAX_BY_TYPE.items():
        max_by_type.setdefault(type_key, max_int)
    for name_lower, max_int in _STUDY_CLOSURE_DOC_MAX_BY_NAME_LOWER.items():
        max_by_name_lower.setdefault(name_lower, max_int)
    for type_key, max_int in _STUDY_CLOSURE_DOC_MAX_BY_TYPE.items():
        max_by_type.setdefault(type_key, max_int)
    for name_lower, max_int in _PROTOCOL_DEVIATION_DOC_MAX_BY_NAME_LOWER.items():
        max_by_name_lower.setdefault(name_lower, max_int)
    for type_key, max_int in _PROTOCOL_DEVIATION_DOC_MAX_BY_TYPE.items():
        max_by_type.setdefault(type_key, max_int)
    for name_lower, max_int in _NOTIFICATION_DOC_MAX_BY_NAME_LOWER.items():
        max_by_name_lower.setdefault(name_lower, max_int)
    for type_key, max_int in _NOTIFICATION_DOC_MAX_BY_TYPE.items():
        max_by_type.setdefault(type_key, max_int)

    if not max_by_name_lower and not max_by_type:
        return

    counts: Dict[str, int] = {}
    counts_by_type: Dict[str, int] = {}
    for d in documents:
        if not isinstance(d, dict):
            continue
        name = str(d.get("name", "")).strip()
        type_key = str(d.get("type", "")).strip().lower()
        if not name:
            if type_key:
                counts_by_type[type_key] = counts_by_type.get(type_key, 0) + 1
            continue
        counts[name] = counts.get(name, 0) + 1
        if type_key:
            counts_by_type[type_key] = counts_by_type.get(type_key, 0) + 1

    for name, count in counts.items():
        max_allowed = max_by_name.get(name)
        if max_allowed is None:
            max_allowed = max_by_name_lower.get(name.lower())
        if max_allowed is not None and count > max_allowed:
            raise HTTPException(
                status_code=400,
                detail=f"Only {max_allowed} file(s) allowed for '{name}'.",
            )

    for type_key, count in counts_by_type.items():
        max_allowed = max_by_type.get(type_key)
        if max_allowed is not None and count > max_allowed:
            raise HTTPException(
                status_code=400,
                detail=f"Only {max_allowed} file(s) allowed for document type '{type_key}'.",
            )


def _warn_incomplete_requirement_metadata(
    requirements: List[Dict[str, Any]],
    *,
    irb_id: Optional[int],
    context: str,
) -> None:
    """Warn when requirement rows are missing metadata needed for strict upload enforcement."""
    if not isinstance(requirements, list) or not requirements:
        return

    missing_type: List[str] = []
    missing_max_files: List[str] = []

    for r in requirements:
        if not isinstance(r, dict):
            continue
        name = str(r.get("name", "")).strip() or "<unnamed>"
        r_type = str(r.get("type", "")).strip()
        r_max = r.get("max_files")

        if not r_type:
            missing_type.append(name)
        if r_max is None:
            missing_max_files.append(name)

    if missing_type or missing_max_files:
        logger.warning(
            "IRB requirement metadata incomplete in %s (irb_id=%s): missing_type=%s missing_max_files=%s",
            context,
            irb_id,
            len(missing_type),
            len(missing_max_files),
        )


def _legacy_progress_percent_from_doc_count(documents: Any) -> int:
    """Previous UI heuristic (docs / 10); kept only when IRB is unknown."""
    if not isinstance(documents, list):
        return 0
    n = len(documents)
    return min(100, max(0, int(round((n / 10) * 100))))


def _infer_submission_type(pkg: SitePackage) -> str:
    # 1) If model ever has this attribute populated, prefer it.
    direct = getattr(pkg, "submissionType", None)
    if isinstance(direct, str) and direct.strip():
        return direct.strip()

    # 2) Check audit trail details for create/update payloads.
    trail = list(getattr(pkg, "auditTrail", None) or [])
    for item in reversed(trail):
        if not isinstance(item, dict):
            continue
        details = item.get("details")
        if not isinstance(details, dict):
            continue
        st = details.get("submissionType") or details.get("submission_type")
        if isinstance(st, str) and st.strip():
            return st.strip()

    # 3) Infer from known workflow-specific document names/types.
    docs = list(getattr(pkg, "documents", None) or [])
    for d in docs:
        if not isinstance(d, dict):
            continue
        t = str(d.get("type") or "").strip().lower()
        n = str(d.get("name") or "").strip().lower()
        if t in _NOTIFICATION_DOC_MAX_BY_TYPE or n in _NOTIFICATION_DOC_TYPE_BY_NAME_LOWER:
            return "Notification"
        if t in _PROTOCOL_DEVIATION_DOC_MAX_BY_TYPE or n in _PROTOCOL_DEVIATION_DOC_TYPE_BY_NAME_LOWER:
            return "Protocol Deviation / Violation Package"
        if t in _STUDY_CLOSURE_DOC_MAX_BY_TYPE or n in _STUDY_CLOSURE_DOC_TYPE_BY_NAME_LOWER:
            return "Study Closure Package"
        if t in _SAFETY_REPORT_DOC_MAX_BY_TYPE or n in _SAFETY_REPORT_DOC_TYPE_BY_NAME_LOWER:
            return "Safety Report Submission / Package"
        if t in _CONTINUING_REVIEW_DOC_MAX_BY_TYPE or n in _CONTINUING_REVIEW_DOC_TYPE_BY_NAME_LOWER:
            return "Continuing Review / (Annual Renewal)"
        if t in _AMENDMENT_DOC_MAX_BY_TYPE or n in _AMENDMENT_DOC_TYPE_BY_NAME_LOWER:
            return "Amendment / Modification Package"

    return "Initial Submission"


async def _progress_for_site_package(db: AsyncSession, pkg: SitePackage) -> int:
    submission_type = _infer_submission_type(pkg)
    workflow_reqs = await _workflow_requirements_for_submission_type_async(submission_type)
    if workflow_reqs:
        return _compute_site_package_progress_percent(workflow_reqs, pkg.documents or [])
    if submission_type == "Notification":
        # Notification has no predefined requirement list.
        return _compute_site_package_progress_percent([], pkg.documents or [])

    if pkg.irb_id is None:
        return _legacy_progress_percent_from_doc_count(pkg.documents or [])
    result = await db.execute(select(IRB).where(IRB.id == int(pkg.irb_id)))
    irb = result.scalar_one_or_none()
    if not irb:
        return _legacy_progress_percent_from_doc_count(pkg.documents or [])
    reqs = await resolve_requirements_for_irb(db, irb)
    return _compute_site_package_progress_percent(reqs, pkg.documents or [])


def _serialize_site_package_single(
    pkg: SitePackage,
    progress_percent: Optional[int] = None,
) -> Dict[str, Any]:
    # Node/Mongoose GET/PATCH responses return references as raw IDs
    # (not populated), so we keep `study`/`site` as IDs and `createdBy`
    # as the stored string.
    out: Dict[str, Any] = {
        "_id": str(pkg.id),
        "study": str(pkg.study_id),
        "site": str(pkg.site_id) if pkg.site_id else None,
        "irb_id": int(getattr(pkg, "irb_id", None)) if getattr(pkg, "irb_id", None) is not None else None,
        "ethicsBoard": pkg.ethicsBoard,
        "packageName": pkg.packageName,
        "description": pkg.description,
        "priority": pkg.priority,
        "submissionType": _infer_submission_type(pkg),
        "expectedSubmissionDate": pkg.expectedSubmissionDate,
        "contactPerson": _serialize_contact_person(pkg.contactPerson),
        "notes": pkg.notes,
        "documents": pkg.documents or [],
        "createdBy": pkg.createdBy,
        "lastUpdated": pkg.lastUpdated or pkg.updatedAt,
        "auditTrail": pkg.auditTrail or [],
        "status": pkg.status,
        "isDeleted": pkg.isDeleted,
        "createdAt": pkg.createdAt,
        "updatedAt": pkg.updatedAt,
    }
    if progress_percent is not None:
        out["progress"] = progress_percent
    return out


def _serialize_site_package_list(
    pkg: SitePackage,
    study_map: Dict[str, Study],
    site_map: Dict[str, Site],
    user_map: Dict[str, User],
    progress_percent: Optional[int] = None,
) -> Dict[str, Any]:
    # Node/Mongoose list response populates:
    # - study -> studyId/title/name
    # - site  -> siteId/name/siteCode
    # - createdBy -> name/email
    study = study_map.get(str(pkg.study_id))
    site = site_map.get(str(pkg.site_id)) if pkg.site_id else None
    user = user_map.get(pkg.createdBy) if pkg.createdBy else None

    study_obj = (
        {
            "_id": str(study.id),
            "studyId": study.study_id,
            "title": study.name,
            "name": study.name,
        }
        if study
        else str(pkg.study_id)
    )
    site_obj = (
        {
            "_id": str(site.id),
            "siteId": site.site_id,
            "name": site.name,
            "siteCode": site.code,
        }
        if site
        else (str(pkg.site_id) if pkg.site_id else None)
    )
    created_by_obj = (
        {"_id": user.user_id, "name": user.name, "email": user.email}
        if user
        else pkg.createdBy
    )

    out: Dict[str, Any] = {
        "_id": str(pkg.id),
        "study": study_obj,
        "site": site_obj,
        "ethicsBoard": pkg.ethicsBoard,
        "packageName": pkg.packageName,
        "description": pkg.description,
        "priority": pkg.priority,
        "submissionType": _infer_submission_type(pkg),
        "expectedSubmissionDate": pkg.expectedSubmissionDate,
        "contactPerson": _serialize_contact_person(pkg.contactPerson),
        "notes": pkg.notes,
        "documents": pkg.documents or [],
        "createdBy": created_by_obj,
        "lastUpdated": pkg.lastUpdated or pkg.updatedAt,
        "auditTrail": pkg.auditTrail or [],
        "status": pkg.status,
        "isDeleted": pkg.isDeleted,
        "createdAt": pkg.createdAt,
        "updatedAt": pkg.updatedAt,
    }
    if progress_percent is not None:
        out["progress"] = progress_percent
    return out


async def create_site_package(
    payload: Any,
    db: AsyncSession,
    current_user: Optional[dict],
) -> Dict[str, Any]:
    study_uuid = await _resolve_study_uuid(db, payload.study)
    site_uuid = await _resolve_site_uuid(db, payload.site)

    # Prefer IRB-derived Ethics Board if `irb_id` is provided.
    ethics_board_value = getattr(payload, "ethicsBoard", None)
    irb_id_value = getattr(payload, "irb_id", None)
    irb: Optional[IRB] = None
    if irb_id_value is not None:
        result = await db.execute(select(IRB).where(IRB.id == int(irb_id_value)))
        irb = result.scalar_one_or_none()
        if not irb:
            raise HTTPException(status_code=404, detail="IRB not found")
        ethics_board_value = irb.name

    if not ethics_board_value:
        raise HTTPException(
            status_code=400,
            detail="Ethics board is required (provide `irb_id` or `ethicsBoard`)",
        )

    if irb is not None:
        requirements = await resolve_requirements_for_irb(db, irb)
        _warn_incomplete_requirement_metadata(requirements, irb_id=int(irb.id), context="create_site_package")
        _enforce_document_limits(requirements, payload.documents or [])

    user_id = (current_user or {}).get("user_id") or payload.createdBy

    now = _now_utc()
    audit_trail: List[Dict[str, Any]] = [
        {
            "action": "created",
            **({"user": user_id} if user_id else {}),
            # mode="json" keeps datetimes as ISO strings so the JSON column can serialize.
            "details": payload.model_dump(mode="json") if hasattr(payload, "model_dump") else payload,
            "timestamp": now.isoformat(),
        }
    ]

    contact_person_value = payload.contactPerson
    if hasattr(contact_person_value, "model_dump"):
        contact_person_value = contact_person_value.model_dump(exclude_none=True)

    pkg = SitePackage(
        study_id=study_uuid,
        site_id=site_uuid,
        irb_id=int(irb_id_value) if irb_id_value is not None else None,
        ethicsBoard=ethics_board_value,
        packageName=payload.packageName,
        description=payload.description,
        priority=payload.priority or "Medium",
        expectedSubmissionDate=payload.expectedSubmissionDate,
        contactPerson=contact_person_value or {},
        notes=payload.notes,
        documents=payload.documents or [],
        createdBy=user_id,
        lastUpdated=now,
        auditTrail=audit_trail,
        status="Draft",
        isDeleted=False,
    )

    db.add(pkg)
    await db.commit()
    await db.refresh(pkg)
    prog = await _progress_for_site_package(db, pkg)
    return _serialize_site_package_single(pkg, progress_percent=prog)


async def list_site_packages(
    query: Any,
    db: AsyncSession,
) -> Dict[str, Any]:
    page = int(query.page or 1)
    limit = int(query.limit or 20)
    skip = (page - 1) * limit

    if not getattr(query, "study", None) or not getattr(query, "site", None):
        raise HTTPException(
            status_code=400,
            detail="Both `study` and `site` query parameters are required",
        )

    study_uuid = await _resolve_study_uuid(db, query.study)
    site_uuid = await _resolve_site_uuid(db, query.site)

    filters = [
        SitePackage.isDeleted.is_(False),
        SitePackage.study_id == study_uuid,
        SitePackage.site_id == site_uuid,
    ]

    if query.status:
        filters.append(SitePackage.status == query.status)

    where_clause = and_(*filters) if filters else SitePackage.isDeleted.is_(False)

    total_q = await db.execute(select(func.count()).select_from(SitePackage).where(where_clause))
    total = int(total_q.scalar_one())

    items_q = await db.execute(
        select(SitePackage)
        .where(where_clause)
        .order_by(SitePackage.updatedAt.desc())
        .offset(skip)
        .limit(limit)
    )
    items = items_q.scalars().all()

    study_ids = {p.study_id for p in items}
    site_ids = {p.site_id for p in items if p.site_id}
    created_by_ids = {p.createdBy for p in items if p.createdBy}

    study_map: Dict[str, Study] = {}
    if study_ids:
        study_q = await db.execute(select(Study).where(Study.id.in_(study_ids)))
        study_map = {str(s.id): s for s in study_q.scalars().all()}

    site_map: Dict[str, Site] = {}
    if site_ids:
        site_q = await db.execute(select(Site).where(Site.id.in_(site_ids)))
        site_map = {str(s.id): s for s in site_q.scalars().all()}

    user_map: Dict[str, User] = {}
    if created_by_ids:
        user_q = await db.execute(select(User).where(User.user_id.in_(created_by_ids)))
        user_map = {u.user_id: u for u in user_q.scalars().all()}

    irb_ids = {p.irb_id for p in items if getattr(p, "irb_id", None) is not None}
    req_by_irb_id: Dict[int, List[Dict[str, Any]]] = {}
    if irb_ids:
        irb_q = await db.execute(select(IRB).where(IRB.id.in_(irb_ids)))
        for irb in irb_q.scalars().all():
            req_by_irb_id[int(irb.id)] = await resolve_requirements_for_irb(db, irb)

    serialized = []
    for p in items:
        progress_pct: Optional[int] = None
        submission_type = _infer_submission_type(p)
        workflow_reqs = await _workflow_requirements_for_submission_type_async(submission_type)
        if workflow_reqs:
            progress_pct = _compute_site_package_progress_percent(
                workflow_reqs,
                p.documents or [],
            )
        elif submission_type == "Notification":
            progress_pct = _compute_site_package_progress_percent(
                [],
                p.documents or [],
            )
        elif p.irb_id is not None and int(p.irb_id) in req_by_irb_id:
            progress_pct = _compute_site_package_progress_percent(
                req_by_irb_id[int(p.irb_id)],
                p.documents or [],
            )
        else:
            progress_pct = _legacy_progress_percent_from_doc_count(p.documents or [])
        serialized.append(
            _serialize_site_package_list(
                pkg=p,
                study_map=study_map,
                site_map=site_map,
                user_map=user_map,
                progress_percent=progress_pct,
            )
        )

    pages = int((total + limit - 1) // limit) if limit else 1
    return {
        "success": True,
        "data": serialized,
        "pagination": {"page": page, "limit": limit, "total": total, "pages": pages},
    }


def _require_study_site_for_scope(
    study_value: Optional[str],
    site_value: Optional[str],
) -> None:
    if not study_value or not site_value:
        raise HTTPException(
            status_code=400,
            detail="Both `study` and `site` query parameters are required",
        )


def _assert_package_in_scope(
    pkg: SitePackage,
    study_uuid: UUID,
    site_uuid: UUID,
) -> None:
    if pkg.study_id != study_uuid or pkg.site_id != site_uuid:
        raise HTTPException(status_code=404, detail="Package not found")


async def get_site_package_by_id(
    pkg_id: str,
    db: AsyncSession,
    study_value: Optional[str] = None,
    site_value: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    _require_study_site_for_scope(study_value, site_value)
    study_uuid = await _resolve_study_uuid(db, study_value or "")
    site_uuid = await _resolve_site_uuid(db, site_value or "")

    try:
        pkg_uuid = UUID(str(pkg_id))
    except Exception:
        # Keep behavior close to "not found"
        return None

    result = await db.execute(
        select(SitePackage).where(SitePackage.id == pkg_uuid).where(SitePackage.isDeleted.is_(False))
    )
    pkg = result.scalar_one_or_none()
    if not pkg:
        return None
    _assert_package_in_scope(pkg, study_uuid, site_uuid)
    prog = await _progress_for_site_package(db, pkg)
    return _serialize_site_package_single(pkg, progress_percent=prog)


async def delete_site_package(
    pkg_id: str,
    db: AsyncSession,
    study_value: Optional[str] = None,
    site_value: Optional[str] = None,
) -> bool:
    _require_study_site_for_scope(study_value, site_value)
    study_uuid = await _resolve_study_uuid(db, study_value or "")
    site_uuid = await _resolve_site_uuid(db, site_value or "")

    try:
        pkg_uuid = UUID(str(pkg_id))
    except Exception:
        return False

    result = await db.execute(select(SitePackage).where(SitePackage.id == pkg_uuid))
    pkg = result.scalar_one_or_none()
    if not pkg or pkg.isDeleted:
        return False

    if pkg.study_id != study_uuid or pkg.site_id != site_uuid:
        return False

    pkg.isDeleted = True
    await db.commit()
    return True


async def update_site_package(
    pkg_id: str,
    updates: Any,
    db: AsyncSession,
    current_user: Optional[dict],
    study_value: Optional[str] = None,
    site_value: Optional[str] = None,
) -> Dict[str, Any]:
    _require_study_site_for_scope(study_value, site_value)
    study_uuid = await _resolve_study_uuid(db, study_value or "")
    site_uuid = await _resolve_site_uuid(db, site_value or "")

    try:
        pkg_uuid = UUID(str(pkg_id))
    except Exception:
        raise HTTPException(status_code=404, detail="Package not found")

    result = await db.execute(select(SitePackage).where(SitePackage.id == pkg_uuid).where(SitePackage.isDeleted.is_(False)))
    pkg = result.scalar_one_or_none()
    if not pkg:
        raise HTTPException(status_code=404, detail="Package not found")

    _assert_package_in_scope(pkg, study_uuid, site_uuid)

    updates_dict = updates.model_dump(exclude_unset=True) if hasattr(updates, "model_dump") else (updates or {})
    # JSON-safe copy for the audit trail (datetimes as ISO strings).
    updates_json = (
        updates.model_dump(exclude_unset=True, mode="json")
        if hasattr(updates, "model_dump")
        else (updates or {})
    )
    allowed = {
        "ethicsBoard",
        "packageName",
        "description",
        "priority",
        "expectedSubmissionDate",
        "contactPerson",
        "notes",
        "status",
        "documents",
        "irb_id",
    }

    for key, value in updates_dict.items():
        if key in allowed and value is not None:
            if hasattr(value, "model_dump"):
                value = value.model_dump(exclude_none=True)
            if key == "irb_id":
                setattr(pkg, "irb_id", int(value))
                continue
            if key == "documents":
                setattr(pkg, "documents", list(value))
                continue
            setattr(pkg, key, value)

    if getattr(pkg, "irb_id", None) is not None:
        irb_result = await db.execute(select(IRB).where(IRB.id == int(pkg.irb_id)))
        irb = irb_result.scalar_one_or_none()
        if irb:
            requirements = await resolve_requirements_for_irb(db, irb)
            _warn_incomplete_requirement_metadata(requirements, irb_id=int(irb.id), context="update_site_package")
            _enforce_document_limits(requirements, pkg.documents or [])

    now = _now_utc()
    pkg.lastUpdated = now

    user_id = (current_user or {}).get("user_id") or pkg.createdBy
    audit_item = {
        "action": "updated",
        **({"user": user_id} if user_id else {}),
        "details": updates_json,
        "timestamp": now.isoformat(),
    }
    pkg.auditTrail = list(pkg.auditTrail or []) + [audit_item]

    await db.commit()
    await db.refresh(pkg)
    prog = await _progress_for_site_package(db, pkg)
    return _serialize_site_package_single(pkg, progress_percent=prog)


async def upload_site_packages_document(
    file: UploadFile,
    study_id: str,
    site_id: str,
    document_type: str,
    db: AsyncSession,  # not used, parity with signature
) -> Dict[str, Any]:
    if not file:
        raise HTTPException(status_code=400, detail="No file")

    content = await file.read()
    filename = file.filename or "upload"
    content_type = file.content_type or "application/octet-stream"
    container = "crm-templates"

    doc_type_clean = (document_type or "unknown").strip().lower()
    study_clean = (study_id or "unknown").strip()
    site_clean = (site_id or "unknown").strip()
    from uuid import uuid4

    unique_name = f"{uuid4()}_{filename}"
    blob_path = f"site-packages/{study_clean}/{site_clean}/{doc_type_clean}/{unique_name}"

    url, size = _upload_bytes_to_azure(content, container=container, blob_name=blob_path, content_type=content_type)

    return {
        "name": filename,
        "storedName": unique_name,
        "fileUrl": url,
        "fileSize": size,
        "mimeType": content_type,
        "blobPath": blob_path,
        "container": container,
    }

def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _clean_protocol_text(value: Any) -> str:
    raw = _safe_text(value)
    if not raw:
        return ""
    # Remove PDF extraction artifacts: hard line breaks, hyphenated wraps, repeated spaces.
    cleaned = raw.replace("\r", "\n")
    cleaned = re.sub(r"-\s*\n\s*", "", cleaned)
    cleaned = re.sub(r"\n{2,}", "\n", cleaned)
    cleaned = re.sub(r"[^\S\n]*+\n\s*+", " ", cleaned)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def _build_standard_irb_payload(
    *,
    profile_data: Dict[str, Any],
    study: Study,
    site: Site,
    irb: Optional[IRB],
    required_docs: List[str],
    now: datetime,
    irb_admin_info: Optional[IRBAdministrativeInfo] = None,
    protocol_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    institution_name = _safe_text(profile_data.get("hospital_name")) or _safe_text(getattr(site, "site_name", None))
    city = _safe_text(profile_data.get("city"))
    state = _safe_text(profile_data.get("state"))
    postal_code = _safe_text(profile_data.get("postal_code"))
    city_state_zip = ", ".join([x for x in [city, state, postal_code] if x])

    admin_country = _safe_text(getattr(irb_admin_info, "country", None))
    # Regulatory defaults should follow IRB jurisdiction first.
    country = admin_country or _safe_text(profile_data.get("country")) or "India"
    committee_name = _safe_text(getattr(irb, "name", None)) or "Institutional Ethics Committee"

    # IRB Chairperson: prefer the value saved on the IRB Administrative form.
    admin_chair_name = _safe_text(getattr(irb_admin_info, "chair_name", None))
    chairperson_name = admin_chair_name or "IRB Chairperson"

    # "Primary Contact" block in the closing should still be the study-side
    # coordinator/PI (not the IRB chair), since they handle queries.
    coordinator_name = (
        _safe_text(profile_data.get("site_coordinator_name"))
        or _safe_text(profile_data.get("pi_name"))
        or chairperson_name
    )
    contact_email = (
        _safe_text(profile_data.get("site_coordinator_email"))
        or _safe_text(profile_data.get("pi_email"))
        or _safe_text(profile_data.get("contact_email"))
    )
    contact_phone = _safe_text(profile_data.get("pi_phone")) or _safe_text(profile_data.get("site_phone"))
    # Header / "on behalf of" / signature block should display the Hospital Name.
    sponsor_name = institution_name
    protocol_number = (
        _clean_protocol_text((protocol_metadata or {}).get("protocol_number"))
        or _safe_text(profile_data.get("protocol_number"))
        or _safe_text(getattr(study, "study_id", None))
    )
    study_title = (
        _clean_protocol_text((protocol_metadata or {}).get("study_title"))
        or _safe_text(profile_data.get("study_title"))
        or _safe_text(getattr(study, "name", None))
    )
    study_description = (
        _clean_protocol_text((protocol_metadata or {}).get("study_description"))
        or _safe_text(profile_data.get("primary_objective"))
    )
    submission_version = (
        _clean_protocol_text((protocol_metadata or {}).get("version"))
        or _safe_text(profile_data.get("protocol_version"))
        or "v1.0"
    )
    pi_name = _safe_text(profile_data.get("pi_name"))
    pi_designation = _safe_text(profile_data.get("pi_designation"))
    pi_department = _safe_text(profile_data.get("pi_department"))
    pi_title = " - ".join([x for x in [pi_designation, pi_department] if x]) or pi_designation or "Principal Investigator"

    enclosed_document_list = [
        {"title": _safe_text(name), "version": "", "format": "PDF"}
        for name in required_docs
        if _safe_text(name)
    ]

    return {
        "submission_date": now.strftime("%d %B %Y"),
        "irb_committee_name": committee_name,
        "irb_chair_name": chairperson_name,
        "institution_name": institution_name,
        "institution_address_line1": _safe_text(profile_data.get("address_line_1")),
        "institution_address_line2": _safe_text(profile_data.get("address_line_2")),
        "city_state_zip": city_state_zip,
        "country": country,
        "salutation": "Chairperson, Institutional Ethics Committee",
        "sponsor_institution_name": sponsor_name,
        "protocol_number": protocol_number,
        "study_title": study_title,
        "study_description": study_description,
        "submission_type": _safe_text(profile_data.get("submission_mode")) or "Initial Submission",
        "study_site_name": institution_name,
        "review_type": _safe_text(profile_data.get("review_type")) or "Full Board Review",
        "submission_version": submission_version,
        "primary_contact_name": coordinator_name,
        "primary_contact_email": contact_email,
        "primary_contact_phone": contact_phone,
        "enclosed_document_list": enclosed_document_list,
        "pi_name": pi_name,
        "pi_title": pi_title,
        "pi_department": pi_department,
        "signature_stamp": "[Auto-generated by CRM]",
        "co_pi_name": "",
        "co_pi_title": "",
        "system_submission_id": f"{_safe_text(getattr(study, 'study_id', None))}-{_safe_text(getattr(site, 'site_id', None))}-{now.strftime('%Y%m%d')}",
        "irb_assigned_ref": _safe_text(profile_data.get("iec_ref_no")),
    }


def _generate_standard_irb_cover_letter_docx(payload: Dict[str, Any]) -> bytes:
    """
    Build the standardized IRB cover letter .docx in-process using python-docx.

    The layout/content mirrors the Node generator in
    Frontend-CRM/scripts/irb-cover-letter, but runs without a Node runtime so it
    works in any backend deployment (including containers without Node.js).
    """
    from app.modules.site_packages.cover_letter_generator import (
        generate_cover_letter_docx_bytes,
    )

    try:
        return generate_cover_letter_docx_bytes(payload)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Standard IRB cover letter generation failed: {exc}",
        ) from exc


_NOTIFICATION_DEFAULT_ATTACHMENTS: List[str] = [
    "Updated Form FDA 1572 / Regulatory Statement of Investigator",
    "Current Curriculum Vitae (CV) of {{new_pi}} (Signed & Dated)",
    "Active Medical License (License No: {{License_Number}} | Expiry: {{License_Expiry}})",
    "Good Clinical Practice (GCP) Training Certificate",
    "Executed Financial Disclosure Form (FDF)",
    "Copy of the IRB/IEC Submission Receipt for this PI Change",
]


_PROTOCOL_DEVIATION_DEFAULT_ATTACHMENTS: List[str] = [
    "Deviation Report Form",
    "Root Cause Analysis",
    "CAPA (Corrective and Preventive Actions)",
    "Impact on Subject Safety/Data Integrity",
    "PI Acknowledgement / Memo to File (MTF)",
]


def _build_notification_cover_letter_payload(
    *,
    profile_data: Dict[str, Any],
    study: Study,
    site: Site,
    required_docs: List[str],
    now: datetime,
    protocol_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    institution_name = (
        _safe_text(profile_data.get("hospital_name"))
        or _safe_text(profile_data.get("site_name"))
        or _safe_text(getattr(site, "name", None))
        or "King George Clinical Research Site"
    )
    crc_name = (
        _safe_text(profile_data.get("site_coordinator_name"))
        or _safe_text(profile_data.get("authorized_signatory_name"))
        or "{{CRC_Name}}"
    )
    current_pi_name = (
        _safe_text(profile_data.get("pi_name"))
        or _safe_text(profile_data.get("principal_investigator_name"))
        or "{{Current_PI}}"
    )
    incoming_pi_name = (
        _safe_text((protocol_metadata or {}).get("new_pi"))
        or _safe_text(profile_data.get("new_pi"))
        or "{{new_pi}}"
    )
    protocol_number = (
        _clean_protocol_text((protocol_metadata or {}).get("protocol_number"))
        or _safe_text(profile_data.get("protocol_number"))
        or _safe_text(getattr(study, "study_id", None))
        or "{{Protocol_Number}}"
    )
    site_number = _safe_text(getattr(site, "site_id", None)) or _safe_text(getattr(site, "code", None)) or "{{Site_Number}}"
    attachment_names = [_safe_text(name) for name in required_docs if _safe_text(name)]

    return {
        "date": now.strftime("%d-%b-%Y"),
        "package_id": f"NOTIF-PI-{now.strftime('%Y')}-{site_number}",
        "sponsor_or_cro_name": "{{Sponsor_or_CRO_Name}}",
        "protocol_number": protocol_number,
        "site_number": site_number,
        "institution_name": institution_name,
        "effective_date_of_transfer": "{{Effective_Date_of_Transfer}}",
        "outgoing_pi_name": current_pi_name,
        "reason_for_change": "{{Reason_for_Change_e.g., relocation}}",
        "incoming_pi_name": incoming_pi_name,
        "crc_name": crc_name,
        "attachments": attachment_names or list(_NOTIFICATION_DEFAULT_ATTACHMENTS),
    }


def _package_id_segment(value: Any, fallback: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", _safe_text(value)).strip("-")
    return cleaned or fallback


def _build_protocol_deviation_cover_letter_payload(
    *,
    profile_data: Dict[str, Any],
    study: Study,
    site: Site,
    required_docs: List[str],
    now: datetime,
    protocol_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    metadata = protocol_metadata or {}
    institution_name = (
        _safe_text(profile_data.get("hospital_name"))
        or _safe_text(profile_data.get("site_name"))
        or _safe_text(getattr(site, "name", None))
        or "{{Site_Name}}"
    )
    crc_name = (
        _safe_text(profile_data.get("site_coordinator_name"))
        or _safe_text(profile_data.get("authorized_signatory_name"))
        or "{{CRC_Name}}"
    )
    investigator_name = (
        _safe_text(metadata.get("investigator_name"))
        or _safe_text(profile_data.get("pi_name"))
        or _safe_text(profile_data.get("principal_investigator_name"))
        or "{{Investigator_Name}}"
    )
    protocol_number = (
        _clean_protocol_text(metadata.get("protocol_number"))
        or _safe_text(profile_data.get("protocol_number"))
        or _safe_text(getattr(study, "study_id", None))
        or "{{Protocol_Number}}"
    )
    site_number = (
        _safe_text(getattr(site, "site_id", None))
        or _safe_text(getattr(site, "code", None))
        or "{{Site_Number}}"
    )
    unique_id = "-".join(
        part
        for part in [
            _package_id_segment(site_number, ""),
            now.strftime("%Y%m%d"),
        ]
        if part
    ) or now.strftime("%Y%m%d")
    attachment_names = [_safe_text(name) for name in required_docs if _safe_text(name)]
    capa_steps = metadata.get("capa_steps")
    if isinstance(capa_steps, list):
        capa_lines = [_safe_text(step) for step in capa_steps if _safe_text(step)]
    else:
        capa_lines = []

    return {
        "date": now.strftime("%d-%b-%Y"),
        "notification_type": "Operational - Protocol Deviation Notification",
        "package_id": f"NOTIF-DEV-{_package_id_segment(protocol_number, 'Protocol_Number')}-{unique_id}",
        "sponsor_or_cro_name": _safe_text(metadata.get("sponsor_or_cro_name")) or "{{Sponsor_or_CRO_Name}}",
        "protocol_number": protocol_number,
        "site_number": site_number,
        "site_name": institution_name,
        "subject_id": _safe_text(metadata.get("subject_id")) or "{{Subject_ID}}",
        "date_of_occurrence": _safe_text(metadata.get("date_of_occurrence")) or "{{Date_of_Occurrence}}",
        "deviation_classification": _safe_text(metadata.get("deviation_classification"))
        or "[ ] Minor / Non-Critical  [X] Major / Critical",
        "procedural_category": _safe_text(metadata.get("procedural_category"))
        or "[e.g., Out-of-Window Visit / Informed Consent / Lab Eligibility]",
        "deviation_description": _safe_text(metadata.get("deviation_description"))
        or "{{Dynamic_Deviation_Description_e.g., Subject missed the Cycle 3 Day 15 safety lab draws due to unexpected hospitalization for an unrelated event. Labs were drawn 4 days outside of the protocol-permissible window (+/- 2 days).}}",
        "impact_statement": _safe_text(metadata.get("impact_statement")) or "[did / did not]",
        "capa_steps": capa_lines
        or [
            "{{CAPA_Step_1_e.g., Site personnel re-trained on protocol visit window calculators.}}",
            "{{CAPA_Step_2_e.g., Pre-scheduling window alerts added to the local site calendar.}}",
        ],
        "attachments": attachment_names or list(_PROTOCOL_DEVIATION_DEFAULT_ATTACHMENTS),
        "crc_name": crc_name,
        "investigator_name": investigator_name,
    }


def _generate_notification_cover_letter_docx(payload: Dict[str, Any]) -> bytes:
    from io import BytesIO

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    data = payload or {}
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    def add_paragraph(text: str = "", *, bold: bool = False, align=None, size: float = 10.5):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Arial"
        r.font.size = Pt(size)
        return p

    add_paragraph("[SITE INSTITUTION LETTERHEAD]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_paragraph("ONCOLOGY CLINICAL RESEARCH CENTER", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph("")
    add_paragraph(f"Date: {data.get('date') or '{{Date}}'}")
    add_paragraph("Notification Type: Regulatory - Change of Principal Investigator")
    add_paragraph(f"Package ID: {data.get('package_id') or '{{Package_ID}}'}")
    add_paragraph("")
    add_paragraph("TO:", bold=True)
    add_paragraph("Global Clinical Operations & Study Management Team")
    add_paragraph(str(data.get("sponsor_or_cro_name") or "{{Sponsor_or_CRO_Name}}"))
    add_paragraph("")
    add_paragraph("SUBJECT: OFFICIAL NOTIFICATION OF CHANGE IN PRINCIPAL INVESTIGATOR", bold=True)
    add_paragraph(f"         Protocol Number:  {data.get('protocol_number') or '{{Protocol_Number}}'}", bold=True)
    add_paragraph(f"         Site Number:      {data.get('site_number') or '{{Site_Number}}'}", bold=True)
    add_paragraph(f"         Institution:      {data.get('institution_name') or '{{Institution_Name}}'}", bold=True)
    add_paragraph("")
    add_paragraph("Dear Study Management Team,")
    add_paragraph("")
    add_paragraph(
        "This letter serves as formal notification of a change in study leadership and institutional "
        "oversight for the above-referenced protocol at our site."
    )
    add_paragraph("")
    add_paragraph(
        f"Effective as of {data.get('effective_date_of_transfer')}, {data.get('outgoing_pi_name')} will step down "
        f"as the Principal Investigator for this trial due to {data.get('reason_for_change')}."
    )
    add_paragraph("")
    add_paragraph(
        f"{data.get('incoming_pi_name')} has agreed to assume full responsibility for the conduct of the "
        "study, including patient safety, protocol compliance, and data integrity, in accordance "
        "with ICH GCP E6(R2/R3) guidelines, FDA regulations, and local institutional policies."
    )
    add_paragraph("")
    add_paragraph("ATTACHED REGULATORY DOCUMENTATION", bold=True)
    add_paragraph(
        "To maintain compliance and update the trial records, the following essential documents "
        "concerning the incoming Investigator are attached to this notification package:"
    )
    for idx, name in enumerate(data.get("attachments") or _NOTIFICATION_DEFAULT_ATTACHMENTS, start=1):
        add_paragraph(f"{idx}. {name}")
    add_paragraph("")
    add_paragraph("SITE DELEGATION LOG STATUS", bold=True)
    add_paragraph(
        "Please be advised that the digital Delegation of Authority (DoA) Log within the eISF "
        f"has been updated. {data.get('incoming_pi_name')} has reviewed, signed, and re-authorized the "
        "delegated duties for all active sub-investigators and study coordinators listed in the "
        "current site personnel register."
    )
    add_paragraph("")
    add_paragraph("ACTION & ACKNOWLEDGEMENT REQUIRED", bold=True)
    add_paragraph(
        "Please update the Master Site Directory within your Clinical Trial Management System (CTMS) "
        "and route the attached essential documents to the electronic Trial Master File (eTMF)."
    )
    add_paragraph("")
    add_paragraph(
        "Kindly acknowledge receipt and system approval of this investigator transfer by executing "
        "the electronic sign-off block in the portal within 5 business days."
    )
    add_paragraph("")
    add_paragraph("Sincerely,")
    add_paragraph("")
    add_paragraph("[Electronic Signature via Portal]")
    add_paragraph("____________________________________________")
    add_paragraph(f"{data.get('crc_name')}, Lead Clinical Research Coordinator")
    add_paragraph(f"On Behalf of {data.get('institution_name')}")
    add_paragraph("")
    add_paragraph("[Electronic Counter-Signature via Portal]")
    add_paragraph("____________________________________________")
    add_paragraph(f"{data.get('incoming_pi_name')}, Incoming Principal Investigator")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _generate_protocol_deviation_cover_letter_docx(payload: Dict[str, Any]) -> bytes:
    from io import BytesIO

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    data = payload or {}
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    def add_paragraph(text: str = "", *, bold: bool = False, align=None, size: float = 10.5):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        r = p.add_run(text)
        r.bold = bold
        r.font.name = "Arial"
        r.font.size = Pt(size)
        return p

    add_paragraph("[SITE INSTITUTION LETTERHEAD]", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_paragraph("CLINICAL RESEARCH COMPLIANCE UNIT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_paragraph("=" * 80, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph("")
    add_paragraph(f"Date: {data.get('date') or '{{Current_Date}}'}")
    add_paragraph(f"Notification Type: {data.get('notification_type') or 'Operational - Protocol Deviation Notification'}")
    add_paragraph(f"Package ID: {data.get('package_id') or '{{Package_ID}}'}")
    add_paragraph("")
    add_paragraph("TO:", bold=True)
    add_paragraph("Assigned Clinical Research Associate (CRA) & Study Management Team")
    add_paragraph(str(data.get("sponsor_or_cro_name") or "{{Sponsor_or_CRO_Name}}"))
    add_paragraph("")
    add_paragraph("SUBJECT: OFFICIAL NOTIFICATION OF PROTOCOL DEVIATION", bold=True)
    add_paragraph(f"         Protocol Number:  {data.get('protocol_number') or '{{Protocol_Number}}'}", bold=True)
    add_paragraph(f"         Site Number:      {data.get('site_number') or '{{Site_Number}}'}", bold=True)
    add_paragraph(f"         Institution:      {data.get('site_name') or '{{Site_Name}}'}", bold=True)
    add_paragraph("")
    add_paragraph("Dear Study Management Team,")
    add_paragraph("")
    add_paragraph(
        "This letter serves as formal notification of a protocol deviation that occurred at our "
        "facility during the conduct of the above-referenced clinical trial."
    )
    add_paragraph("")
    add_paragraph("EVENT SUMMARY", bold=True)
    add_paragraph(f"Subject ID:                {data.get('subject_id') or '{{Subject_ID}}'}")
    add_paragraph(f"Date of Occurrence:        {data.get('date_of_occurrence') or '{{Date_of_Occurrence}}'}")
    add_paragraph(f"Deviation Classification:  {data.get('deviation_classification') or '[ ] Minor / Non-Critical  [X] Major / Critical'}")
    add_paragraph(f"Procedural Category:       {data.get('procedural_category') or '[e.g., Out-of-Window Visit / Informed Consent / Lab Eligibility]'}")
    add_paragraph("")
    add_paragraph("Description of the Event:", bold=True)
    add_paragraph(str(data.get("deviation_description") or "{{Dynamic_Deviation_Description}}"))
    add_paragraph("")
    add_paragraph("IMPACT AND CORRECTIVE ACTIONS", bold=True)
    add_paragraph(
        "An immediate internal review was performed by our research team. It has been determined "
        f"that this event {data.get('impact_statement') or '[did / did not]'} compromise subject safety, "
        "data integrity, or the ethical conduct of the trial."
    )
    add_paragraph("")
    add_paragraph(
        "To prevent recurrence of this specific issue, the following Corrective and Preventive "
        "Action (CAPA) steps have been implemented at the site level:"
    )
    for step in data.get("capa_steps") or []:
        add_paragraph(str(step))
    add_paragraph("")
    add_paragraph("ATTACHED REGULATORY DOCUMENTATION", bold=True)
    add_paragraph(
        "The following source metrics and tracking files are bundled within this notification "
        "package for your review:"
    )
    for idx, name in enumerate(data.get("attachments") or _PROTOCOL_DEVIATION_DEFAULT_ATTACHMENTS, start=1):
        add_paragraph(f"{idx}. {name}")
    add_paragraph("")
    add_paragraph("COMPLIANCE & REGULATORY STATEMENT", bold=True)
    add_paragraph(
        "In accordance with ICH GCP E6(R2/R3) Section 4.5.3, the investigator should document and "
        "explain any deviation from the approved protocol. This notification fulfills the site's "
        "obligation for prompt, transparent reporting of operational variations to the Sponsor/CRO."
    )
    add_paragraph("")
    add_paragraph("ACTION & ACKNOWLEDGEMENT REQUIRED", bold=True)
    add_paragraph(
        "Please route this notification package directly into the master integration channel. The "
        "assigned CRA must review the attached documentation, verify the categorization inside the "
        "Sponsor CTMS Deviation Module, and map the package to the electronic Trial Master File (eTMF)."
    )
    add_paragraph("")
    add_paragraph(
        "Kindly execute the portal electronic sign-off block below to acknowledge receipt and "
        "reconciliation of this event."
    )
    add_paragraph("")
    add_paragraph("Sincerely,")
    add_paragraph("")
    add_paragraph("[Electronic Signature via Portal]")
    add_paragraph("____________________________________________")
    add_paragraph(f"{data.get('crc_name') or '{{CRC_Name}}'}, Lead Clinical Research Coordinator")
    add_paragraph(f"On Behalf of {data.get('site_name') or '{{Site_Name}}'}")
    add_paragraph("")
    add_paragraph("[Electronic Counter-Signature via Portal]")
    add_paragraph("____________________________________________")
    add_paragraph(f"{data.get('investigator_name') or '{{Investigator_Name}}'}, Principal Investigator")
    add_paragraph("=" * 80, align=WD_ALIGN_PARAGRAPH.CENTER)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def generate_site_package_document(
    *,
    study_id: str,
    site_id: str,
    document_type: str,
    document_name: Optional[str],
    output_format: str = "docx",
    required_document_names: Optional[List[str]] = None,
    protocol_metadata: Optional[Dict[str, Any]] = None,
    db: AsyncSession,
) -> Dict[str, Any]:
    normalized_type = (document_type or "").strip().lower()
    if normalized_type not in {"tmc_cover_letter", "cover_letter", "notification_cover_letter", "protocol_deviation_cover_letter"}:
        raise HTTPException(status_code=400, detail="Unsupported document type for generation")

    normalized_output_format = (output_format or "docx").strip().lower()
    if normalized_output_format not in {"docx", "pdf"}:
        raise HTTPException(status_code=400, detail="Unsupported output format. Use 'docx' or 'pdf'.")

    study_uuid = await _resolve_study_uuid(db, study_id)
    site_uuid = await _resolve_site_uuid(db, site_id)

    study_result = await db.execute(select(Study).where(Study.id == study_uuid))
    study = study_result.scalar_one_or_none()
    if not study:
        raise HTTPException(status_code=404, detail="Study not found")
    site_result = await db.execute(select(Site).where(Site.id == site_uuid))
    site = site_result.scalar_one_or_none()
    if not site:
        raise HTTPException(status_code=404, detail="Site not found")

    profile_result = await db.execute(select(SiteProfile).where(SiteProfile.site_id == site.id))
    profile = profile_result.scalar_one_or_none()
    if not profile:
        raise HTTPException(status_code=404, detail="Site profile not found")

    # Build a read-only profile view with dynamic/system fields for generation.
    # Avoid mutating SQLAlchemy model attributes that may not exist as columns.
    profile_data: Dict[str, Any] = {
        col.name: getattr(profile, col.name, None)
        for col in profile.__table__.columns
    }

    profile_data["study_title"] = getattr(study, "name", None)
    profile_data["protocol_number"] = getattr(study, "study_id", None)
    profile_data["contact_email"] = (
        profile_data.get("site_coordinator_email")
        or profile_data.get("pi_email")
    )

    # Dynamic/system placeholders that are required by the template.
    now = datetime.now()
    profile_data["generated_date"] = now.strftime("%d-%m-%Y")
    
    # Try to get Registration ID and Chairperson name from IRB Administrative Info
    registration_id = None
    mapping_result = await db.execute(
        select(SiteIRBMapping).where(SiteIRBMapping.site_id == site.id)
    )
    mapping = mapping_result.scalar_one_or_none()
    irb = None
    admin_info: Optional[IRBAdministrativeInfo] = None
    if mapping:
        irb_result = await db.execute(select(IRB).where(IRB.id == mapping.irb_id))
        irb = irb_result.scalar_one_or_none()
        admin_info_result = await db.execute(
            select(IRBAdministrativeInfo).where(IRBAdministrativeInfo.irb_id == mapping.irb_id)
        )
        admin_info = admin_info_result.scalar_one_or_none()
        if admin_info and admin_info.registration_id:
            registration_id = admin_info.registration_id

    existing_ref = profile_data.get("iec_ref_no")
    if registration_id:
        profile_data["iec_ref_no"] = registration_id
    elif not existing_ref and normalized_type == "tmc_cover_letter":
        site_code = (getattr(site, "site_id", None) or str(site.id)).strip()
        profile_data["iec_ref_no"] = f"TMC-IEC/{site_code}/{now.strftime('%Y%m%d')}"

    profile_data["rec_ref"] = profile_data.get("iec_ref_no") or profile_data.get("protocol_number")
    profile_data["iras_id"] = profile_data.get("protocol_number")
    profile_data["protocol_version"] = profile_data.get("protocol_version") or "1.0"

    if normalized_type == "notification_cover_letter":
        required_docs = [
            _safe_text(name)
            for name in (required_document_names or _NOTIFICATION_DEFAULT_ATTACHMENTS)
            if _safe_text(name)
        ]
        required_docs_source = "uploaded_documents" if required_document_names else "notification_defaults"
        logger.info(
            "Generating notification PI-change cover letter with %s attachments (source=%s)",
            len(required_docs),
            required_docs_source,
        )
        payload = _build_notification_cover_letter_payload(
            profile_data=profile_data,
            study=study,
            site=site,
            required_docs=required_docs,
            now=now,
            protocol_metadata=protocol_metadata,
        )
        generated_docx = _generate_notification_cover_letter_docx(payload)
    elif normalized_type == "protocol_deviation_cover_letter":
        required_docs = [
            _safe_text(name)
            for name in (required_document_names or _PROTOCOL_DEVIATION_DEFAULT_ATTACHMENTS)
            if _safe_text(name)
        ]
        required_docs_source = "uploaded_documents" if required_document_names else "protocol_deviation_defaults"
        logger.info(
            "Generating protocol deviation cover letter with %s attachments (source=%s)",
            len(required_docs),
            required_docs_source,
        )
        payload = _build_protocol_deviation_cover_letter_payload(
            profile_data=profile_data,
            study=study,
            site=site,
            required_docs=required_docs,
            now=now,
            protocol_metadata=protocol_metadata,
        )
        generated_docx = _generate_protocol_deviation_cover_letter_docx(payload)
    else:
        required_docs, required_docs_source = await _resolve_tmc_required_documents(
            db,
            site,
            preferred_names=required_document_names,
        )
        logger.info(
            "Generating standardized IRB cover letter (python-docx) with %s required documents (source=%s)",
            len(required_docs),
            required_docs_source,
        )
        payload = _build_standard_irb_payload(
            profile_data=profile_data,
            study=study,
            site=site,
            irb=irb,
            required_docs=required_docs,
            now=now,
            irb_admin_info=admin_info,
            protocol_metadata=protocol_metadata,
        )
        generated_docx = _generate_standard_irb_cover_letter_docx(payload)

    out_path = await asyncio.to_thread(_make_named_tempfile, ".docx")
    pdf_path = await asyncio.to_thread(_make_named_tempfile, ".pdf")

    try:
        out_path.write_bytes(generated_docx)

        if normalized_output_format == "pdf":
            try:
                docx_to_pdf(out_path, pdf_path, method="libreoffice")
                data = pdf_path.read_bytes()
            except Exception as conversion_error:
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to convert generated DOCX to PDF: {conversion_error}",
                ) from conversion_error
        else:
            data = out_path.read_bytes()
    finally:
        try:
            out_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            pdf_path.unlink(missing_ok=True)
        except Exception:
            pass

    default_names = {
        "notification_cover_letter": "PI_Change_Notification_Cover_Letter",
        "protocol_deviation_cover_letter": "Protocol_Deviation_Notification_Cover_Letter",
    }
    default_name = default_names.get(normalized_type, "Standard_IRB_Cover_Letter")
    output_name = (document_name or default_name).strip().replace(" ", "_")
    output_ext = ".pdf" if normalized_output_format == "pdf" else ".docx"
    if output_name.lower().endswith(".pdf") or output_name.lower().endswith(".docx"):
        output_name = str(Path(output_name).with_suffix(output_ext))
    else:
        output_name = f"{output_name}{output_ext}"

    mime_type = (
        "application/pdf"
        if normalized_output_format == "pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    return {
        "filename": output_name,
        "mime_type": mime_type,
        "content": data,
        "study_id": study_id,
        "site_id": site_id,
        "required_docs_count": len(required_docs),
        "required_docs_source": required_docs_source,
    }


async def generate_cover_letter_assets(
    *,
    study_id: str,
    site_id: str,
    document_type: str,
    document_name: Optional[str],
    required_document_names: Optional[List[str]],
    protocol_metadata: Optional[Dict[str, Any]],
    db: AsyncSession,
) -> Dict[str, Any]:
    normalized_type = (document_type or "").strip().lower()
    if normalized_type not in {"cover_letter", "tmc_cover_letter", "notification_cover_letter", "protocol_deviation_cover_letter"}:
        raise HTTPException(status_code=400, detail="Unsupported document type for cover-letter asset generation")

    docx_result = await generate_site_package_document(
        study_id=study_id,
        site_id=site_id,
        document_type=normalized_type,
        document_name=document_name,
        output_format="docx",
        required_document_names=required_document_names,
        protocol_metadata=protocol_metadata,
        db=db,
    )

    docx_bytes = docx_result["content"]
    output_name = str(docx_result.get("filename") or "Cover_Letter.docx")
    base_name = Path(output_name).stem
    docx_filename = f"{base_name}.docx"
    pdf_filename = f"{base_name}.pdf"

    tmp_docx_path = await asyncio.to_thread(_make_named_tempfile, ".docx", docx_bytes)
    tmp_pdf_path = await asyncio.to_thread(_make_named_tempfile, ".pdf")

    try:
        try:
            await asyncio.to_thread(docx_to_pdf, tmp_docx_path, tmp_pdf_path, "libreoffice")
            pdf_bytes = tmp_pdf_path.read_bytes()
        except Exception as conversion_error:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to convert generated DOCX to PDF: {conversion_error}",
            ) from conversion_error
    finally:
        try:
            tmp_docx_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            tmp_pdf_path.unlink(missing_ok=True)
        except Exception:
            pass

    docx_upload = await _upload_generated_bytes_to_azure(
        content=docx_bytes,
        filename=docx_filename,
        study_id=study_id,
        site_id=site_id,
        document_type=f"{normalized_type}_docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    pdf_upload = await _upload_generated_bytes_to_azure(
        content=pdf_bytes,
        filename=pdf_filename,
        study_id=study_id,
        site_id=site_id,
        document_type=f"{normalized_type}_pdf",
        content_type="application/pdf",
    )

    return {
        "name": document_name or base_name,
        "type": normalized_type,
        "version": 1,
        "docx": docx_upload,
        "pdf": pdf_upload,
        "required_docs_count": int(docx_result.get("required_docs_count") or 0),
        "required_docs_source": str(docx_result.get("required_docs_source") or "none"),
    }


async def get_cover_letter_onlyoffice_config(
    *,
    pkg_id: str,
    study_value: str,
    site_value: str,
    document_name: Optional[str],
    current_user: Optional[dict],
    db: AsyncSession,
) -> Dict[str, Any]:
    _require_study_site_for_scope(study_value, site_value)
    study_uuid = await _resolve_study_uuid(db, study_value)
    site_uuid = await _resolve_site_uuid(db, site_value)

    try:
        pkg_uuid = UUID(str(pkg_id))
    except Exception:
        raise HTTPException(status_code=404, detail="Site package not found")

    result = await db.execute(select(SitePackage).where(SitePackage.id == pkg_uuid).where(SitePackage.isDeleted.is_(False)))
    pkg = result.scalar_one_or_none()
    if not pkg:
        raise HTTPException(status_code=404, detail="Site package not found")
    _assert_package_in_scope(pkg, study_uuid, site_uuid)

    if not isDraftStatus(getattr(pkg, "status", None)):
        raise HTTPException(status_code=400, detail="Only Draft site packages can be edited")

    requested_name = (document_name or "").strip()
    documents = list(pkg.documents or [])
    selected_index, selected_doc = _select_cover_letter_document(
        documents=documents,
        requested_name=requested_name,
        require_editable_docx=True,
    )

    if selected_doc is None:
        raise HTTPException(status_code=404, detail="Editable cover letter DOCX not found")

    from app.utils.onlyoffice_utils import create_document_config

    doc_name = str(selected_doc.get("name") or "Cover Letter").strip() or "Cover Letter"
    doc_title = doc_name if doc_name.lower().endswith(".docx") else f"{doc_name}.docx"
    docx_url = _resolve_cover_docx_url(selected_doc)
    internal_base = _resolve_onlyoffice_internal_base(settings.backend_internal_url)
    callback_query = f"study={study_value}&site={site_value}&document_name={doc_name}"
    callback_url = f"{internal_base}/api/site-packages/{pkg_id}/cover-letter-onlyoffice-callback?{callback_query}"
    # Cache-bust document fetch so ONLYOFFICE always requests the latest file bytes.
    # Prefer semantic freshness markers (document version / updatedAt), with a time fallback.
    updated_at_raw = selected_doc.get("updatedAt")
    updated_at_ms = 0
    if updated_at_raw:
        try:
            updated_at_ms = int(datetime.fromisoformat(str(updated_at_raw).replace("Z", "+00:00")).timestamp() * 1000)
        except Exception:
            updated_at_ms = 0
    cache_bust = updated_at_ms or int(time.time() * 1000)
    document_url = (
        f"{internal_base}/api/site-packages/{pkg_id}/cover-letter-onlyoffice-document"
        f"?study={study_value}&site={site_value}&document_name={doc_name}&v={cache_bust}"
    )

    version_value = int(selected_doc.get("version") or 1)
    # ONLYOFFICE caches by `document.key`; if this key is reused, stale content can appear.
    # Include version + updated marker + high-resolution timestamp to guarantee freshness.
    document_key = (
        f"sitepkg-{pkg_id}-{selected_index}-v{version_value}-u{cache_bust}-n{time.time_ns()}"
    )

    user_id = str((current_user or {}).get("user_id") or pkg.createdBy or "site-package-user")
    user_name = str((current_user or {}).get("name") or user_id)

    config = create_document_config(
        document_url=document_url,
        callback_url=callback_url,
        document_key=document_key,
        document_title=doc_title,
        user_id=user_id,
        user_name=user_name,
        mode="edit",
        file_type="docx",
        document_type="word",
    )

    return {
        "config": config,
        "document": {
            "name": doc_name,
            "version": version_value,
            "docxFileUrl": docx_url,
            "pdfFileUrl": str(selected_doc.get("pdfFileUrl") or selected_doc.get("fileUrl") or ""),
        },
    }


async def get_cover_letter_onlyoffice_status(
    *,
    pkg_id: str,
    study_value: str,
    site_value: str,
    document_name: Optional[str],
    db: AsyncSession,
) -> Dict[str, Any]:
    _require_study_site_for_scope(study_value, site_value)
    study_uuid = await _resolve_study_uuid(db, study_value)
    site_uuid = await _resolve_site_uuid(db, site_value)

    try:
        pkg_uuid = UUID(str(pkg_id))
    except Exception:
        raise HTTPException(status_code=404, detail="Site package not found")

    result = await db.execute(select(SitePackage).where(SitePackage.id == pkg_uuid).where(SitePackage.isDeleted.is_(False)))
    pkg = result.scalar_one_or_none()
    if not pkg:
        raise HTTPException(status_code=404, detail="Site package not found")
    _assert_package_in_scope(pkg, study_uuid, site_uuid)

    requested_name = (document_name or "").strip()
    documents = list(pkg.documents or [])
    _, selected_doc = _select_cover_letter_document(
        documents=documents,
        requested_name=requested_name,
        require_editable_docx=False,
    )
    if selected_doc is None:
        raise HTTPException(status_code=404, detail="Cover letter document not found")

    updated_at_value = selected_doc.get("updatedAt")
    if not updated_at_value:
        updated_at_value = getattr(pkg, "lastUpdated", None)
        if isinstance(updated_at_value, datetime):
            updated_at_value = updated_at_value.isoformat()

    return {
        "name": str(selected_doc.get("name") or document_name or "Cover Letter"),
        "version": int(selected_doc.get("version") or 1),
        "updatedAt": str(updated_at_value or ""),
        "pdfFileUrl": str(selected_doc.get("pdfFileUrl") or selected_doc.get("fileUrl") or ""),
        "docxFileUrl": str(selected_doc.get("docxFileUrl") or ""),
    }


def isDraftStatus(status: Any) -> bool:
    return str(status or "Draft").strip().lower() == "draft"


async def handle_cover_letter_onlyoffice_callback(
    *,
    pkg_id: str,
    study_value: str,
    site_value: str,
    document_name: Optional[str],
    callback_payload: Dict[str, Any],
    db: AsyncSession,
) -> Dict[str, int]:
    status = int(callback_payload.get("status") or 0)
    # 2 = MustSave, 6 = MustForceSave in ONLYOFFICE callback contract.
    if status not in {2, 6}:
        return {"error": 0}

    file_url = str(callback_payload.get("url") or "").strip()
    if not file_url:
        logger.warning("Site package ONLYOFFICE callback missing url for pkg_id=%s", sfmt(pkg_id))
        return {"error": 0}

    _require_study_site_for_scope(study_value, site_value)
    study_uuid = await _resolve_study_uuid(db, study_value)
    site_uuid = await _resolve_site_uuid(db, site_value)

    try:
        pkg_uuid = UUID(str(pkg_id))
    except Exception:
        raise HTTPException(status_code=404, detail="Site package not found")

    result = await db.execute(select(SitePackage).where(SitePackage.id == pkg_uuid).where(SitePackage.isDeleted.is_(False)))
    pkg = result.scalar_one_or_none()
    if not pkg:
        raise HTTPException(status_code=404, detail="Site package not found")
    _assert_package_in_scope(pkg, study_uuid, site_uuid)

    requested_name = (document_name or "").strip()
    documents = list(pkg.documents or [])
    selected_index, selected_doc = _select_cover_letter_document(
        documents=documents,
        requested_name=requested_name,
        require_editable_docx=False,
    )

    if selected_doc is None or selected_index < 0:
        raise HTTPException(status_code=404, detail="Cover letter document not found")

    rewritten_url = _rewrite_onlyoffice_callback_url_if_needed(file_url)
    async with httpx.AsyncClient(timeout=httpx.Timeout(120.0, connect=30.0), follow_redirects=True, trust_env=True) as client:
        response = await client.get(rewritten_url)
        response.raise_for_status()
        docx_bytes = response.content

    tmp_docx_path = await asyncio.to_thread(_make_named_tempfile, ".docx", docx_bytes)
    tmp_pdf_path = await asyncio.to_thread(_make_named_tempfile, ".pdf")

    try:
        try:
            await asyncio.to_thread(docx_to_pdf, tmp_docx_path, tmp_pdf_path, "libreoffice")
            pdf_bytes = tmp_pdf_path.read_bytes()
        except Exception as conversion_error:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to convert edited DOCX to PDF: {conversion_error}",
            ) from conversion_error
    finally:
        try:
            tmp_docx_path.unlink(missing_ok=True)
        except Exception:
            pass
        try:
            tmp_pdf_path.unlink(missing_ok=True)
        except Exception:
            pass

    current_name = str(selected_doc.get("name") or "Cover Letter").strip() or "Cover Letter"
    current_type = str(selected_doc.get("type") or "cover_letter").strip().lower() or "cover_letter"
    current_version = int(selected_doc.get("version") or 1)
    next_version = current_version + 1

    base_name = Path(current_name).stem.replace(" ", "_") or "Cover_Letter"
    docx_filename = f"{base_name}_v{next_version}.docx"
    pdf_filename = f"{base_name}_v{next_version}.pdf"

    docx_upload = await _upload_generated_bytes_to_azure(
        content=docx_bytes,
        filename=docx_filename,
        study_id=str(study_value),
        site_id=str(site_value),
        document_type=f"{current_type}_docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    pdf_upload = await _upload_generated_bytes_to_azure(
        content=pdf_bytes,
        filename=pdf_filename,
        study_id=str(study_value),
        site_id=str(site_value),
        document_type=f"{current_type}_pdf",
        content_type="application/pdf",
    )

    now = _now_utc()
    updated_doc = dict(selected_doc)
    updated_doc.update(
        {
            "type": current_type,
            "version": next_version,
            "fileUrl": pdf_upload["fileUrl"],
            "pdfFileUrl": pdf_upload["fileUrl"],
            "pdfBlobPath": pdf_upload.get("blobPath"),
            "docxFileUrl": docx_upload["fileUrl"],
            "docxBlobPath": docx_upload.get("blobPath"),
            "fileSize": pdf_upload.get("fileSize") or 0,
            "mimeType": "application/pdf",
            "docxMimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "updatedAt": now.isoformat(),
        }
    )

    documents[selected_index] = updated_doc
    pkg.documents = documents
    pkg.lastUpdated = now
    audit_item = {
        "action": "cover_letter_edited",
        "details": {
            "documentName": current_name,
            "version": next_version,
            "docxFileUrl": docx_upload["fileUrl"],
            "pdfFileUrl": pdf_upload["fileUrl"],
        },
        "timestamp": now.isoformat(),
    }
    pkg.auditTrail = list(pkg.auditTrail or []) + [audit_item]

    await db.commit()
    return {"error": 0}


async def get_cover_letter_onlyoffice_document(
    *,
    pkg_id: str,
    study_value: str,
    site_value: str,
    document_name: Optional[str],
    db: AsyncSession,
) -> Tuple[bytes, str, str]:
    _require_study_site_for_scope(study_value, site_value)
    study_uuid = await _resolve_study_uuid(db, study_value)
    site_uuid = await _resolve_site_uuid(db, site_value)

    try:
        pkg_uuid = UUID(str(pkg_id))
    except Exception:
        raise HTTPException(status_code=404, detail="Site package not found")

    result = await db.execute(select(SitePackage).where(SitePackage.id == pkg_uuid).where(SitePackage.isDeleted.is_(False)))
    pkg = result.scalar_one_or_none()
    if not pkg:
        raise HTTPException(status_code=404, detail="Site package not found")
    _assert_package_in_scope(pkg, study_uuid, site_uuid)

    requested_name = (document_name or "").strip()
    documents = list(pkg.documents or [])
    _, selected_doc = _select_cover_letter_document(
        documents=documents,
        requested_name=requested_name,
        require_editable_docx=True,
    )

    if selected_doc is None:
        raise HTTPException(status_code=404, detail="Editable cover letter DOCX not found")

    docx_url = _resolve_cover_docx_url(selected_doc)
    data, media_type = await fetch_remote_blob_bytes_for_client(file_url=docx_url)
    media_lower = str(media_type or "").lower()
    if "wordprocessingml" in media_lower or docx_url.lower().endswith(".docx"):
        data = _remove_docx_protection_bytes(data)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    doc_name = str(selected_doc.get("name") or "Cover Letter").strip() or "Cover Letter"
    return data, media_type or "application/vnd.openxmlformats-officedocument.wordprocessingml.document", doc_name


async def _build_requirement_lookup_cache(
    pkg: SitePackage,
    db: AsyncSession,
) -> Tuple[Dict[str, str], Dict[str, int], Dict[str, int]]:
    """
    Build and cache requirement metadata lookups from IRB.

    Returns: (req_type_by_name_lower, req_max_by_type, req_max_by_name_lower)

    This function is extracted to be reusable across multiple document uploads
    for the same package, avoiding redundant IRB resolution and processing.
    """
    req_type_by_name_lower: Dict[str, str] = {}
    req_max_by_type: Dict[str, int] = {}
    req_max_by_name_lower: Dict[str, int] = {}

    # Always include amendment/modification workflow constraints.
    req_type_by_name_lower.update(_AMENDMENT_DOC_TYPE_BY_NAME_LOWER)
    req_max_by_name_lower.update(_AMENDMENT_DOC_MAX_BY_NAME_LOWER)
    req_max_by_type.update(_AMENDMENT_DOC_MAX_BY_TYPE)
    req_type_by_name_lower.update(_CONTINUING_REVIEW_DOC_TYPE_BY_NAME_LOWER)
    req_max_by_name_lower.update(_CONTINUING_REVIEW_DOC_MAX_BY_NAME_LOWER)
    req_max_by_type.update(_CONTINUING_REVIEW_DOC_MAX_BY_TYPE)
    req_type_by_name_lower.update(_SAFETY_REPORT_DOC_TYPE_BY_NAME_LOWER)
    req_max_by_name_lower.update(_SAFETY_REPORT_DOC_MAX_BY_NAME_LOWER)
    req_max_by_type.update(_SAFETY_REPORT_DOC_MAX_BY_TYPE)
    req_type_by_name_lower.update(_STUDY_CLOSURE_DOC_TYPE_BY_NAME_LOWER)
    req_max_by_name_lower.update(_STUDY_CLOSURE_DOC_MAX_BY_NAME_LOWER)
    req_max_by_type.update(_STUDY_CLOSURE_DOC_MAX_BY_TYPE)
    req_type_by_name_lower.update(_PROTOCOL_DEVIATION_DOC_TYPE_BY_NAME_LOWER)
    req_max_by_name_lower.update(_PROTOCOL_DEVIATION_DOC_MAX_BY_NAME_LOWER)
    req_max_by_type.update(_PROTOCOL_DEVIATION_DOC_MAX_BY_TYPE)
    req_type_by_name_lower.update(_NOTIFICATION_DOC_TYPE_BY_NAME_LOWER)
    req_max_by_name_lower.update(_NOTIFICATION_DOC_MAX_BY_NAME_LOWER)
    req_max_by_type.update(_NOTIFICATION_DOC_MAX_BY_TYPE)

    if getattr(pkg, "irb_id", None) is None:
        return req_type_by_name_lower, req_max_by_type, req_max_by_name_lower

    irb_result = await db.execute(select(IRB).where(IRB.id == int(pkg.irb_id)))
    irb = irb_result.scalar_one_or_none()
    if not irb:
        return req_type_by_name_lower, req_max_by_type, req_max_by_name_lower

    reqs = await resolve_requirements_for_irb(db, irb)
    _warn_incomplete_requirement_metadata(reqs, irb_id=int(irb.id), context="requirement_lookup_cache")

    for r in reqs:
        if not isinstance(r, dict):
            continue
        r_name = str(r.get("name", "")).strip()
        r_type = str(r.get("type", "")).strip().lower() or None
        r_max_raw = r.get("max_files")
        if r_max_raw is None:
            continue
        try:
            r_max = int(r_max_raw)
        except (TypeError, ValueError):
            continue
        if r_max < 1:
            continue
        if r_name:
            req_max_by_name_lower[r_name.lower()] = r_max
            if r_type:
                req_type_by_name_lower[r_name.lower()] = r_type
        if r_type:
            req_max_by_type[r_type] = r_max

    return req_type_by_name_lower, req_max_by_type, req_max_by_name_lower


def _create_document_info_dict(
    document_name_final: str,
    normalized_type: Optional[str],
    url: str,
    size: int,
    content_type: str,
    source_final: str,
    now: datetime,
    uploader_id: Optional[str],
    tmf_document_id: Optional[str],
) -> Dict[str, Any]:
    """
    Build the document info dictionary to store in SitePackage.
    Extracted to separate concerns and make the upload handler more readable.
    """
    document_info: Dict[str, Any] = {
        "name": document_name_final,
        **({"type": normalized_type} if normalized_type else {}),
        "fileUrl": url,
        "fileSize": size,
        "mimeType": content_type,
        "source": source_final,
        "uploadedAt": now.isoformat(),
        **({"uploadedBy": uploader_id} if uploader_id else {}),
    }
    if tmf_document_id:
        document_info["tmfDocumentId"] = tmf_document_id
    return document_info


def _create_audit_trail_item(
    uploader_id: Optional[str],
    document_name_final: str,
    url: str,
    now: datetime,
) -> Dict[str, Any]:
    """
    Build an audit trail entry for document upload.
    Extracted to separate concerns and make the upload handler more readable.
    """
    return {
        "action": "document_uploaded",
        **({"user": uploader_id} if uploader_id else {}),
        "details": {"documentName": document_name_final, "fileUrl": url},
        "timestamp": now.isoformat(),
    }


async def upload_site_package_document_to_package(
    pkg_id: str,
    file: UploadFile,
    document_name: Optional[str],
    document_type: Optional[str],
    tmf_document_id: Optional[str],
    source: Optional[str],
    db: AsyncSession,
    current_user: Optional[dict],
    study_value: Optional[str] = None,
    site_value: Optional[str] = None,
) -> Dict[str, Any]:
    if not file:
        raise HTTPException(status_code=400, detail="No file uploaded")

    _require_study_site_for_scope(study_value, site_value)
    study_uuid = await _resolve_study_uuid(db, study_value or "")
    site_uuid = await _resolve_site_uuid(db, site_value or "")

    try:
        pkg_uuid = UUID(str(pkg_id))
    except Exception:
        raise HTTPException(status_code=404, detail="Site package not found")

    result = await db.execute(
        select(SitePackage).where(SitePackage.id == pkg_uuid).where(SitePackage.isDeleted.is_(False))
    )
    pkg = result.scalar_one_or_none()
    if not pkg:
        raise HTTPException(status_code=404, detail="Site package not found")

    _assert_package_in_scope(pkg, study_uuid, site_uuid)

    content = await file.read()
    filename = file.filename or "upload"
    content_type = file.content_type or "application/octet-stream"

    container = _get_azure_sitepackages_container()
    blob_path = f"{pkg_id}/{int(time.time() * 1000)}-{filename}"

    url, size = _upload_bytes_to_azure(
        content,
        container=container,
        blob_name=blob_path,
        content_type=content_type,
    )

    document_name_final = document_name or filename
    source_final = source or "upload"
    now = _now_utc()

    # Build requirement lookup cache (reusable if multiple documents uploaded for same package)
    req_type_by_name_lower, req_max_by_type, req_max_by_name_lower = await _build_requirement_lookup_cache(pkg, db)

    # Determine normalized document type and max file limit
    normalized_type = (document_type or "").strip().lower() or None
    if not normalized_type:
        normalized_type = req_type_by_name_lower.get(document_name_final.strip().lower())

    max_files: int = 10_000
    if normalized_type and normalized_type in req_max_by_type:
        max_files = req_max_by_type[normalized_type]
    else:
        max_files = req_max_by_name_lower.get(document_name_final.strip().lower(), 10_000)

    # Validate document count against limit
    existing_docs = list(pkg.documents or [])
    if normalized_type:
        existing_for_type = [d for d in existing_docs if str(d.get("type") or "").lower() == normalized_type]
    else:
        # Fallback to per-documentName when no type is provided
        existing_for_type = [d for d in existing_docs if str(d.get("name") or "") == document_name_final]

    if max_files == 1 and existing_for_type:
        # Replace existing file(s) for this type/name
        if normalized_type:
            existing_docs = [d for d in existing_docs if str(d.get("type") or "").lower() != normalized_type]
        else:
            existing_docs = [d for d in existing_docs if str(d.get("name") or "") != document_name_final]
    else:
        if len(existing_for_type) + 1 > max_files:
            raise HTTPException(status_code=400, detail=f"Max {max_files} files allowed for this document")

    # Build document info and audit trail
    uploader_id = (current_user or {}).get("user_id") or pkg.createdBy
    document_info = _create_document_info_dict(
        document_name_final=document_name_final,
        normalized_type=normalized_type,
        url=url,
        size=size,
        content_type=content_type,
        source_final=source_final,
        now=now,
        uploader_id=uploader_id,
        tmf_document_id=tmf_document_id,
    )
    audit_item = _create_audit_trail_item(
        uploader_id=uploader_id,
        document_name_final=document_name_final,
        url=url,
        now=now,
    )

    # Update package with new document and audit entry
    pkg.documents = existing_docs + [document_info]
    pkg.lastUpdated = now
    pkg.auditTrail = list(pkg.auditTrail or []) + [audit_item]

    await db.commit()
    await db.refresh(pkg)

    pkg_prog = await _progress_for_site_package(db, pkg)
    return {
        "success": True,
        "data": {
            "document": document_info,
            "package": _serialize_site_package_single(pkg, progress_percent=pkg_prog),
        },
        "message": "Document uploaded successfully",
    }


def _assert_url_allowed_for_blob_proxy(file_url: str) -> None:
    """Restrict server-side fetch to HTTPS Azure Blob URLs (SSRF protection)."""
    from urllib.parse import urlparse

    parsed = urlparse((file_url or "").strip())
    if parsed.scheme.lower() != "https":
        raise HTTPException(status_code=400, detail="Only HTTPS blob URLs are allowed")
    host = (parsed.hostname or "").lower()
    if not host.endswith(".blob.core.windows.net"):
        raise HTTPException(status_code=400, detail="URL host is not permitted for download")


def _encode_azure_blob_https_url_for_get(file_url: str) -> str:
    """
    Percent-encode each path segment so the HTTP request matches Azure SAS expectations.
    Spaces, parentheses, etc. in blob names otherwise cause 403 AuthenticationFailed /
    Signature did not match when the signed string does not match the request path encoding.
    """
    from urllib.parse import quote, unquote, urlparse, urlunparse

    raw = (file_url or "").strip()
    if not raw:
        return raw
    p = urlparse(raw)
    if not p.scheme or not p.netloc or p.scheme.lower() != "https":
        return raw
    parts: List[str] = []
    for seg in p.path.split("/"):
        if seg == "":
            parts.append(seg)
        else:
            parts.append(quote(unquote(seg), safe=""))
    new_path = "/".join(parts)
    return urlunparse((p.scheme, p.netloc, new_path, "", p.query, p.fragment))


def _parse_azure_blob_url_container_and_name(file_url: str) -> Optional[Tuple[str, str, str]]:
    """Return (storage_account_name, container_name, blob_name) for *.blob.core.windows.net URLs."""
    from urllib.parse import unquote, urlparse

    p = urlparse((file_url or "").strip())
    if not p.scheme or p.scheme.lower() != "https" or not p.netloc:
        return None
    host = (p.hostname or "").lower()
    if not host.endswith(".blob.core.windows.net"):
        return None
    account = host[: -len(".blob.core.windows.net")]
    path_parts = [x for x in p.path.split("/") if x]
    if len(path_parts) < 2:
        return None
    container_name = path_parts[0]
    blob_name = "/".join(path_parts[1:])
    blob_name = unquote(blob_name)
    return account, container_name, blob_name


def _download_blob_via_same_storage_account(file_url: str) -> Optional[Tuple[bytes, str]]:
    """
    If the URL's storage account matches AZURE_STORAGE_CONNECTION_STRING, download using
    the account key and ignore SAS (fixes expired or signature-mismatched SAS for our blobs).
    """
    conn = _get_azure_connection_string()
    if not conn:
        return None
    parsed = _parse_azure_blob_url_container_and_name(file_url)
    if not parsed:
        return None
    url_account, container_name, blob_name = parsed
    cs_parts = _parse_connection_string(conn)
    cfg_account = (cs_parts.get("AccountName") or "").strip().lower()
    if not cfg_account or cfg_account != url_account.strip().lower():
        return None
    try:
        from azure.storage.blob import BlobServiceClient
    except Exception:
        return None
    try:
        bsc = BlobServiceClient.from_connection_string(conn)
        bc = bsc.get_blob_client(container=container_name, blob=blob_name)
        downloader = bc.download_blob()
        data = downloader.readall()
        ct = "application/octet-stream"
        if downloader.properties and downloader.properties.content_settings:
            raw_ct = downloader.properties.content_settings.content_type
            if raw_ct:
                ct = raw_ct.split(";")[0].strip()
        return data, ct
    except Exception as e:
        logger.warning(
            "Same-account SDK blob download failed (will try HTTPS): %s",
            e,
            exc_info=False,
        )
        return None


async def fetch_remote_blob_bytes_for_client(file_url: str) -> Tuple[bytes, str]:
    """
    Download bytes from a SAS URL server-side so the browser never hits Azure CORS on
    ISF / CRM template blobs (used for ZIP bundling of URL-only rows).
    """
    _assert_url_allowed_for_blob_proxy(file_url)
    stripped = (file_url or "").strip()

    def _normalize_content_type(data: bytes, content_type: str) -> str:
        ct = (content_type or "application/octet-stream").split(";")[0].strip().lower()
        # Some blob uploads are stored as octet-stream; sniff PDF bytes to preserve inline preview support.
        if ct == "application/octet-stream" and data and data[:5] == b"%PDF-":
            return "application/pdf"
        return ct or "application/octet-stream"

    # Same storage account as backend config: use account key (no SAS required).
    sdk_result = await asyncio.to_thread(_download_blob_via_same_storage_account, stripped)
    if sdk_result is not None:
        data, ct = sdk_result
        return data, _normalize_content_type(data, ct)

    encoded_url = _encode_azure_blob_https_url_for_get(stripped)
    url_candidates = [encoded_url]
    if encoded_url != stripped:
        url_candidates.append(stripped)

    # Azure occasionally rejects non-browser clients; trust_env helps corporate proxies.
    headers = {
        "User-Agent": "Mozilla/5.0 (compatible; CRM-SitePackage-Zip/1.0; +https://github.com/)",
        "Accept": "*/*",
    }
    try:
        async with httpx.AsyncClient(
            timeout=httpx.Timeout(120.0, connect=30.0),
            follow_redirects=True,
            trust_env=True,
        ) as client:
            for i, request_url in enumerate(url_candidates):
                try:
                    resp = await client.get(request_url, headers=headers)
                    resp.raise_for_status()
                    ct = (resp.headers.get("content-type") or "application/octet-stream").split(";")[0].strip()
                    data = resp.content
                    return data, _normalize_content_type(data, ct)
                except httpx.HTTPStatusError as e:
                    if e.response.status_code == 403 and i < len(url_candidates) - 1:
                        continue
                    raise
    except httpx.HTTPStatusError as e:
        snippet = ""
        try:
            snippet = (e.response.text or "")[:1200].strip()
        except Exception:
            pass
        expired = (
            "AuthorizationFailure" in snippet
            or "SignatureExpired" in snippet
            or "signature not valid" in snippet.lower()
        )
        sig_mismatch = "Signature did not match" in snippet or "AuthenticationFailed" in snippet
        if e.response.status_code == 403 and expired:
            detail = (
                "Blob access denied (403). The download link may have expired — remove the document and "
                "re-import it from ISF, then try Download ZIP again."
            )
        elif e.response.status_code == 403 and sig_mismatch:
            detail = (
                "Blob access denied (403): SAS signature did not match (often expired link or renamed blob). "
                "Re-import the document from ISF, then try Download ZIP again."
            )
        elif e.response.status_code == 403:
            detail = f"Blob access denied (403). {snippet or e.response.reason_phrase}"
        else:
            detail = f"Blob returned HTTP {e.response.status_code}. {snippet or e.response.reason_phrase}"
        logger.warning("fetch_remote_blob_bytes_for_client HTTP error: %s", detail[:500])
        raise HTTPException(status_code=502, detail=detail) from e
    except httpx.RequestError as e:
        logger.warning("fetch_remote_blob_bytes_for_client request error: %s", e)
        raise HTTPException(
            status_code=502,
            detail=f"Could not reach blob storage from the server: {e!s}. Check network or proxy settings.",
        ) from e


async def extract_protocol_metadata_from_isf_import(
    *,
    file_url: str,
    max_pages: int = 5,
) -> Dict[str, Any]:
    if not file_url or not str(file_url).strip():
        raise HTTPException(status_code=400, detail="file_url is required")

    data, content_type = await fetch_remote_blob_bytes_for_client(file_url)
    if not data:
        raise HTTPException(status_code=400, detail="Protocol file is empty")

    looks_like_pdf = data[:5] == b"%PDF-" or "pdf" in (content_type or "").lower() or str(file_url).lower().endswith(".pdf")
    if not looks_like_pdf:
        raise HTTPException(
            status_code=400,
            detail="Metadata extraction currently supports PDF protocol files only",
        )

    from app.modules.site_packages.protocol_metadata_extractor import extract_metadata

    tmp_path = await asyncio.to_thread(_make_named_tempfile, ".pdf", data)

    try:
        metadata = await asyncio.to_thread(extract_metadata, str(tmp_path), int(max_pages or 5))
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Protocol metadata extraction failed: {exc}") from exc
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass

    # Extra cleaning pass before values are injected into templates.
    metadata["study_description"] = _clean_protocol_text(metadata.get("study_description")) or None
    metadata["study_title"] = _clean_protocol_text(metadata.get("study_title")) or None
    metadata["protocol_number"] = _clean_protocol_text(metadata.get("protocol_number")) or None
    metadata["version"] = _clean_protocol_text(metadata.get("version")) or None
    metadata["date"] = _clean_protocol_text(metadata.get("date")) or None
    metadata["_missing_fields"] = [
        k
        for k in ("protocol_number", "study_title", "study_description", "version", "date")
        if not metadata.get(k)
    ]
    return metadata


def presign_read(file_url: str) -> str:
    if not file_url:
        raise HTTPException(status_code=400, detail="fileUrl is required")

    conn = _get_azure_connection_string()
    if not conn:
        # Fallback: if we can't generate a SAS, just return the URL we were given.
        return file_url

    try:
        from azure.storage.blob import BlobSasPermissions, generate_blob_sas
    except Exception as e:
        logger.warning("azure-storage-blob not available (%s); returning original URL", e)
        return file_url

    # Example:
    # https://{account}.blob.core.windows.net/{container}/{blob_path}?...
    try:
        from urllib.parse import urlparse, unquote

        parsed = urlparse(file_url)
        path_parts = [p for p in parsed.path.split("/") if p]
        if len(path_parts) < 2:
            raise ValueError("Could not parse container/blob from fileUrl")
        container_name = path_parts[0]
        blob_name = "/".join(path_parts[1:])
        blob_name = unquote(blob_name)

        parts = _parse_connection_string(conn)
        account_name = parts.get("AccountName")
        account_key = parts.get("AccountKey")
        if not account_name or not account_key:
            raise ValueError("Could not parse AccountName/AccountKey")

        now = datetime.now(timezone.utc)
        start = now - timedelta(minutes=15)
        expiry = now + timedelta(hours=1)
        sas_token = generate_blob_sas(
            account_name=account_name,
            container_name=container_name,
            blob_name=blob_name,
            account_key=account_key,
            permission=BlobSasPermissions(read=True),
            start=start,
            expiry=expiry,
        )
        return f"https://{account_name}.blob.core.windows.net/{container_name}/{blob_name}?{sas_token}"
    except HTTPException:
        raise
    except Exception as e:
        logger.error("presign_read failed: %s", e, exc_info=True)
        # Match Node style: return a 500 with the exception message
        raise HTTPException(status_code=500, detail=str(e))


def _upload_bytes_to_azure(
    data: bytes,
    *,
    container: str,
    blob_name: str,
    content_type: str,
) -> Tuple[str, int]:
    """
    Upload bytes to Azure Blob Storage.
    Returns (url, size).

    Falls back to a mock URL if Azure isn't configured.
    """
    conn = _get_azure_connection_string()
    if not conn:
        url = f"https://mock-storage.local/{int(time.time() * 1000)}-{blob_name}"
        return url, len(data)

    try:
        from azure.storage.blob import BlobServiceClient
        from azure.storage.blob import ContentSettings
    except Exception as e:
        logger.warning("azure-storage-blob not available (%s); using mock URL", e)
        return f"https://mock-storage.local/{int(time.time() * 1000)}-{blob_name}", len(data)

    try:
        service_client = BlobServiceClient.from_connection_string(conn)
        container_client = service_client.get_container_client(container)
        try:
            container_client.create_container()
        except Exception:
            pass

        blob_client = container_client.get_blob_client(blob=blob_name)
        blob_client.upload_blob(
            data,
            overwrite=True,
            content_settings=ContentSettings(content_type=content_type or "application/octet-stream"),
        )
        return blob_client.url, len(data)
    except Exception as e:
        logger.error("Azure upload failed: %s", e, exc_info=True)
        # Match Node: throw so caller returns a 500 in the controller.
        raise


async def _upload_generated_bytes_to_azure(
    *,
    content: bytes,
    filename: str,
    study_id: str,
    site_id: str,
    document_type: str,
    content_type: str,
) -> Dict[str, Any]:
    """
    Upload generated cover-letter assets (DOCX/PDF) and return document metadata.
    """
    container = _get_azure_sitepackages_container()
    safe_filename = (filename or "generated-file").strip() or "generated-file"
    blob_path = (
        f"site-packages/generated/{study_id}/{site_id}/"
        f"{(document_type or 'generated').strip().lower()}/{int(time.time() * 1000)}-{safe_filename}"
    )

    url, size = _upload_bytes_to_azure(
        content,
        container=container,
        blob_name=blob_path,
        content_type=content_type,
    )

    return {
        "name": safe_filename,
        "fileUrl": url,
        "fileSize": int(size or 0),
        "mimeType": content_type,
        "blobPath": blob_path,
        "container": container,
    }


"""
Verify PER-SIGNER PLACEHOLDER signature placement: {{SIGNATURE_<ROLE>}}.

Covers the whole feature on a real create->convert->sign path:
  PART 1  create-time preservation: a template authored with {{SIGNATURE_DIRECTOR}},
          {{SIGNATURE_PI}}, {{SIGNATURE_VP}} run through replace_placeholders_in_docx
          (the exact create path) still contains those tokens; the PDF the converter
          produces still shows them.
  PART 2  per-signer stamping: each signer's mark lands AT its own placeholder
          (Director@DIRECTOR, PI@PI, VP@VP), the literal token is COVERED (white rect
          over its box), and the signer's name is drawn at that spot.
  PART 3  precedence + fallback: the template also has a 'For Site' Signature line, but
          the Director (site family) is placed at its PLACEHOLDER, not the block line
          (placeholder wins). A template with NO {{SIGNATURE_*}} resolves no placeholder
          anchor -> the label/fallback path is used.
  PART 4  compliance shape: the produced PDF is hashable exactly as record_signature
          hashes it (placement only).

Run in the backend container. Throwaway; cleans up its temp dir.
"""
import asyncio
import base64
import hashlib
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, "/app")

from app.modules.agreements.utils.pdf_export import (  # noqa: E402
    _placeholder_tokens_for_role, _resolve_placeholder_anchor,
    _resolve_signature_anchor, embed_signature_into_pdf,
)
from app.modules.agreements.services import doc_convert  # noqa: E402
from app.utils.signature_image import render_signature_png_bytes  # noqa: E402
from app.utils.docx_placeholder_replace import replace_placeholders_in_docx  # noqa: E402

results = []


def ok(label, detail=""):
    results.append((label, detail))


class StubProfile:
    hospital_name = "Mercy General Hospital"
    authorized_signatory_name = "Dr. Jane Site"
    authorized_signatory_title = "Site Director"
    primary_contracting_entity = "Dizzaroo Sponsor Inc."


def author_template(path):
    """A template authored with per-signer placeholders + a 'For Site' label block."""
    import docx
    d = docx.Document()
    d.add_heading("Clinical Trial Agreement", level=1)
    d.add_paragraph("This agreement is entered into by the parties below.")
    d.add_heading("10. Signatures", level=2)
    # A label-style block too, so we can prove the placeholder WINS over it.
    d.add_paragraph("For Site")
    d.add_paragraph("Signed for and on behalf of: Mercy General Hospital")
    d.add_paragraph("Name: Dr. Jane Site")
    d.add_paragraph("Title: Site Director")
    d.add_paragraph("Signature:")
    d.add_paragraph("Date:")
    d.add_paragraph("")
    # Per-signer placeholders — one per signer, each on its own line.
    d.add_paragraph("Site Director: {{SIGNATURE_DIRECTOR}}")
    d.add_paragraph("Principal Investigator: {{SIGNATURE_PI}}")
    d.add_paragraph("Sponsor VP: {{SIGNATURE_VP}}")
    d.save(path)


def extract_placeholders(docx_path):
    import docx
    import re
    d = docx.Document(docx_path)
    txt = "\n".join(p.text for p in d.paragraphs)
    for t in d.tables:
        for r in t.rows:
            for cell in r.cells:
                txt += "\n" + cell.text
    return set(re.findall(r"\{\{\s*([^{}]+?)\s*\}\}", txt))


async def main():
    work = tempfile.mkdtemp(prefix="phsign_")
    tpl = str(Path(work) / "template.docx")
    out = str(Path(work) / "created.docx")
    author_template(tpl)

    # --- token candidate mapping (reuses role table) ---
    assert "SIGNATURE_DIRECTOR" in _placeholder_tokens_for_role("director")
    assert "SIGNATURE_PI" in _placeholder_tokens_for_role("pi")
    assert "SIGNATURE_VP" in _placeholder_tokens_for_role("vp")
    # pi also gains aliases derived from the existing role table (e.g. INVESTIGATOR).
    assert "SIGNATURE_INVESTIGATOR" in _placeholder_tokens_for_role("pi")
    ok("Role->placeholder map honors {{SIGNATURE_<ROLE>}} + reuses role-table aliases",
       "pi -> " + ",".join(_placeholder_tokens_for_role("pi")))

    # --- PART 1: run the REAL create path, confirm placeholders survive ---
    replace_placeholders_in_docx(Path(tpl), StubProfile(), Path(out),
                                 sponsor_signatory_name="Sam Sponsor")
    survived = extract_placeholders(out)
    for tok in ("SIGNATURE_DIRECTOR", "SIGNATURE_PI", "SIGNATURE_VP"):
        assert tok in survived, (f"{tok} did not survive create", survived)
    ok("PART 1: {{SIGNATURE_DIRECTOR/PI/VP}} PRESERVED through create (not expanded)",
       ", ".join(sorted(t for t in survived if t.startswith("SIGNATURE_"))))

    pdf = await doc_convert.to_pdf(out)
    assert Path(pdf).exists()
    import pdfplumber
    with pdfplumber.open(pdf) as pp:
        pre_text = "\n".join((p.extract_text() or "") for p in pp.pages)
    assert "{{SIGNATURE_DIRECTOR}}" in pre_text, "placeholder must render in the pre-sign PDF"
    ok("Created PDF shows the literal {{SIGNATURE_*}} anchors before signing", "")

    # --- PART 2/3: resolve each role's placeholder anchor (distinct spots) ---
    a_dir = _resolve_placeholder_anchor(pdf, "director")
    a_pi = _resolve_placeholder_anchor(pdf, "pi")
    a_vp = _resolve_placeholder_anchor(pdf, "vp")
    print("  placeholder anchors: director=%s pi=%s vp=%s" % (a_dir, a_pi, a_vp))
    assert a_dir and a_pi and a_vp, ("all three placeholders must resolve", a_dir, a_pi, a_vp)
    ys = {round(a_dir[3]), round(a_pi[3]), round(a_vp[3])}
    assert len(ys) == 3, ("each placeholder must be at its OWN distinct spot", ys)
    ok("PART 2: Director/PI/VP each resolve to their OWN {{SIGNATURE_<ROLE>}} anchor",
       f"y: dir={a_dir[3]:.1f}, pi={a_pi[3]:.1f}, vp={a_vp[3]:.1f}")

    # PART 3 precedence: Director is site-family; the label path would target the For
    # Site 'Signature:' line — confirm the placeholder anchor is a DIFFERENT spot, and
    # that embed actually uses the placeholder (Director's name lands at the token y).
    line_dir = _resolve_signature_anchor(pdf, "director")
    assert line_dir is not None, "label/line anchor should exist for the For Site block"
    assert round(line_dir[2]) != round(a_dir[3]), \
        ("placeholder and label spots must differ to prove precedence", line_dir, a_dir)

    def png(name):
        return base64.b64encode(render_signature_png_bytes(name)).decode("ascii")

    cur = pdf
    signers = [("Alice Director", "Site Director", "director", a_dir),
               ("Bob Investigator", "Principal Investigator", "pi", a_pi),
               ("Vera Veep", "VP Clinical", "vp", a_vp)]
    for nm, title, role, _anchor in signers:
        cur = await embed_signature_into_pdf(cur, png(nm), nm, title, role=role,
                                             stack_offset=0, signed_on="2026-06-04")

    # Re-extract: literal tokens COVERED (a white rect over each box) + name at the spot.
    import pdfplumber as _pp
    with _pp.open(cur) as pdf2:
        words, rects = [], []
        for pidx, page in enumerate(pdf2.pages):
            ph = float(page.height)
            for w in page.extract_words() or []:
                words.append((pidx, w["text"], ph - float(w["bottom"]), float(w["x0"])))
            for r in page.rects or []:
                rects.append((pidx, float(r["x0"]), float(r["x1"]),
                              ph - float(r["bottom"]), ph - float(r["top"])))

    def name_y(token):
        ys2 = [y for (_, t, y, _x) in words if token.lower() in t.lower()]
        return ys2[0] if ys2 else None

    def covered(anchor):
        pidx, x0, x1, ybot, ytop = anchor
        cx, cy = (x0 + x1) / 2.0, (ybot + ytop) / 2.0
        return any(rp == pidx and rx0 - 2 <= cx <= rx1 + 2 and ry0 - 2 <= cy <= ry1 + 2
                   for (rp, rx0, rx1, ry0, ry1) in rects)

    for (nm, _t, _r, anchor), short in zip(signers, ("Alice", "Bob", "Vera")):
        ny = name_y(short)
        assert ny is not None, (f"{nm} caption must appear", short)
        assert abs(ny - anchor[3]) < 20.0, (f"{nm} must land AT its placeholder", ny, anchor[3])
        assert covered(anchor), (f"{nm}'s placeholder token must be covered by a rect", anchor)
    ok("PART 2: each signature stamped AT its placeholder; token COVERED; name shown",
       f"Alice@{name_y('Alice'):.1f}, Bob@{name_y('Bob'):.1f}, Vera@{name_y('Vera'):.1f}")

    # --- PART 3 fallback: a template with NO {{SIGNATURE_*}} resolves no placeholder ---
    from reportlab.pdfgen import canvas as rlc
    from reportlab.lib.pagesizes import letter
    plain = str(Path(work) / "plain.pdf")
    c = rlc.Canvas(plain, pagesize=letter)
    c.drawString(72, 700, "Body with no signature placeholders and no labels.")
    c.showPage(); c.save()
    assert _resolve_placeholder_anchor(plain, "director") is None, "no-placeholder template must yield None"
    fb = await embed_signature_into_pdf(plain, png("No Anchor"), "No Anchor", "Signer",
                                        role="director", stack_offset=0, signed_on="2026-06-04")
    with _pp.open(fb) as pdf3:
        ftext = "\n".join((p.extract_text() or "") for p in pdf3.pages)
    assert "Electronically signed by No Anchor" in ftext, ("must fall back + still stamp", ftext[-200:])
    ok("PART 3: no placeholder -> falls through to label/labelled-fallback, still stamps", "")

    # --- PART 4 compliance shape ---
    h = hashlib.sha256(Path(cur).read_bytes()).hexdigest()
    ok("PART 4: produced PDF hashable as-is (record_signature hashes these bytes)", h[:16] + "…")

    import shutil
    shutil.rmtree(work, ignore_errors=True)
    for f in Path(pdf).parent.glob("created*_signed*.pdf"):
        try:
            f.unlink()
        except Exception:
            pass

    print("=" * 92)
    print("PER-SIGNER PLACEHOLDER SIGNATURE PLACEMENT — VERIFICATION")
    print("=" * 92)
    for label, detail in results:
        print(f"PASS  {label}" + (f"\n        -> {detail}" if detail else ""))
    print("=" * 92)
    print(f"{len(results)} checks passed")


if __name__ == "__main__":
    asyncio.run(main())

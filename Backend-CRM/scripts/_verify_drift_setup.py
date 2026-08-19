"""CHECK 2 — #6 source-drift: fixture setup.

Builds a real agreement whose latest .docx has two tagged content-control fields:
  - PI_NAME       : source-backed (DEFAULT_TOKEN_MAP -> site_profile.pi_name),
                    set to an EDITED value different from the SiteProfile's pi_name.
  - SPONSOR_NAME  : free-text (no source) — must NOT be flagged.
Prints AGREEMENT_ID + DOC_PATH so the GET/POST source-drift endpoints can be curled.
"""
import asyncio
import tempfile
import uuid

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.db import AsyncSessionLocal
from app.models import (
    Study, Site, StudySite, StudyTemplate, Agreement, AgreementDocument,
    AgreementStatus, TemplateType,
)
from app.models.agreement import AgreementSignatureSource
from app.models.site_status import SiteProfile

MARK = "DRIFTTEST"
SOURCE_PI_NAME = "Dr Source Truth"
EDITED_PI_NAME = "Dr Edited In Doc"


def _tmpfile(suffix: str, prefix: str = "tmp") -> str:
    """NamedTemporaryFile(delete=False) path. Called via asyncio.to_thread from async code."""
    tmp = tempfile.NamedTemporaryFile(delete=False, prefix=prefix, suffix=suffix)
    tmp.close()
    return tmp.name


def _cc(paragraph, tag, value):
    p = paragraph._p
    sdt = OxmlElement("w:sdt")
    pr = OxmlElement("w:sdtPr")
    tag_el = OxmlElement("w:tag"); tag_el.set(qn("w:val"), tag); pr.append(tag_el)
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r"); t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve"); t.text = value; run.append(t)
    content.append(run); sdt.append(pr); sdt.append(content); p.append(sdt)


async def main():
    async with AsyncSessionLocal() as db:
        study = Study(id=uuid.uuid4(), study_id=f"{MARK}-STU-{uuid.uuid4().hex[:6]}", name=f"{MARK} Study")
        site = Site(id=uuid.uuid4(), site_id=f"{MARK}-SITE-{uuid.uuid4().hex[:6]}", name=f"{MARK} Site")
        db.add_all([study, site]); await db.flush()
        ss = StudySite(id=uuid.uuid4(), study_id=study.id, site_id=site.id)
        db.add(ss); await db.flush()
        # SiteProfile is the SOURCE for PI_NAME
        db.add(SiteProfile(id=uuid.uuid4(), site_id=site.id, pi_name=SOURCE_PI_NAME))
        tpl = StudyTemplate(id=uuid.uuid4(), study_id=study.id, template_name=f"{MARK} tpl",
                            template_type=TemplateType.CTA, is_active="true", field_mappings={})
        db.add(tpl); await db.flush()
        ag = Agreement(id=uuid.uuid4(), site_id=site.id, study_id=study.id, study_site_id=ss.id,
                       title=f"{MARK} agreement", status=AgreementStatus.DRAFT, is_legacy="false",
                       signature_source=AgreementSignatureSource.INTERNAL, agreement_type=TemplateType.CTA)
        db.add(ag); await db.flush()

        # Build the DOCX with content controls; PI_NAME edited != source, SPONSOR_NAME free-text.
        doc_path = await asyncio.to_thread(_tmpfile, ".docx", "drift_")
        doc = Document()
        p1 = doc.add_paragraph("PI: "); _cc(p1, "PI_NAME", EDITED_PI_NAME)
        p2 = doc.add_paragraph("Sponsor: "); _cc(p2, "SPONSOR_NAME", "Acme Pharma")
        doc.save(doc_path)

        db.add(AgreementDocument(id=uuid.uuid4(), agreement_id=ag.id, version_number=1,
                                 is_signed_version="false", created_from_template_id=tpl.id,
                                 document_file_path=doc_path, document_file_url=None))
        await db.commit()

        print(f"AGREEMENT_ID={ag.id}")
        print(f"DOC_PATH={doc_path}")
        print(f"SOURCE_PI_NAME={SOURCE_PI_NAME!r}  EDITED_PI_NAME={EDITED_PI_NAME!r}")


if __name__ == "__main__":
    asyncio.run(main())

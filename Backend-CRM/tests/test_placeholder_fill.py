"""Unit tests for the shared placeholder-fill capability.

`app/modules/agreements/services/placeholder_fill.py` is the single home for the
``{{token}}`` machinery (lifted out of the retired ``types/cta/service.py`` during
the 2026-06-18 legacy retirement). These tests re-establish the coverage that was
lost when ``test_cta_workflow.py`` was deleted, but against the SHARED module.

Design: no live DB. Source resolution is exercised with a tiny fake async session
(a queue of rows) and SimpleNamespace stubs, so the tests are deterministic and
schema-tolerant — matching the module's own ``getattr``-based, exception-swallowing
behaviour. DOCX behaviours use real python-docx files in ``tmp_path``; the
content-control (``w:sdt``) drift fields are built directly via oxml.

`asyncio_mode = auto` (pytest.ini) runs the async tests without markers.
"""
from types import SimpleNamespace

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.modules.agreements.services import placeholder_fill as pf


# ---------------------------------------------------------------------------
# Test doubles
# ---------------------------------------------------------------------------

class _Result:
    def __init__(self, row):
        self._row = row

    def scalar_one_or_none(self):
        return self._row


class _FakeSession:
    """Async session stub: each ``execute`` pops the next queued row's result."""
    def __init__(self, *rows):
        self._q = list(rows)
        self.calls = 0

    async def execute(self, *_a, **_k):
        self.calls += 1
        return _Result(self._q.pop(0) if self._q else None)


def _make_docx(path, paragraphs=None, table_cells=None):
    """Build a DOCX with the given body paragraphs and a one-row table whose
    cells hold ``table_cells`` text. Returns the saved path (str)."""
    doc = Document()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    if table_cells:
        table = doc.add_table(rows=1, cols=len(table_cells))
        for cell, text in zip(table.rows[0].cells, table_cells):
            cell.text = text
    doc.save(str(path))
    return str(path)


def _add_content_control(doc, tag, value):
    """Append a block-level tagged content control (w:sdt) to the body — the
    on-disk shape produced when a source-backed placeholder is filled. ``tag`` is
    stored verbatim; the drift readers upper-case it on read."""
    sdt = OxmlElement('w:sdt')
    sdt_pr = OxmlElement('w:sdtPr')
    tag_el = OxmlElement('w:tag')
    tag_el.set(qn('w:val'), tag)
    sdt_pr.append(tag_el)
    sdt.append(sdt_pr)

    content = OxmlElement('w:sdtContent')
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = value
    r.append(t)
    p.append(r)
    content.append(p)
    sdt.append(content)

    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        sect_pr.addprevious(sdt)
    else:
        body.append(sdt)


def _make_cc_docx(path, **tag_values):
    """Build a DOCX whose body holds one tagged content control per kwarg."""
    doc = Document()
    for tag, value in tag_values.items():
        _add_content_control(doc, tag, value)
    doc.save(str(path))
    return str(path)


# ---------------------------------------------------------------------------
# 1. scan_placeholders_in_docx
# ---------------------------------------------------------------------------

async def test_scan_finds_tokens_dedup_and_order(tmp_path):
    path = _make_docx(
        tmp_path / "scan.docx",
        paragraphs=[
            "Site: {{site_name}}",
            "Coordinator emails {{ pi_email }} weekly.",  # whitespace variant
            "Confirm {{site_name}} again.",               # duplicate -> not repeated
            "Effective {{Effective_Date}}.",              # case preserved verbatim
        ],
        table_cells=["Study {{study_name}}"],
    )
    tokens = await pf.scan_placeholders_in_docx(path)

    # unique, first-occurrence order; whitespace stripped; case kept verbatim
    assert tokens == ["site_name", "pi_email", "Effective_Date", "study_name"]


async def test_scan_whitespace_variants_collapse_to_same_token(tmp_path):
    path = _make_docx(
        tmp_path / "ws.docx",
        paragraphs=["{{site_name}}", "{{   site_name   }}"],
    )
    tokens = await pf.scan_placeholders_in_docx(path)
    assert tokens == ["site_name"]


async def test_scan_no_tokens_returns_empty(tmp_path):
    path = _make_docx(tmp_path / "plain.docx", paragraphs=["No placeholders here."])
    assert await pf.scan_placeholders_in_docx(path) == []


# ---------------------------------------------------------------------------
# 2. resolve_field_value — the 4 sources + graceful unknowns
# ---------------------------------------------------------------------------

async def test_resolve_agreement_source_reads_attribute():
    ag = SimpleNamespace(title="My CTA", site_id=None, study_id=None)
    assert await pf.resolve_field_value("agreement.title", ag, _FakeSession()) == "My CTA"


async def test_resolve_site_profile_field():
    ag = SimpleNamespace(site_id="site-1")
    db = _FakeSession(SimpleNamespace(pi_name="Dr. Who"))
    assert await pf.resolve_field_value("site_profile.pi_name", ag, db) == "Dr. Who"


async def test_resolve_site_profile_address_is_composed():
    ag = SimpleNamespace(site_id="site-1")
    db = _FakeSession(SimpleNamespace(
        address_line_1="123 Main", city="Boston", state="MA", country="USA",
    ))
    assert await pf.resolve_field_value("site_profile.address", ag, db) == "123 Main, Boston, MA, USA"


async def test_resolve_study_field():
    ag = SimpleNamespace(site_id=None, study_id="study-1")
    db = _FakeSession(SimpleNamespace(name="ONCO Study"))
    assert await pf.resolve_field_value("study.name", ag, db) == "ONCO Study"


async def test_resolve_irb_field_two_hop():
    # IRB resolution is a two-query hop: site->mapping, then mapping->IRB row.
    ag = SimpleNamespace(site_id="site-1")
    db = _FakeSession(SimpleNamespace(irb_id="irb-1"), SimpleNamespace(name="Central IRB"))
    assert await pf.resolve_field_value("irb.name", ag, db) == "Central IRB"


async def test_resolve_unknown_source_returns_none_not_error():
    ag = SimpleNamespace(site_id="site-1", study_id="study-1")
    assert await pf.resolve_field_value("weird.field", ag, _FakeSession()) is None


async def test_resolve_malformed_path_returns_none():
    ag = SimpleNamespace(site_id="site-1")
    assert await pf.resolve_field_value("nodothere", ag, _FakeSession()) is None
    assert await pf.resolve_field_value("", ag, _FakeSession()) is None


async def test_resolve_site_profile_missing_site_id_returns_none_without_db():
    ag = SimpleNamespace(site_id=None)
    db = _FakeSession()
    assert await pf.resolve_field_value("site_profile.pi_name", ag, db) is None
    assert db.calls == 0  # short-circuits before touching the DB


async def test_resolve_missing_row_returns_none():
    ag = SimpleNamespace(site_id="site-1")
    db = _FakeSession(None)  # no SiteProfile row found
    assert await pf.resolve_field_value("site_profile.pi_name", ag, db) is None


async def test_resolve_unknown_field_on_present_row_returns_none():
    ag = SimpleNamespace(site_id="site-1")
    db = _FakeSession(SimpleNamespace(pi_name="Dr. Who"))  # no `nonexistent` attr
    assert await pf.resolve_field_value("site_profile.nonexistent", ag, db) is None


# ---------------------------------------------------------------------------
# 3. resolve_field_mapping_values — UPPER-keyed, known resolve, unknown dropped
# ---------------------------------------------------------------------------

async def test_field_mapping_values_resolve_known_drop_unknown():
    # agreement-sourced mappings need no DB; site/study/irb default tokens all
    # resolve to None here (stub has no site_id/study_id), so they drop out.
    ag = SimpleNamespace(site_id=None, study_id=None, title="T1", pi="Dr. A")
    field_mappings = {
        "doc_title": "agreement.title",
        "pi_name": "agreement.pi",
        "unknown_tok": "agreement.does_not_exist",  # unresolved -> omitted, no error
    }
    out = await pf.resolve_field_mapping_values(field_mappings, ag, _FakeSession())

    assert out["DOC_TITLE"] == "T1"          # UPPER-cased key
    assert out["PI_NAME"] == "Dr. A"
    assert "UNKNOWN_TOK" not in out           # unmapped left out, not errored
    # default site/study tokens can't resolve against this stub -> absent
    assert "SITE_NAME" not in out
    assert "STUDY_NAME" not in out


async def test_field_mapping_keys_are_trimmed_and_uppercased():
    ag = SimpleNamespace(site_id=None, study_id=None, title="Trimmed")
    out = await pf.resolve_field_mapping_values({"  Doc_Title  ": "agreement.title"}, ag, _FakeSession())
    assert out.get("DOC_TITLE") == "Trimmed"


async def test_field_mapping_handles_none_mappings():
    ag = SimpleNamespace(site_id=None, study_id=None)
    # None field_mappings + a stub that resolves nothing -> empty dict, no crash
    assert await pf.resolve_field_mapping_values(None, ag, _FakeSession()) == {}


# ---------------------------------------------------------------------------
# 4. apply_placeholders_to_docx — literal find/replace incl. whitespace variant
# ---------------------------------------------------------------------------

async def test_apply_replaces_both_brace_variants_and_tables(tmp_path):
    path = _make_docx(
        tmp_path / "apply.docx",
        paragraphs=["Site: {{site_name}}", "PI: {{ pi_name }}"],  # tight + spaced
        table_cells=["Study {{study_name}}"],
    )
    mapping = {"site_name": "Acme Site", "pi_name": "Dr. Smith", "study_name": "ONCO-1"}
    out_path = await pf.apply_placeholders_to_docx(path, mapping)

    assert out_path.endswith("_filled.docx")
    doc = Document(out_path)
    body = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(
        c.text for t in doc.tables for row in t.rows for c in row.cells
    )
    full = body + "\n" + table_text
    assert "Acme Site" in full
    assert "Dr. Smith" in full          # {{ pi_name }} (spaced variant) replaced
    assert "ONCO-1" in table_text       # table cell replaced
    assert "{{" not in full             # nothing left unfilled


async def test_apply_leaves_unmapped_tokens_intact(tmp_path):
    path = _make_docx(tmp_path / "partial.docx", paragraphs=["{{site_name}} / {{unmapped}}"])
    out_path = await pf.apply_placeholders_to_docx(path, {"site_name": "Acme"})
    text = "\n".join(p.text for p in Document(out_path).paragraphs)
    assert "Acme" in text
    assert "{{unmapped}}" in text       # token with no mapping is untouched


# ---------------------------------------------------------------------------
# 5. Source-drift: read / detect / set against tagged content controls
# ---------------------------------------------------------------------------

def test_read_filled_field_values_upper_keys(tmp_path):
    path = _make_cc_docx(tmp_path / "read.docx", DOC_TITLE="Alpha", site_name="Beta")
    out = pf.read_filled_field_values(path)
    assert out == {"DOC_TITLE": "Alpha", "SITE_NAME": "Beta"}  # tags upper-cased


def test_read_filled_field_values_missing_file_returns_empty():
    assert pf.read_filled_field_values("/no/such/file.docx") == {}


async def test_detect_drift_in_sync(tmp_path):
    path = _make_cc_docx(tmp_path / "sync.docx", DOC_TITLE="Live Title")
    ag = SimpleNamespace(title="Live Title", site_id=None, study_id=None)
    result = await pf.detect_source_drift(path, {"DOC_TITLE": "agreement.title"}, ag, _FakeSession())
    assert result["in_sync"] == ["DOC_TITLE"]
    assert result["drifted"] == []


async def test_detect_drift_edited_value_flagged(tmp_path):
    path = _make_cc_docx(tmp_path / "drift.docx", DOC_TITLE="Edited Title")
    ag = SimpleNamespace(title="Live Title", site_id=None, study_id=None)
    result = await pf.detect_source_drift(path, {"DOC_TITLE": "agreement.title"}, ag, _FakeSession())
    assert result["in_sync"] == []
    assert len(result["drifted"]) == 1
    drift = result["drifted"][0]
    assert drift["token"] == "DOC_TITLE"
    assert drift["doc_value"] == "Edited Title"
    assert drift["source_value"] == "Live Title"


async def test_detect_drift_ignores_free_text_no_source_fields(tmp_path):
    # SPONSOR_NAME has no source mapping (and isn't in DEFAULT_TOKEN_MAP) -> a
    # free-text field; it must NOT be flagged as in_sync or drifted.
    path = _make_cc_docx(tmp_path / "freetext.docx", DOC_TITLE="Live Title", SPONSOR_NAME="Acme Pharma")
    ag = SimpleNamespace(title="Live Title", site_id=None, study_id=None)
    result = await pf.detect_source_drift(path, {"DOC_TITLE": "agreement.title"}, ag, _FakeSession())
    assert result["in_sync"] == ["DOC_TITLE"]
    drifted_tokens = {d["token"] for d in result["drifted"]}
    assert "SPONSOR_NAME" not in result["in_sync"]
    assert "SPONSOR_NAME" not in drifted_tokens


def test_set_filled_field_values_updates_target_only(tmp_path):
    path = _make_cc_docx(tmp_path / "set.docx", DOC_TITLE="Old", SPONSOR_NAME="KeepMe")
    out_path = pf.set_filled_field_values(path, {"doc_title": "New Value"})  # lower key normalises
    assert out_path.endswith("_resynced.docx")
    values = pf.read_filled_field_values(out_path)
    assert values["DOC_TITLE"] == "New Value"    # pulled the source value in
    assert values["SPONSOR_NAME"] == "KeepMe"    # untouched free-text field


def test_set_filled_field_values_empty_updates_is_noop(tmp_path):
    path = _make_cc_docx(tmp_path / "noop.docx", DOC_TITLE="Same")
    assert pf.set_filled_field_values(path, {}) == path

"""
Document Comparison Utility

Compares DOCX documents to detect changes (additions, deletions, modifications).
Used for tracking changes during agreement review/negotiation.

IMPORTANT: Uses lxml to extract ALL text including from content controls (w:sdt elements),
because python-docx's paragraph.text property only reads direct w:r runs and misses
text inside content controls entirely.
"""

import logging
import re
from pathlib import Path
from typing import List, Optional
from docx import Document

logger = logging.getLogger(__name__)

# WordprocessingML namespace
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W_T = f'{{{WORD_NS}}}t'


class DocumentChange:
    """Represents a single change detected in a document."""
    def __init__(
        self,
        change_type: str,  # 'ADDED', 'DELETED', 'MODIFIED'
        old_text: Optional[str] = None,
        new_text: Optional[str] = None,
        paragraph_index: Optional[int] = None,
        table_index: Optional[int] = None,
        row_index: Optional[int] = None,
        cell_index: Optional[int] = None,
        section_reference: Optional[str] = None,
    ):
        self.change_type = change_type
        self.old_text = old_text
        self.new_text = new_text
        self.paragraph_index = paragraph_index
        self.table_index = table_index
        self.row_index = row_index
        self.cell_index = cell_index
        self.section_reference = section_reference

    def to_dict(self):
        return {
            "change_type": self.change_type,
            "old_text": self.old_text,
            "new_text": self.new_text,
            "paragraph_index": self.paragraph_index,
            "table_index": self.table_index,
            "row_index": self.row_index,
            "cell_index": self.cell_index,
            "section_reference": self.section_reference,
        }


def _get_element_full_text(element) -> str:
    """
    Extract ALL text from an XML element by iterating all w:t descendants.

    This is the CORRECT way to extract text — it handles:
    - Normal runs: <w:r><w:t>text</w:t></w:r>
    - Content controls: <w:sdt><w:sdtContent><w:r><w:t>text</w:t></w:r></w:sdtContent></w:sdt>
    - Nested structures of any depth

    python-docx's paragraph.text ONLY reads direct w:r runs and misses
    content inside w:sdt (content controls), so we must use lxml directly.
    """
    texts = []
    for t_elem in element.iter(W_T):
        if t_elem.text:
            texts.append(t_elem.text)
        # Also get tail text if any
        if t_elem.tail:
            texts.append(t_elem.tail.strip())
    return ''.join(texts).strip()


def _extract_all_paragraphs(doc: Document) -> List[str]:
    """
    Extract text from all paragraphs in a document.
    Includes text inside content controls (w:sdt elements).
    Returns list of non-empty paragraph texts.
    """
    texts = []
    for para in doc.paragraphs:
        text = _get_element_full_text(para._element)
        texts.append(text)  # Keep empty ones to preserve index alignment
    return texts


def _extract_all_paragraph_texts_flat(doc: Document) -> List[str]:
    """
    Extract all non-empty paragraph texts for comparison.
    """
    return [t for t in _extract_all_paragraphs(doc) if t]


def compare_documents(
    original_docx_path: Path,
    modified_docx_path: Path,
) -> List[DocumentChange]:
    """
    Compare two DOCX documents and detect changes.

    Uses lxml text extraction to include text inside content controls (w:sdt),
    which python-docx's paragraph.text misses entirely.
    """
    try:
        original_doc = Document(original_docx_path)
        modified_doc = Document(modified_docx_path)

        # Extract all paragraph texts using proper full-text extraction
        original_paras = _extract_all_paragraphs(original_doc)
        modified_paras = _extract_all_paragraphs(modified_doc)

        # Log for debugging
        logger.info(
            f"Comparison: original has {len(original_paras)} paragraphs "
            f"({len([t for t in original_paras if t])} non-empty), "
            f"modified has {len(modified_paras)} paragraphs "
            f"({len([t for t in modified_paras if t])} non-empty)"
        )

        # Log first few paragraphs from each doc to debug
        orig_sample = [t for t in original_paras if t][:5]
        mod_sample = [t for t in modified_paras if t][:5]
        logger.info(f"Original first 5 paragraphs: {orig_sample}")
        logger.info(f"Modified first 5 paragraphs: {mod_sample}")

        changes = []

        # --- Paragraph-level comparison ---
        changes.extend(_compare_paragraph_lists(original_paras, modified_paras))

        # --- Table-level comparison ---
        changes.extend(_compare_tables(original_doc, modified_doc))

        logger.info(f"Detected {len(changes)} changes between documents")
        return changes

    except Exception as e:
        logger.error(f"Error comparing documents: {str(e)}", exc_info=True)
        raise Exception(f"Failed to compare documents: {str(e)}")


def _compare_paragraph_lists(
    original_paras: List[str],
    modified_paras: List[str],
) -> List[DocumentChange]:
    """
    Compare two lists of paragraph texts and detect changes.
    Uses difflib for accurate line-by-line diffing.
    """
    import difflib
    changes = []

    # Use difflib SequenceMatcher for proper diff
    matcher = difflib.SequenceMatcher(
        None,
        original_paras,
        modified_paras,
        autojunk=False
    )

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue  # No change

        elif tag == 'replace':
            # Lines changed — emit MODIFIED for pairs, or DELETED+ADDED for extras
            orig_block = [t for t in original_paras[i1:i2] if t]
            mod_block = [t for t in modified_paras[j1:j2] if t]
            max_len = max(len(orig_block), len(mod_block))

            for k in range(max_len):
                orig_text = orig_block[k] if k < len(orig_block) else None
                mod_text = mod_block[k] if k < len(mod_block) else None

                if orig_text and mod_text:
                    changes.append(DocumentChange(
                        change_type="MODIFIED",
                        old_text=orig_text,
                        new_text=mod_text,
                        paragraph_index=i1 + k,
                        section_reference=_extract_section_reference(mod_text),
                    ))
                elif orig_text and not mod_text:
                    changes.append(DocumentChange(
                        change_type="DELETED",
                        old_text=orig_text,
                        paragraph_index=i1 + k,
                        section_reference=_extract_section_reference(orig_text),
                    ))
                elif not orig_text and mod_text:
                    changes.append(DocumentChange(
                        change_type="ADDED",
                        new_text=mod_text,
                        paragraph_index=j1 + k,
                        section_reference=_extract_section_reference(mod_text),
                    ))

        elif tag == 'delete':
            for k, text in enumerate(original_paras[i1:i2]):
                if text:
                    changes.append(DocumentChange(
                        change_type="DELETED",
                        old_text=text,
                        paragraph_index=i1 + k,
                        section_reference=_extract_section_reference(text),
                    ))

        elif tag == 'insert':
            for k, text in enumerate(modified_paras[j1:j2]):
                if text:
                    changes.append(DocumentChange(
                        change_type="ADDED",
                        new_text=text,
                        paragraph_index=j1 + k,
                        section_reference=_extract_section_reference(text),
                    ))

    return changes


def _compare_tables(
    original_doc: Document,
    modified_doc: Document,
) -> List[DocumentChange]:
    """Compare tables between two documents using full text extraction."""
    changes = []

    original_tables = original_doc.tables
    modified_tables = modified_doc.tables

    if not original_tables and not modified_tables:
        return changes

    max_tables = max(len(original_tables), len(modified_tables))

    for table_idx in range(max_tables):
        orig_table = original_tables[table_idx] if table_idx < len(original_tables) else None
        mod_table = modified_tables[table_idx] if table_idx < len(modified_tables) else None

        if not orig_table and mod_table:
            for row_idx, row in enumerate(mod_table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = _get_element_full_text(cell._element)
                    if cell_text:
                        changes.append(DocumentChange(
                            change_type="ADDED",
                            new_text=cell_text,
                            table_index=table_idx,
                            row_index=row_idx,
                            cell_index=cell_idx,
                        ))
        elif orig_table and not mod_table:
            for row_idx, row in enumerate(orig_table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = _get_element_full_text(cell._element)
                    if cell_text:
                        changes.append(DocumentChange(
                            change_type="DELETED",
                            old_text=cell_text,
                            table_index=table_idx,
                            row_index=row_idx,
                            cell_index=cell_idx,
                        ))
        elif orig_table and mod_table:
            max_rows = max(len(orig_table.rows), len(mod_table.rows))
            for row_idx in range(max_rows):
                orig_row = orig_table.rows[row_idx] if row_idx < len(orig_table.rows) else None
                mod_row = mod_table.rows[row_idx] if row_idx < len(mod_table.rows) else None

                if orig_row and mod_row:
                    max_cells = max(len(orig_row.cells), len(mod_row.cells))
                    for cell_idx in range(max_cells):
                        orig_text = _get_element_full_text(orig_row.cells[cell_idx]._element) if cell_idx < len(orig_row.cells) else None
                        mod_text = _get_element_full_text(mod_row.cells[cell_idx]._element) if cell_idx < len(mod_row.cells) else None

                        if orig_text and mod_text and orig_text != mod_text:
                            changes.append(DocumentChange(
                                change_type="MODIFIED",
                                old_text=orig_text,
                                new_text=mod_text,
                                table_index=table_idx,
                                row_index=row_idx,
                                cell_index=cell_idx,
                            ))
                        elif orig_text and not mod_text:
                            changes.append(DocumentChange(
                                change_type="DELETED",
                                old_text=orig_text,
                                table_index=table_idx,
                                row_index=row_idx,
                                cell_index=cell_idx,
                            ))
                        elif not orig_text and mod_text:
                            changes.append(DocumentChange(
                                change_type="ADDED",
                                new_text=mod_text,
                                table_index=table_idx,
                                row_index=row_idx,
                                cell_index=cell_idx,
                            ))

    return changes


def _extract_section_reference(text: str) -> Optional[str]:
    """Extract section reference from text."""
    if not text:
        return None
    patterns = [
        r'(?:Section|Clause|Article|Paragraph)\s+(\d+(?:\.\d+)?)',
        r'(\d+)\.\s+[A-Z]',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0)
    return None

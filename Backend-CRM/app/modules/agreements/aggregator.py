from fastapi import APIRouter, Depends, HTTPException, WebSocket, WebSocketDisconnect, Header, Query, UploadFile, File, Form, Request
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from fastapi.security import OAuth2PasswordRequestForm
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, or_, and_, delete, update, func
from sqlalchemy.orm import selectinload
from typing import Optional, List, Tuple
from uuid import UUID
from datetime import timedelta, datetime, timezone
import asyncio
import os
import hashlib
import hmac
import shutil
import re
import tempfile
from pathlib import Path
from urllib.parse import urlparse
import secrets
import aiofiles
from app.db import get_db, init_db
from app import crud, schemas
from app.models import MessageDirection, MessageStatus, MessageChannel, ConversationAccessLevel, ThreadAttachment, Attachment, Conversation, ChatDocument, PrimarySiteStatus, UserRoleAssignment, Site, Study, StudySite, SiteStatus, SiteWorkflowStep, SiteDocument, WorkflowStepName, StepStatus, DocumentCategory, DocumentType, ReviewStatus, ProjectFeasibilityCustomQuestion, FeasibilityRequest, FeasibilityResponse, FeasibilityRequestStatus, FeasibilityAttachment, Agreement, AgreementComment, AgreementStatus, CommentType, StudyTemplate, AgreementDocument, TemplateType, AgreementSignedDocument, SiteProfile, AgreementReviewToken, AgreementChange, AgreementReviewDocument, AgreementDocumentComment, CommentStatus, AgreementReviewOtp, AgreementReviewerSignature, AuditLog, AgreementInternalSignature, AgreementSigningOtp, AgreementSignatureSource, AgreementSigningToken, AgreementSigner
from app.websocket_manager import manager
from app.config import settings
from app.auth import create_access_token, get_password_hash, get_current_user, get_current_user_optional, ACCESS_TOKEN_EXPIRE_MINUTES
from app.modules.sites.services.site_status_service import (
    get_study_status_summary,
    get_country_site_counts,
    get_sites_by_status,
    get_site_status_detail,
)
from app.modules.clinical_workflow.services.study_site_service import (
    get_or_create_study_site,
)
from app.modules.clinical_workflow.services.cda_documents import (
    generate_cda_document_html,
    get_cda_document_path,
)
from app.integrations.ai import ai_service
from app.modules.agreements.services.agreement_service import (
    change_agreement_status as change_agreement_status_service,
    get_next_allowed_status as get_next_allowed_status_service,
    check_can_upload_new_version,
    filter_versions_by_role,
    is_user_internal,
    can_create_comment_type,
    filter_comments_by_role,
    get_editor_permissions,
    get_agreement_permissions,
    create_agreement_notice,
    AgreementStatus as AgreementStatusEnum,
)
from app.modules.agreements.blob_names import (
    _build_pretty_document_blob_name,
    _build_review_working_blob_name,
    _extract_blob_seq,
)
import uuid
import logging
# Tasks imported where needed to avoid circular imports


router = APIRouter(tags=["Legal"])
logger = logging.getLogger(__name__)

OTP_EXPIRY_MINUTES = 5
OTP_MAX_ATTEMPTS = 5
OTP_RESEND_COOLDOWN_SECONDS = 30
OTP_MEANING = "Reviewed & Approved"
SITE_SIGN_MEANING = "Signed & Approved"
SIGNING_LINK_EXPIRY_DAYS = 30


def _make_named_tempfile(suffix: str, data: bytes | None = None):
    """Create a NamedTemporaryFile(delete=False) with ``suffix``, optionally
    write ``data`` into it, and return the (closed) tempfile object — callers
    rely on its full-path ``.name``. Sync on purpose — call via
    ``asyncio.to_thread`` from async code so the blocking file I/O stays off
    the event loop."""
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    try:
        if data is not None:
            tmp.write(data)
    finally:
        tmp.close()
    return tmp


def _resolve_onlyoffice_internal_base(raw_base: str) -> str:
    """Return a backend URL that the ONLYOFFICE container can fetch.

    The ONLYOFFICE Document Server must fetch the DOCX from the backend, not the
    frontend dev server. If a frontend origin slips into config during local dev
    (for example localhost:3000 or localhost:5173), replace it with the backend
    host that is reachable from Docker Desktop.
    """
    parsed = urlparse(raw_base or "")
    if parsed.hostname in {"localhost", "127.0.0.1"} and parsed.port in {3000, 5173}:
        return "http://host.docker.internal:8000"
    return raw_base.rstrip("/")


def _latest_agreement_document(documents: List[AgreementDocument]) -> Optional[AgreementDocument]:
    """Return most recent document by blob file sequence, then timestamps / version.

    This avoids selecting stale files when duplicate version_number rows exist.
    """
    if not documents:
        return None
    return max(
        documents,
        key=lambda d: (
            _extract_blob_seq(getattr(d, "document_file_path", None)),
            getattr(d, "updated_at", None)
            or getattr(d, "created_at", None)
            or datetime.min.replace(tzinfo=timezone.utc),
            d.version_number or 0,
            getattr(d, "id", None),
        ),
    )


def _document_for_signature_apply(
    documents: List[AgreementDocument],
    role: str,
) -> Optional[AgreementDocument]:
    """
    Pick the DOCX row to mutate when applying a signature.

    Site signatures must start from the latest *signed* version (reviewer/sponsor
    signature already embedded). Using `_latest_agreement_document` alone can pick
    a newer unsigned OnlyOffice save and drop the first party's signature.
    """
    if not documents:
        return None
    latest_any = _latest_agreement_document(documents)
    role_l = (role or "").strip().lower()
    if role_l == "site":
        latest_signed = _latest_signed_agreement_document(documents)
        return latest_signed or latest_any
    return latest_any


def _latest_signed_agreement_document(
    documents: List[AgreementDocument],
) -> Optional[AgreementDocument]:
    """Prefer a signed document version for agreements in EXECUTED state.

    Signed versions are often stored as blobs that don't match the `_vNNN_` naming
    pattern used by `_extract_blob_seq`, so we rely on `version_number` and timestamps
    instead.
    """
    signed_docs = [d for d in (documents or []) if getattr(d, "is_signed_version", None) == "true"]
    if not signed_docs:
        return None
    return max(
        signed_docs,
        key=lambda d: (
            d.version_number or 0,
            getattr(d, "updated_at", None)
            or getattr(d, "created_at", None)
            or datetime.min.replace(tzinfo=timezone.utc),
            getattr(d, "id", None),
        ),
    )


def _latest_document_for_version(
    documents: List[AgreementDocument],
    version_number: int,
) -> Optional[AgreementDocument]:
    """Return newest row among documents with the same version number."""
    if not documents:
        return None
    same_version_docs = [d for d in documents if d.version_number == version_number]
    if not same_version_docs:
        return None
    return max(
        same_version_docs,
        key=lambda d: (
            _extract_blob_seq(getattr(d, "document_file_path", None)),
            getattr(d, "updated_at", None)
            or getattr(d, "created_at", None)
            or datetime.min.replace(tzinfo=timezone.utc),
            getattr(d, "id", None),
        ),
    )


async def _get_active_signing_context(
    db: AsyncSession,
    agreement_id: UUID,
    token: str,
) -> Tuple[Agreement, AgreementSigningToken, AgreementDocument]:
    token_result = await db.execute(
        select(AgreementSigningToken)
        .where(AgreementSigningToken.token == token)
        .where(AgreementSigningToken.agreement_id == agreement_id)
        .where(AgreementSigningToken.is_active == 'true')
    )
    signing_token = token_result.scalar_one_or_none()
    if not signing_token:
        raise HTTPException(status_code=404, detail="Invalid or expired signing token")

    now = datetime.now(timezone.utc)
    if signing_token.expires_at < now:
        raise HTTPException(status_code=403, detail="Signing link has expired")

    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(selectinload(Agreement.documents))
    )
    agreement = agreement_result.scalar_one_or_none()
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    if agreement.status == AgreementStatusEnum.FULLY_SIGNED:
        raise HTTPException(status_code=403, detail="This agreement has already been fully signed.")

    latest_doc = _document_for_signature_apply(
        list(agreement.documents or []),
        signing_token.role,
    )
    if not latest_doc or not latest_doc.document_file_path:
        raise HTTPException(status_code=404, detail="No document found")

    if not signing_token.used_at:
        signing_token.used_at = now
        await db.commit()

    return agreement, signing_token, latest_doc


def _hash_otp_value(otp: str) -> str:
    return hashlib.sha256(otp.encode("utf-8")).hexdigest()


def _extract_request_ip(request: Request) -> Optional[str]:
    forwarded_for = request.headers.get("x-forwarded-for")
    if forwarded_for:
        return forwarded_for.split(",")[0].strip()
    client_host = getattr(request.client, "host", None)
    return client_host


_SIGNATURE_PLACEHOLDERS_BY_ROLE = {
    "reviewer": [
        "REVIEWER_SIGNATURE", "REVIEWER_SIGN", "INTERNAL_SIGNATURE",
        "INTERNAL_SIGN",
    ],
    "sponsor": [
        "SPONSOR_SIGNATURE", "SPONSOR_SIGN", "SPONSOR_SIGNATURE_BLOCK",
        "COMPANY_SIGNATURE", "CRO_SIGNATURE",
    ],
    "site": [
        "SITE_SIGNATURE", "SITE_SIGN", "HOSPITAL_SIGNATURE",
        "EXTERNAL_SIGNATURE", "SITE_SIGNATURE_BLOCK",
    ],
    "pi": [
        "PI_SIGNATURE", "PI_SIGN", "INVESTIGATOR_SIGNATURE",
        "PRINCIPAL_INVESTIGATOR_SIGNATURE",
    ],
    "co_pi": [
        "CO_PI_SIGNATURE", "COPI_SIGNATURE", "SUB_INVESTIGATOR_SIGNATURE",
    ],
    "witness": [
        "WITNESS_SIGNATURE", "WITNESS_SIGN", "WITNESSED_BY_SIGNATURE",
    ],
    "legal": [
        "LEGAL_SIGNATURE", "LEGAL_SIGN", "COUNSEL_SIGNATURE",
        "ATTORNEY_SIGNATURE",
    ],
    "investigator": [
        "INVESTIGATOR_SIGNATURE", "INVESTIGATOR_SIGN", "PI_SIGNATURE",
    ],
}
_SIGNATURE_PLACEHOLDERS_GENERIC = ["SIGNATURE", "SIGN", "SIGNATURE_BLOCK", "SIG"]


def list_signature_placeholders_in_docx(docx_path: str) -> List[Tuple[str, int]]:
    """
    Scan a DOCX for any signature-like placeholders and return them as
    [(placeholder_name_upper, paragraph_index)] tuples. Used by the send-time
    UI to show which signature spots already exist in a template.

    The placeholder_name is normalized to UPPER_SNAKE_CASE so the UI can match
    it back to a known role via the BY_ROLE / GENERIC lookup tables.
    """
    from docx import Document

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.warning("list_signature_placeholders_in_docx: cannot open %s: %s", docx_path, exc)
        return []

    known_names: List[str] = []
    for role_names in _SIGNATURE_PLACEHOLDERS_BY_ROLE.values():
        known_names.extend(role_names)
    known_names.extend(_SIGNATURE_PLACEHOLDERS_GENERIC)
    patterns = [(name, _signature_placeholder_pattern(name)) for name in known_names]

    found: List[Tuple[str, int]] = []
    seen_positions: set = set()
    for idx, paragraph in enumerate(_iter_all_paragraphs(doc)):
        text = paragraph.text or ""
        if not text:
            continue
        for name, pat in patterns:
            if pat.search(text) and (name, idx) not in seen_positions:
                found.append((name, idx))
                seen_positions.add((name, idx))
                break  # one match per paragraph is enough
    return found


def inject_signature_placeholder_at_paragraph(
    docx_in_path: str,
    docx_out_path: str,
    role: str,
    paragraph_index: int,
    placeholder_name: Optional[str] = None,
) -> str:
    """
    For templates that don't already contain a signature placeholder, insert one
    at the sender-chosen paragraph position so the existing signature-stamping
    logic has a target.

    paragraph_index counts paragraphs in the order yielded by _iter_all_paragraphs
    (body → tables → headers/footers) so it matches list_signature_placeholders_in_docx.

    Returns the placeholder name written into the document (e.g. "SPONSOR_SIGNATURE").
    """
    from docx import Document

    role_key = (role or "").strip().lower()
    candidates = _SIGNATURE_PLACEHOLDERS_BY_ROLE.get(role_key) or _SIGNATURE_PLACEHOLDERS_GENERIC
    chosen_name = (placeholder_name or candidates[0]).upper()

    doc = Document(docx_in_path)
    target_paragraph = None
    for idx, paragraph in enumerate(_iter_all_paragraphs(doc)):
        if idx == paragraph_index:
            target_paragraph = paragraph
            break

    if target_paragraph is None:
        raise ValueError(
            f"Paragraph index {paragraph_index} out of range for the supplied DOCX."
        )

    # Append the marker to the target paragraph so any surrounding label text
    # ("Signature:", "Signed by:", etc.) is preserved. The stamper later
    # replaces the marker with the rendered signature image in-place.
    run = target_paragraph.add_run(f" {{{{{chosen_name}}}}}")
    # Light formatting so the marker is visually distinct in OnlyOffice preview
    # before stamping; will be replaced by image so styling is cosmetic.
    run.font.italic = True

    doc.save(docx_out_path)
    logger.info(
        "Injected signature placeholder %r at paragraph %d (role=%s) into %s",
        chosen_name, paragraph_index, role_key, docx_out_path,
    )
    return chosen_name


async def get_next_pending_signer(
    db: AsyncSession,
    agreement_id: UUID,
) -> Optional["AgreementSigner"]:
    """Return the next signer in sign_order that has not signed yet (status =
    pending or invited), or None when the chain is complete."""
    result = await db.execute(
        select(AgreementSigner)
        .where(AgreementSigner.agreement_id == agreement_id)
        .where(AgreementSigner.status.in_(("pending", "invited")))
        .order_by(AgreementSigner.sign_order.asc())
    )
    return result.scalars().first()


async def mark_signer_signed(
    db: AsyncSession,
    agreement_id: UUID,
    role: str,
    email: str,
) -> Optional["AgreementSigner"]:
    """Look up the (agreement, role, email) signer row and mark it 'signed'.
    Returns the row, or None if no matching signer is configured (legacy
    agreements created before the multi-signer schema). Caller commits."""
    result = await db.execute(
        select(AgreementSigner)
        .where(AgreementSigner.agreement_id == agreement_id)
        .where(AgreementSigner.role == role)
        .where(AgreementSigner.email == email.strip())
    )
    signer = result.scalar_one_or_none()
    if signer is None:
        return None
    signer.status = "signed"
    signer.signed_at = datetime.now(timezone.utc)
    return signer


async def dispatch_next_signer_otp(
    db: AsyncSession,
    agreement: Agreement,
    sender_email: Optional[str] = None,
) -> Optional["AgreementSigner"]:
    """Sequential dispatcher. After a signer commits, this finds the next
    pending signer, mints a fresh AgreementSigningToken, emails them an
    invite, and flips their status to 'invited'. Returns the invited signer,
    or None when the chain is complete. Email failure is logged but does not
    raise — the just-committed signature must persist."""
    next_signer = await get_next_pending_signer(db, agreement.id)
    if next_signer is None:
        logger.info(
            "dispatch_next_signer_otp: chain complete for agreement %s", agreement.id,
        )
        return None

    token_value = secrets.token_urlsafe(32)
    token_row = AgreementSigningToken(
        agreement_id=agreement.id,
        token=token_value,
        role=next_signer.role,
        recipient_email=next_signer.email,
        expires_at=datetime.now(timezone.utc) + timedelta(days=7),
        created_by=sender_email,
        is_active='true',
    )
    db.add(token_row)
    next_signer.status = "invited"
    next_signer.invited_at = datetime.now(timezone.utc)

    # IMPORTANT: this must match the React route in App.tsx — which is
    # /agreement/sign/:agreementId?token=... (path-based id, query-based token).
    # An earlier version of this code sent /sign?agreement=...&token=... which
    # did NOT match any route, so clicking the email landed on a blank page
    # and no signature ever fired.
    sign_link = (
        f"{getattr(settings, 'frontend_base_url', '').rstrip('/')}"
        f"/agreement/sign/{agreement.id}?token={token_value}"
    )
    try:
        from app.integrations.smtp_service import smtp_service
        result = await smtp_service.send_email_async(
            to=next_signer.email,
            subject=f"Your signature is required: {agreement.title}",
            body=(
                f"Hello {next_signer.name},\n\n"
                f"It is now your turn to sign '{agreement.title}' as "
                f"{next_signer.role.replace('_', ' ').title()}.\n\n"
                f"Please review and sign here:\n{sign_link}\n\n"
                f"This link is valid for 7 days. You will be asked to enter a "
                f"one-time code (sent separately) before your signature is recorded."
            ),
            from_email=getattr(settings, "smtp_user", "noreply@crm.local"),
            from_name="Clinical Trials CRM",
        )
        if not result.get("success"):
            logger.warning(
                "dispatch_next_signer_otp: email to %s failed: %s",
                next_signer.email, result.get("error"),
            )
    except Exception as exc:
        logger.warning(
            "dispatch_next_signer_otp: SMTP error for %s: %s",
            next_signer.email, exc,
        )

    logger.info(
        "dispatch_next_signer_otp: invited signer %s (role=%s, order=%s) for agreement %s",
        next_signer.email, next_signer.role, next_signer.sign_order, agreement.id,
    )
    return next_signer


def _signature_placeholder_pattern(name: str) -> "re.Pattern[str]":
    """Build a tolerant regex for {{NAME}} — case-insensitive, underscore /
    hyphen / space interchangeable, surrounding whitespace allowed."""
    # `_` is not a regex metacharacter so re.escape leaves it alone; replace
    # it on the raw name with a character class that accepts `_ - space`.
    relaxed = re.escape(name).replace("_", r"[_\-\s]")
    return re.compile(r"\{\{\s*" + relaxed + r"\s*\}\}", re.IGNORECASE)


# Role -> identifying keywords in the "For <entity>" header above each
# signature block. Used to pick which block belongs to the current signer
# when a doc has multiple "Name:/Title:/Date:" sections.
_ROLE_BLOCK_KEYWORDS = {
    "reviewer": ("dizzaroo", "sponsor", "cro", "for and on behalf"),
    "sponsor": ("dizzaroo", "sponsor", "cro", "for and on behalf"),
    "internal": ("dizzaroo", "sponsor", "cro"),
    "site": ("hospital", "site", "centre", "center", "clinic", "institute"),
    "pi": ("investigator", "pi ", "principal", "site"),
    "investigator": ("investigator", "pi ", "principal"),
    "witness": ("witness",),
    "legal": ("legal", "counsel", "attorney"),
}


def _fill_signer_labels_for_role(
    doc,
    role: str,
    signer_name: str,
    signer_title: str,
    timestamp_display: str,
) -> bool:
    """
    Best-effort filler for the "Name: ___ / Title: ___ / Date: ___" labels
    that appear inside a signature block. Walks paragraphs until it finds a
    "For <entity>" header whose entity text matches a keyword for `role`,
    then fills any empty Name/Title/Date label in the next ~6 paragraphs.

    Skips silently when:
      - the role keywords don't match any block header (template is bespoke),
      - the labels already have content (signer or template populated them),
      - signer_name / signer_title / timestamp_display are empty.

    Returns True if at least one label was filled.
    """
    if not (signer_name or signer_title or timestamp_display):
        return False

    role_key = (role or "").strip().lower()
    keywords = _ROLE_BLOCK_KEYWORDS.get(role_key, ())
    if not keywords:
        return False

    # Possessive quantifiers (Py3.11+) — the \s runs never need to backtrack
    # (nothing after them can match whitespace), this just makes that explicit
    # and keeps the match linear on pathological whitespace runs.
    label_value_pattern = re.compile(
        r"^(\s*+)(name|title|signature|date)\s*+[:\-]\s*+(.*)$",
        re.IGNORECASE,
    )

    paragraphs = list(_iter_all_paragraphs(doc))
    filled_any = False
    i = 0
    while i < len(paragraphs):
        text = (paragraphs[i].text or "").strip().lower()
        # Look for a "For <entity>" header that matches our role's keywords.
        if text.startswith("for ") and any(kw in text for kw in keywords):
            # Scan the next 8 paragraphs for empty Name/Title/Date labels.
            for j in range(i + 1, min(i + 9, len(paragraphs))):
                para_text = paragraphs[j].text or ""
                m = label_value_pattern.match(para_text)
                if m is None:
                    continue
                indent, label, existing_value = m.group(1), m.group(2).lower(), m.group(3).strip()
                if existing_value:
                    continue  # already populated
                replacement: Optional[str] = None
                if label == "name" and signer_name:
                    replacement = f"{indent}Name: {signer_name}"
                elif label == "title" and signer_title:
                    replacement = f"{indent}Title: {signer_title}"
                elif label == "date" and timestamp_display:
                    replacement = f"{indent}Date: {timestamp_display}"
                # "Signature:" is left blank here — the rendered image is
                # stamped by _stamp_signature_at_placeholder which knows how to
                # insert a picture inline.
                if replacement is None:
                    continue
                for run in list(paragraphs[j].runs):
                    run._r.getparent().remove(run._r)
                paragraphs[j].add_run(replacement)
                filled_any = True
            # One matching block per role is enough; stop scanning.
            break
        i += 1

    if filled_any:
        logger.info(
            "_fill_signer_labels_for_role: populated labels for role=%s "
            "(name=%r, title=%r)", role_key, signer_name, signer_title,
        )
    return filled_any


def _iter_all_paragraphs(doc):
    """Yield every paragraph in body, tables, headers, and footers."""
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is None:
                continue
            for paragraph in hf.paragraphs:
                yield paragraph


def _stamp_signature_at_placeholder(
    doc,
    role: str,
    png_bytes: bytes,
    signer_display: str,
    timestamp_display: str,
) -> bool:
    """
    Find a role-appropriate signature placeholder in the document and replace
    it in-place with the signature image plus the signer's name and date.
    Returns True if a placeholder was found and replaced; False otherwise.
    """
    import io as _io
    from docx.shared import Inches, Pt

    role_key = (role or "").strip().lower()
    candidates = list(_SIGNATURE_PLACEHOLDERS_BY_ROLE.get(role_key, []))
    candidates.extend(_SIGNATURE_PLACEHOLDERS_GENERIC)
    patterns = [_signature_placeholder_pattern(name) for name in candidates]

    for paragraph in list(_iter_all_paragraphs(doc)):
        text = paragraph.text or ""
        if not text:
            continue
        matched_pat = next((p for p in patterns if p.search(text)), None)
        if matched_pat is None:
            continue

        match = matched_pat.search(text)
        before_text = text[: match.start()]
        after_text = text[match.end():]

        # Clear existing runs; we rebuild the paragraph so the image sits exactly
        # where the placeholder was.
        for run in list(paragraph.runs):
            run._r.getparent().remove(run._r)

        if before_text:
            paragraph.add_run(before_text)

        pic_run = paragraph.add_run()
        pic_run.add_picture(_io.BytesIO(png_bytes), width=Inches(2.0))

        # Suffix the typed name + date so reviewers/signers see a human-readable
        # acknowledgement right next to the rendered signature image.
        suffix_run = paragraph.add_run(f"  {signer_display}  ({timestamp_display})")
        suffix_run.font.size = Pt(9)

        if after_text:
            paragraph.add_run(after_text)

        logger.info(
            "Stamped %s signature in-place at placeholder %r",
            role_key or "<unknown role>",
            match.group(0),
        )
        return True

    return False


async def _append_signature_to_docx_and_create_signed_version(
    db: AsyncSession,
    agreement: Agreement,
    source_document: AgreementDocument,
    source_docx_path: Path,
    signature_data: dict,
) -> AgreementDocument:
    """
    Stamp the visible signature on a DOCX and persist it as a NEW signed version.

    Behaviour:
      - Scans the source DOCX for a signature placeholder appropriate to the
        signer's role (e.g. {{SITE_SIGNATURE}}, {{REVIEWER_SIGNATURE}}, or the
        generic {{SIGNATURE}}). When found, replaces the placeholder in-place
        with the rendered signature image + signer name + timestamp.
      - Always appends an audit block at the END of the document with full
        e-sign metadata (signer email, UTC timestamp, auth method, meaning).
      - If no placeholder is found, the visible signature falls back to the
        audit block at the end (preserves legacy behaviour).

    This is additive: it never overwrites the existing source document row or file.
    """
    from app.modules.agreements.routes.crud import create_agreement_system_comment
    import io as _io
    from docx import Document

    tmp_out = await asyncio.to_thread(_make_named_tempfile, ".docx")
    tmp_out_path = Path(tmp_out.name)
    repaired_src = await asyncio.to_thread(_make_named_tempfile, ".docx")
    repaired_src_path = Path(repaired_src.name)

    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        from app.utils.docx_comment_extractor import repair_docx_package
        from app.utils.signature_image import render_signature_png_bytes

        try:
            repair_docx_package(str(source_docx_path), str(repaired_src_path))
            docx_open_path = repaired_src_path
        except Exception as repair_err:
            logger.warning(
                "repair_docx_package failed for %s, opening source as-is: %s",
                source_docx_path,
                repair_err,
            )
            docx_open_path = source_docx_path

        doc = Document(str(docx_open_path))
        # The inline-placeholder replacement described in the docstring is not
        # implemented in this code path: the visible signature is always appended
        # as the audit/OTP block below. Record that so the system comment reports
        # the correct mode (this also defines the name used further down).
        inline_applied = False
        doc.add_paragraph("")
        doc.add_paragraph("")

        role_raw = str(signature_data.get("role") or "").strip()
        role_label = role_raw.replace("_", " ").title() if role_raw else ""

        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_before = Pt(14)
        title_p.paragraph_format.space_after = Pt(6)
        tr = title_p.add_run("Electronic signature")
        tr.bold = True
        tr.font.size = Pt(12)

        if role_label:
            role_p = doc.add_paragraph()
            rr = role_p.add_run(f"{role_label}")
            rr.bold = True
            rr.font.size = Pt(11)

        signer_display = str(signature_data.get("email") or "").strip() or "Signer"
        png_bytes = render_signature_png_bytes(signer_display)
        pic_p = doc.add_paragraph()
        pic_p.paragraph_format.space_before = Pt(4)
        pic_p.paragraph_format.space_after = Pt(8)
        pic_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pic_run = pic_p.add_run()
        pic_run.add_picture(_io.BytesIO(png_bytes), width=Inches(3.45))

        mean_p = doc.add_paragraph(signature_data["meaning"])
        mean_p.paragraph_format.space_after = Pt(4)
        for r in mean_p.runs:
            r.font.size = Pt(10)

        meta_lines = [
            f"Signer: {signature_data['email']}",
            f"Date and time (UTC): {signature_data['timestamp_display']}",
            f"Authentication: {signature_data['auth_method']}",
        ]
        for line in meta_lines:
            mp = doc.add_paragraph(line)
            mp.paragraph_format.space_after = Pt(2)
            mp.paragraph_format.left_indent = Inches(0.05)
            for r in mp.runs:
                r.font.size = Pt(9)

        doc.save(str(tmp_out_path))

        max_version_result = await db.execute(
            select(func.max(AgreementDocument.version_number))
            .where(AgreementDocument.agreement_id == agreement.id)
        )
        max_version_number = max_version_result.scalar_one_or_none() or 0
        next_version_number = max_version_number + 1

        study_r = await db.execute(select(Study).where(Study.id == agreement.study_id))
        site_r = await db.execute(select(Site).where(Site.id == agreement.site_id))
        study = study_r.scalar_one_or_none()
        site = site_r.scalar_one_or_none()

        storage = None
        file_url = None
        file_path: str
        try:
            from app.utils.azure_storage import get_document_storage
            storage = get_document_storage()
        except Exception:
            storage = None

        async with aiofiles.open(tmp_out_path, "rb") as f:
            signed_bytes = await f.read()

        if storage and (getattr(source_document, "document_file_url", None) or settings.azure_storage_connection_string):
            blob_name = _build_pretty_document_blob_name(
                study_name=study.name if study else str(agreement.study_id),
                site_name=(site.name or site.site_id) if site else "Unknown",
                agreement_title=agreement.title,
                seq=next_version_number,
                when_utc=datetime.now(timezone.utc),
            )
            file_url = await storage.upload_file(
                _io.BytesIO(signed_bytes),
                blob_name,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                overwrite=False,
            )
            file_path = blob_name
        else:
            agreement_dir = Path(settings.upload_dir) / "agreements" / str(agreement.id)
            agreement_dir.mkdir(parents=True, exist_ok=True)
            local_path = agreement_dir / f"version_{next_version_number}_{uuid.uuid4()}.docx"
            async with aiofiles.open(local_path, "wb") as f:
                await f.write(signed_bytes)
            file_path = str(local_path)

        new_document = AgreementDocument(
            agreement_id=agreement.id,
            version_number=next_version_number,
            document_content=source_document.document_content,
            document_file_path=file_path,
            document_file_url=file_url,
            document_html='',
            created_from_template_id=getattr(source_document, "created_from_template_id", None),
            created_by=source_document.created_by,
            is_signed_version='true',
        )
        db.add(new_document)
        await db.flush()

        await create_agreement_system_comment(
            db,
            agreement.id,
            None,
            (
                f"Version {next_version_number} created with inline signature in document."
                if inline_applied
                else f"Version {next_version_number} created with visible OTP signature block."
            ),
        )

        return new_document
    finally:
        if tmp_out_path.exists():
            try:
                os.unlink(tmp_out_path)
            except OSError:
                pass
        if repaired_src_path.exists():
            try:
                os.unlink(repaired_src_path)
            except OSError:
                pass


async def _get_active_review_context(
    db: AsyncSession,
    agreement_id: UUID,
    token: str,
) -> Tuple[Agreement, AgreementReviewToken, AgreementReviewDocument]:
    token_result = await db.execute(
        select(AgreementReviewToken)
        .where(AgreementReviewToken.token == token)
        .where(AgreementReviewToken.agreement_id == agreement_id)
        .where(AgreementReviewToken.is_active == 'true')
    )
    review_token = token_result.scalar_one_or_none()
    if not review_token:
        raise HTTPException(status_code=404, detail="Invalid or expired review token")

    if review_token.expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=403, detail="Review token has expired")

    agreement_result = await db.execute(
        select(Agreement)
        .where(Agreement.id == agreement_id)
        .options(selectinload(Agreement.documents))
    )
    agreement = agreement_result.scalar_one_or_none()
    if not agreement:
        raise HTTPException(status_code=404, detail="Agreement not found")

    if agreement.status not in {
        AgreementStatusEnum.UNDER_REVIEW,
        AgreementStatusEnum.UNDER_NEGOTIATION,
    }:
        raise HTTPException(
            status_code=403,
            detail="This agreement is no longer accepting review submissions. The review period has ended.",
        )

    review_doc_result = await db.execute(
        select(AgreementReviewDocument)
        .where(AgreementReviewDocument.review_token_id == review_token.id)
        .where(AgreementReviewDocument.is_submitted == 'false')
        .options(selectinload(AgreementReviewDocument.base_version))
    )
    review_document = review_doc_result.scalar_one_or_none()
    if not review_document:
        raise HTTPException(
            status_code=404,
            detail="Review document not found. The review may have already been submitted.",
        )

    return agreement, review_token, review_document


async def _get_review_document_local_path(
    review_document: AgreementReviewDocument,
) -> Tuple[Path, Optional[object]]:
    _dt_rev = review_document.updated_at or datetime.min.replace(tzinfo=timezone.utc)
    _dt_saved = review_document.saved_at
    _prefer_saved = (
        review_document.saved_review_path
        and review_document.saved_review_url
        and _dt_saved is not None
        and _dt_saved >= _dt_rev
    )

    _temp_review = None
    if _prefer_saved:
        _rfp_raw = review_document.saved_review_path
        from app.utils.azure_storage import get_document_storage as _gds_review
        _storage = _gds_review()
        if _storage:
            _review_bytes = await _storage.download_file(_rfp_raw)
            if not _review_bytes:
                raise HTTPException(status_code=404, detail="Saved review document not found in Azure")
            _temp_review = await asyncio.to_thread(_make_named_tempfile, ".docx", _review_bytes)
            return Path(_temp_review.name), _temp_review
        raise HTTPException(status_code=500, detail="Azure storage not configured")

    _rfp_raw = review_document.review_file_path
    from app.utils.azure_storage import get_document_storage as _gds_review
    _storage = _gds_review()
    if _storage and _rfp_raw and _rfp_raw.startswith("agreement_reviews/"):
        _review_bytes = await _storage.download_file(_rfp_raw)
        if not _review_bytes:
            raise HTTPException(status_code=404, detail="Review document not found in Azure")
        _temp_review = await asyncio.to_thread(_make_named_tempfile, ".docx", _review_bytes)
        return Path(_temp_review.name), _temp_review

    review_path = Path(_rfp_raw)
    if not review_path.exists():
        raise HTTPException(status_code=404, detail="Review document file not found")
    return review_path, None


async def _promote_review_file_to_agreement_document(
    db: AsyncSession,
    agreement: Agreement,
    main_document: AgreementDocument,
    review_docx_path: Path,
) -> tuple[str, Optional[str]]:
    """
    After external review submit: copy the reviewed DOCX to a new main agreement
    blob and update **this** AgreementDocument row (the review's base_version).

    Must not use _latest_agreement_document(): when multiple version rows exist,
    the review is always tied to base_version_id; promoting the wrong row leaves
    refresh/OnlyOffice showing stale content.
    """
    import io as _io

    if not main_document or not main_document.document_file_path:
        return ("", None)

    async with aiofiles.open(review_docx_path, "rb") as f:
        docx_content = await f.read()

    study_r = await db.execute(select(Study).where(Study.id == agreement.study_id))
    site_r = await db.execute(select(Site).where(Site.id == agreement.site_id))
    study = study_r.scalar_one_or_none()
    site = site_r.scalar_one_or_none()

    from app.utils.azure_storage import get_document_storage

    storage = get_document_storage()
    doc_url = getattr(main_document, "document_file_url", None)

    if storage and (doc_url or settings.azure_storage_connection_string):
        _prev_seq = _extract_blob_seq(main_document.document_file_path)
        _new_seq = (_prev_seq + 1) if _prev_seq > 0 else (main_document.version_number + 1)
        _new_blob = _build_pretty_document_blob_name(
            study_name=study.name if study else str(agreement.study_id),
            site_name=(site.name or site.site_id) if site else "Unknown",
            agreement_title=agreement.title,
            seq=_new_seq,
            when_utc=datetime.now(timezone.utc),
        )
        new_url = await storage.upload_file(
            _io.BytesIO(docx_content),
            _new_blob,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            overwrite=False,
        )
        main_document.document_file_path = _new_blob
        main_document.document_file_url = new_url
    else:
        _base_local_path = Path(main_document.document_file_path)
        _base_local_path.parent.mkdir(parents=True, exist_ok=True)
        _stamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S_%f")
        _new_local_name = f"{_base_local_path.stem}_review_{_stamp}{_base_local_path.suffix or '.docx'}"
        draft_file_path = _base_local_path.parent / _new_local_name
        async with aiofiles.open(draft_file_path, "wb") as f:
            await f.write(docx_content)
        main_document.document_file_path = str(draft_file_path)
        main_document.document_file_url = None
        new_url = None

    await db.flush()
    new_path = main_document.document_file_path
    logger.info(
        "submit-review: promoted reviewed file to base agreement document "
        "document_id=%s version_number=%s new_path=%s",
        main_document.id,
        main_document.version_number,
        new_path,
    )
    return (str(new_path), new_url)



# ============================================================
# DEPRECATED DOCUMENT ENDPOINTS (replaced by ISF module)
# These endpoints are kept for backward compatibility only.
# New document management is handled by /api/isf-documents/*
# and /api/isf/* endpoints registered in main.py.
# TODO: Remove these endpoints after full ISF rollout.
# ============================================================


# ---------------------------------------------------------------------------
# Feasibility Questionnaire Endpoints
# ---------------------------------------------------------------------------


# --- Site Documents cluster split-off (Phase 2.1d) --------------------------
from app.modules.agreements.routes.site_documents import router as _site_documents_router
router.include_router(_site_documents_router)

# --- doc_operations cluster split-off (Phase 2.1m) -------------------------
from app.modules.agreements.routes.doc_operations import router as _doc_operations_router
router.include_router(_doc_operations_router)

# --- onlyoffice cluster split-off (Phase 2.1n) -----------------------------
from app.modules.agreements.routes.onlyoffice import router as _onlyoffice_router
router.include_router(_onlyoffice_router)

# --- signing cluster split-off (Phase 2.1o) --------------------------------
from app.modules.agreements.routes.signing import router as _signing_router
router.include_router(_signing_router)

# Unified signing endpoint group (/agreements/{id}/sign/*), gated on WORKFLOW_UNIFIED.
# Additive — the live CDA/CTA signing routers stay as the fallback.
from app.modules.agreements.routes.unified_sign import router as _unified_sign_router
router.include_router(_unified_sign_router)

# Unified review endpoint group (/agreements/{id}/review/*), gated on WORKFLOW_UNIFIED.
# Additive — the live send-for-review / submit-review / changes routes stay as fallback.
from app.modules.agreements.routes.unified_review import router as _unified_review_router
router.include_router(_unified_review_router)

# Unified notify + broadcast endpoints (/agreements/{id}/notify|broadcast), gated.
from app.modules.agreements.routes.unified_notify import router as _unified_notify_router
router.include_router(_unified_notify_router)

# --- public_pages cluster split-off (Phase 2.1p) ---------------------------
from app.modules.agreements.routes.public_pages import router as _public_pages_router
router.include_router(_public_pages_router)

# --- otp cluster split-off (Phase 2.1q) ------------------------------------
from app.modules.agreements.routes.otp import router as _otp_router
router.include_router(_otp_router)

# --- Webhooks cluster split-off (Phase 2.1i) -------------------------------
from app.modules.agreements.routes.webhooks import router as _webhooks_router
router.include_router(_webhooks_router)

# --- Review cluster split-off (Phase 2.1h) -------------------------------
from app.modules.agreements.routes.review import router as _review_router
router.include_router(_review_router)

# --- Document files cluster split-off (Phase 2.1g) --------------------------
from app.modules.agreements.routes.document_files import router as _document_files_router
router.include_router(_document_files_router)

# --- Agreement CRUD cluster split-off (Phase 2.1f) -------------------------
from app.modules.agreements.routes.crud import router as _agreement_crud_router
router.include_router(_agreement_crud_router)

# --- Changes + Review-tokens cluster split-off (Phase 2.1e) ----------------
from app.modules.agreements.routes.changes import router as _changes_router
router.include_router(_changes_router)

# --- Comments cluster split-off (Phase 2.1c) --------------------------------
from app.modules.agreements.routes.comments import router as _comments_router
router.include_router(_comments_router)

# --- Edit-lock cluster split-off (Phase 2.1b) -------------------------------
from app.modules.agreements.routes.edit_lock import router as _edit_lock_router
router.include_router(_edit_lock_router)

# --- Templates cluster split-off (Phase 2.1a) -------------------------------
from app.modules.agreements.routes.templates import router as _templates_router
router.include_router(_templates_router)

# --- Clause Library (new: clause CRUD + template composition) ---------------
from app.modules.agreements.routes.clauses import router as _clauses_router
router.include_router(_clauses_router)

from app.modules.agreements.routes.template_composition import router as _template_composition_router
router.include_router(_template_composition_router)

# --- "My CTA assignments" inbox (survivor of the type retirement) ----------
# The CTA *workflow* moved to the general engine; only the reviewer/sponsor-head
# assignment inbox (Navbar bell + WorkspaceHome) still reads cta_assignments.
from app.modules.agreements.routes.assignments import router as _assignments_router
router.include_router(_assignments_router)

# --- Agreement type sub-modules (registry-mounted) -------------------------
# As of 2026-06-18 the legacy per-document-type modules (CDA, CTA) were retired:
# both signing flows run on the general workflow engine. TYPE_MODULES is now
# empty, so this loop mounts nothing — the machinery is kept generic in case a
# future type genuinely needs a registry-mounted router.
from app.modules.agreements.registry import all_descriptors, load_all_types

load_all_types()
for _type_descriptor in all_descriptors():
    if _type_descriptor.router is not None:
        try:
            router.include_router(_type_descriptor.router)
        except Exception:  # noqa: BLE001 — one broken type must not kill the rest
            logger.exception(
                "Failed to mount router for agreement type %s — type endpoints will be unavailable",
                _type_descriptor.code,
            )

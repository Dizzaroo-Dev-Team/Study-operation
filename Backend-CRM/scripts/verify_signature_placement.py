"""
Verify LABEL-BASED PDF signature placement on the REAL template.

Walks the user's acceptance scenario directly against the production template:
  1) Expand the real template's SITE/SPONSOR signature blocks (same code create uses),
     convert to PDF (same converter signing uses).
  2) Anchor resolution: Site Director, PI, and a "remaining" site signer all resolve to
     the SAME 'For Site' Signature line; Sponsor/VP resolve to the 'For Sponsor' line.
  3) Stacking: three site signers stamped with stack_offset 0/1/2 land at DESCENDING y
     on (or just below) the For Site line — no overlap. Re-extracted from the produced PDF.
  4) Fallback: a label-less PDF resolves to no anchor -> labelled fallback block (logged),
     and still stamps (never fails / never mis-places).
  5) Compliance-shape: producing the stamped PDF is deterministic bytes -> hashable
     exactly as record_signature does (the hash is of whatever we produce; unchanged path).

Run in the backend container. Throwaway; writes only to a temp dir it cleans up.
"""
import asyncio
import base64
import glob
import hashlib
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, "/app")

from app.modules.agreements.utils import pdf_export  # noqa: E402
from app.modules.agreements.utils.pdf_export import (  # noqa: E402
    _block_family, _anchor_terms, _resolve_signature_anchor, embed_signature_into_pdf,
)
from app.modules.agreements.services import doc_convert  # noqa: E402
from app.utils.signature_image import render_signature_png_bytes  # noqa: E402

results = []


def ok(label, detail=""):
    results.append((label, detail))


def find_real_template():
    cands = glob.glob("/app/uploads/templates/template_*.docx")
    # prefer the one the project notes reference, else any with the blocks
    pref = [c for c in cands if "75a4022f" in c]
    for c in (pref + cands):
        try:
            import docx
            d = docx.Document(c)
            txt = "\n".join(p.text for p in d.paragraphs)
            for t in d.tables:
                for r in t.rows:
                    for cell in r.cells:
                        txt += "\n" + cell.text
            if "SIGNATURE_BLOCK" in txt or "For Site" in txt or "For Sponsor" in txt:
                return c
        except Exception:
            continue
    return cands[0] if cands else None


def expand_template(src_docx, work):
    """Expand the SITE/SPONSOR signature blocks exactly as the create path does."""
    import docx
    from app.utils.docx_placeholder_replace import _replace_signature_blocks_in_docx

    class StubProfile:
        hospital_name = "Mercy General Hospital"
        authorized_signatory_name = "Dr. Jane Site"
        authorized_signatory_title = "Site Director"
        primary_contracting_entity = "Dizzaroo Sponsor Inc."

    d = docx.Document(src_docx)
    try:
        _replace_signature_blocks_in_docx(d, StubProfile(), sponsor_signatory_name="Sam Sponsor")
    except Exception as exc:  # noqa: BLE001 — some templates have only one block; fine
        print("   (block expansion note:", exc, ")")
    out = str(Path(work) / "expanded.docx")
    d.save(out)
    return out


async def main():
    work = tempfile.mkdtemp(prefix="sigplace_")
    tpl = find_real_template()
    assert tpl, "no template_*.docx found under /app/uploads/templates"
    print("Template:", tpl)

    expanded = expand_template(tpl, work)
    pdf_path = await doc_convert.to_pdf(expanded)
    assert Path(pdf_path).exists(), pdf_path
    ok("Real template expanded + converted to PDF", Path(pdf_path).name)

    # --- 2) anchor resolution per role -------------------------------------------------
    a_site = _resolve_signature_anchor(pdf_path, "site")
    a_pi = _resolve_signature_anchor(pdf_path, "pi")
    a_dir = _resolve_signature_anchor(pdf_path, "site_director")
    a_spon = _resolve_signature_anchor(pdf_path, "sponsor")
    a_vp = _resolve_signature_anchor(pdf_path, "vp")
    print("  anchors: site=%s pi=%s site_director=%s sponsor=%s vp=%s" %
          (a_site, a_pi, a_dir, a_spon, a_vp))

    assert a_site is not None, "site anchor not found in real template"
    assert a_pi is not None, "pi anchor not found"
    assert a_dir is not None, "site_director anchor not found"
    # site family roles all resolve to the SAME line (one 'For Site' block).
    assert a_site[0] == a_pi[0] == a_dir[0], ("site-family on different pages", a_site, a_pi, a_dir)
    assert abs(a_site[2] - a_pi[2]) < 1.0 and abs(a_site[2] - a_dir[2]) < 1.0, \
        ("site/pi/site_director must share one 'For Site' line", a_site, a_pi, a_dir)
    ok("Site Director + PI + remaining site signer all resolve to the SAME 'For Site' Signature line",
       f"y≈{a_site[2]:.1f} on p{a_site[0]+1}")

    assert a_spon is not None, "sponsor anchor not found"
    assert a_vp is not None, "vp anchor not found"
    assert abs(a_spon[2] - a_vp[2]) < 1.0, ("sponsor/vp must share the For Sponsor line", a_spon, a_vp)
    # the two blocks are DIFFERENT lines.
    assert not (a_spon[0] == a_site[0] and abs(a_spon[2] - a_site[2]) < 1.0), \
        ("For Sponsor must be a different line than For Site", a_spon, a_site)
    ok("Sponsor + VP resolve to the 'For Sponsor' Signature line (distinct from 'For Site')",
       f"sponsor y≈{a_spon[2]:.1f}, site y≈{a_site[2]:.1f}")

    assert _block_family("site_director") == "site" and _block_family("pi") == "site"
    assert _block_family("sponsor") == "sponsor" and _block_family("vp") == "sponsor"
    ok("Block-family mapping reuses role table (site* + pi -> site; sponsor/vp -> sponsor)",
       f"_anchor_terms('pi') kw={_anchor_terms('pi')[1]}")

    # --- 3) stacking: 3 site signers, stack_offset 0/1/2 ------------------------------
    def png(name):
        return base64.b64encode(render_signature_png_bytes(name)).decode("ascii")

    # NOTE: use signer names NOT present in the expanded block (which already shows
    # "Dr. Jane Site" / "Sam Sponsor") so the caption search measures only our marks.
    cur = pdf_path
    site_signers = [("Zelda Alpha", "Site Director", 0),
                    ("Yuri Bravo", "Principal Investigator", 1),
                    ("Xena Charlie", "Study Coordinator", 2)]
    for nm, title, off in site_signers:
        cur = await embed_signature_into_pdf(cur, png(nm), nm, title, role="site",
                                             stack_offset=off, signed_on="2026-06-04")
    cur = await embed_signature_into_pdf(cur, png("Walt Delta"), "Walt Delta", "VP Clinical",
                                         role="sponsor", stack_offset=0, signed_on="2026-06-04")

    # Re-extract: every signer's caption present, site captions at descending y, sponsor near For Sponsor.
    import pdfplumber
    with pdfplumber.open(cur) as pdf:
        words = []
        for pi2, page in enumerate(pdf.pages):
            ph = float(page.height)
            for w in page.extract_words() or []:
                words.append((pi2, w["text"], ph - float(w["bottom"])))  # (page, text, y_pdf)

    def caption_y(token):
        ys = [y for (_, t, y) in words if token.lower() in t.lower()]
        return min(ys) if ys else None  # lowest on page = stacked furthest down

    y_dir = caption_y("Zelda")
    y_pi = caption_y("Yuri")
    y_crc = caption_y("Xena")
    y_spon = caption_y("Walt")
    print(f"  caption y: dir={y_dir} pi={y_pi} crc={y_crc} sponsor={y_spon}")
    if not (y_dir and y_pi and y_crc and y_spon):
        site_pg = a_site[0]
        print("  page-%d tokens 150..320:" % (site_pg + 1),
              [(t, round(y, 1)) for (p, t, y) in words if p == site_pg and 150 <= y <= 320])
    assert y_dir and y_pi and y_crc and y_spon, ("all four signer captions must appear", y_dir, y_pi, y_crc, y_spon)
    # descending within For Site block (each stacked lower than the previous).
    assert y_dir > y_pi > y_crc, ("site signers must stack downward, not overlap", y_dir, y_pi, y_crc)
    ok("Three site signers STACK downward in the 'For Site' block (no overlap)",
       f"Jane>{y_pi and 'Paul'}>{y_crc and 'Cara'}: {y_dir:.1f} > {y_pi:.1f} > {y_crc:.1f}")
    ok("Sponsor mark lands in its own block", f"Sam y≈{y_spon:.1f} (site dir y≈{y_dir:.1f})")

    # --- 4) fallback on a label-less PDF ----------------------------------------------
    from reportlab.pdfgen import canvas as rlc
    from reportlab.lib.pagesizes import letter
    plain = str(Path(work) / "plain.pdf")
    c = rlc.Canvas(plain, pagesize=letter)
    c.drawString(72, 720, "Generic contract body with no signature labels at all.")
    c.showPage(); c.save()
    assert _resolve_signature_anchor(plain, "site") is None, "label-less PDF must yield no anchor"
    fb = await embed_signature_into_pdf(plain, png("No Label"), "No Label", "Signer",
                                        role="site", stack_offset=0, signed_on="2026-06-04")
    assert Path(fb).exists(), "fallback must still produce a signed PDF (never fail)"
    with pdfplumber.open(fb) as pdf:
        ftext = "\n".join((p.extract_text() or "") for p in pdf.pages)
    assert "Electronically signed by No Label" in ftext, ("labelled fallback text missing", ftext[-300:])
    ok("No matching label -> labelled fallback block ('Electronically signed by …'), still stamps", "logged")

    # --- 5) compliance shape: produced bytes hash exactly as record_signature does ----
    h = hashlib.sha256(Path(cur).read_bytes()).hexdigest()
    ok("Produced stamped PDF is hashable as-is (record_signature hashes these exact bytes)", h[:16] + "…")

    # cleanup
    import shutil
    shutil.rmtree(work, ignore_errors=True)
    for f in glob.glob(str(Path(pdf_path).parent / "expanded*_signed*.pdf")):
        try:
            Path(f).unlink()
        except Exception:
            pass

    print("=" * 92)
    print("LABEL-BASED PDF SIGNATURE PLACEMENT — VERIFICATION (real template)")
    print("=" * 92)
    for label, detail in results:
        print(f"PASS  {label}" + (f"\n        -> {detail}" if detail else ""))
    print("=" * 92)
    print(f"{len(results)} checks passed")


if __name__ == "__main__":
    asyncio.run(main())

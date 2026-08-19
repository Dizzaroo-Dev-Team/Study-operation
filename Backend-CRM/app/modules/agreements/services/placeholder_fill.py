"""Shared placeholder-fill capability.

GENERAL across every agreement/document type — this is the single home for the
``{{token}}`` fill machinery that used to live inside ``types/cta/service.py`` and was
imported by shared routes (the generality leak). Per ADDING_A_TYPE.md ("If two types
need the same helper, lift it into ``services/``"), it now lives here and type modules
import FROM it (type → shared is allowed; shared → type is not).

Capabilities (moved verbatim — behaviour unchanged):
  - scan_placeholders_in_docx        — extract literal {{token}}s from a DOCX
  - ai_detect_additional_placeholders — Gemini-assisted detection of un-braced blanks
  - resolve_field_value              — resolve one 'source.field' against the 4 sources
  - DEFAULT_TOKEN_MAP                — well-known token -> source.field table
  - resolve_field_mapping_values     — resolve all known tokens to UPPER-keyed values
  - apply_placeholders_to_docx       — find/replace {{token}} -> value across a DOCX

All DB writes use the session passed in; never commit here.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Token pattern: {{some_token}} — possessive [^{}]++ (Py3.11+) keeps the scan
# linear (the old \s*(...+?)\s* flanks were ambiguous / super-linear). Callers
# .strip() the captured token, so whitespace inside the braces is still fine.
# ---------------------------------------------------------------------------
_TOKEN_RE = re.compile(r"\{\{([^{}]++)\}\}")

# ---------------------------------------------------------------------------
# Canonical placeholder → profile field resolution table (plan §token table)
# ---------------------------------------------------------------------------
_PROFILE_FIELD_MAP: Dict[str, str] = {
    "pi_name": "pi_name",
    "pi_email": "pi_email",
    "site_name": "site_name",
    "authorized_signatory": "authorized_signatory_name",
    "authorized_signatory_name": "authorized_signatory_name",
}

_ADDRESS_TOKENS = {"site_address"}  # composed from multiple fields


# ---------------------------------------------------------------------------
# 1. scan_placeholders_in_docx
# ---------------------------------------------------------------------------

async def scan_placeholders_in_docx(file_path: str) -> List[str]:
    """Extract literal {{token}} occurrences from a DOCX.

    Uses python-docx to walk paragraphs + tables (body + headers/footers).
    Returns unique tokens preserving order of first occurrence.
    """
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        logger.exception("python-docx is not installed; cannot scan placeholders")
        return []

    try:
        doc = Document(file_path)
    except Exception as exc:
        logger.exception("Failed to open DOCX %s: %s", file_path, exc)
        raise

    seen: dict[str, None] = {}

    def _scan_text(text: str) -> None:
        for match in _TOKEN_RE.finditer(text):
            token = match.group(1).strip()
            seen.setdefault(token, None)

    # Body paragraphs
    for para in doc.paragraphs:
        _scan_text(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_text(para.text)

    # Headers and footers
    for section in doc.sections:
        for container in (section.header, section.footer):
            for para in container.paragraphs:
                _scan_text(para.text)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _scan_text(para.text)

    return list(seen.keys())


# ---------------------------------------------------------------------------
# 2. ai_detect_additional_placeholders
# ---------------------------------------------------------------------------

async def ai_detect_additional_placeholders(
    file_path: str,
    known_tokens: List[str],
) -> List[Dict[str, Any]]:
    """Send DOCX text to Gemini asking it to spot likely placeholder-like
    phrases not already in known_tokens.

    Returns list of {token, suggested_field, confidence}.
    Gracefully returns [] on AI failure — never blocks the workflow.
    """
    try:
        from docx import Document  # type: ignore[import]

        doc = Document(file_path)
        lines: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            lines.append(para.text.strip())
        text_content = "\n".join(lines[:300])  # cap to avoid huge prompts
    except Exception as exc:
        logger.warning("ai_detect_additional_placeholders: could not extract text: %s", exc)
        return []

    if not text_content.strip():
        return []

    try:
        from app.integrations.ai.client import AIClient

        client = AIClient()
        if not client.is_available():
            logger.info("AI service not available; skipping AI placeholder detection")
            return []

        known_str = ", ".join(known_tokens) if known_tokens else "none"
        prompt = f"""You are analysing a Clinical Trial Agreement (CTA) document.

Already-detected curly-brace placeholders: {known_str}

Document text (first 300 lines):
{text_content}

Task: identify additional phrases in the document that look like they should
be filled-in placeholders but were NOT wrapped in {{{{ }}}} by the author.
Look for:
  - Bracketed phrases like [INSERT PI NAME] or [PROTOCOL NUMBER]
  - ALL CAPS labels like STUDY TITLE, IRB NUMBER
  - Obvious blanks like ____________ (underscores used as fill-in lines)
  - Phrases "NAME OF SPONSOR" or "SITE ADDRESS" appearing without a value

For `suggested_field`, use EXACTLY one of these four prefixes (singular):
  - site_profile   (e.g. site_profile.pi_name, site_profile.address_line_1)
  - agreement      (e.g. agreement.effective_date, agreement.title)
  - study          (e.g. study.protocol_number, study.sponsor_name)
  - irb            (e.g. irb.name, irb.unique_code)
Do NOT use plurals like `sites`, `agreements`, `studies`, `irbs` — those
will be rejected.

Output ONLY a single valid JSON array — no prose, no markdown, no trailing
comma. The first character of your response must be `[` and the last
character must be `]`. Each item must follow this exact shape:
[
  {{
    "token": "short_snake_case_name",
    "suggested_field": "site_profile.pi_name or study.sponsor_name etc.",
    "confidence": 0.85
  }}
]

Return `[]` (empty array) if no additional candidates are found.
"""
        result = await client.generate_json(prompt)
        if not isinstance(result, list):
            return []

        out: List[Dict[str, Any]] = []
        for item in result:
            if not isinstance(item, dict):
                continue
            token = str(item.get("token", "")).strip()
            if not token:
                continue
            # Skip already-known tokens
            if token in known_tokens:
                continue
            out.append({
                "token": token,
                "suggested_field": str(item.get("suggested_field", "")),
                "confidence": float(item.get("confidence", 0.5)),
            })
        return out

    except Exception as exc:
        logger.warning(
            "ai_detect_additional_placeholders: AI call failed (non-fatal): %s", exc
        )
        return []


# ---------------------------------------------------------------------------
# 3. resolve_field_value  (single-field resolution for field_mappings entries)
# ---------------------------------------------------------------------------

async def resolve_field_value(
    source_path: str,
    agreement: Any,
    db: AsyncSession,
) -> Optional[str]:
    """Resolve a single field_mappings value like 'site_profile.pi_name' or
    'agreement.title' to its runtime string value.

    source_path format is 'data_source.field_name'. Supported sources:
      - site_profile  -> site_profiles table for agreement.site_id
      - agreement     -> agreements row (title, study_protocol_number, etc.)
      - irb           -> irbs row via site_irb_mapping
      - study         -> studies row
    Returns None if the field cannot be resolved (caller surfaces to UI as
    an unmapped placeholder).
    """
    if not source_path or "." not in source_path:
        return None

    data_source, _, field_name = source_path.partition(".")
    data_source = data_source.strip().lower()
    field_name = field_name.strip()

    try:
        if data_source == "site_profile":
            from app.models import SiteProfile
            if not getattr(agreement, "site_id", None):
                return None
            row_result = await db.execute(
                select(SiteProfile).where(SiteProfile.site_id == agreement.site_id)
            )
            row = row_result.scalar_one_or_none()
            if row is None:
                return None
            # Special-case: compose address from multiple columns
            if field_name == "address":
                parts = [
                    getattr(row, "address_line_1", None),
                    getattr(row, "city", None),
                    getattr(row, "state", None),
                    getattr(row, "country", None),
                ]
                composed = ", ".join(p for p in parts if p)
                return composed or None
            value = getattr(row, field_name, None)
            return str(value) if value is not None else None

        elif data_source == "agreement":
            value = getattr(agreement, field_name, None)
            return str(value) if value is not None else None

        elif data_source == "study":
            from app.models import Study
            if not getattr(agreement, "study_id", None):
                return None
            row_result = await db.execute(
                select(Study).where(Study.id == agreement.study_id)
            )
            row = row_result.scalar_one_or_none()
            if row is None:
                return None
            value = getattr(row, field_name, None)
            return str(value) if value is not None else None

        elif data_source == "irb":
            from app.models.irb import IRB, SiteIRBMapping
            if not getattr(agreement, "site_id", None):
                return None
            mapping_result = await db.execute(
                select(SiteIRBMapping).where(SiteIRBMapping.site_id == agreement.site_id)
            )
            mapping = mapping_result.scalar_one_or_none()
            if mapping is None:
                return None
            irb_result = await db.execute(select(IRB).where(IRB.id == mapping.irb_id))
            irb = irb_result.scalar_one_or_none()
            if irb is None:
                return None
            value = getattr(irb, field_name, None)
            return str(value) if value is not None else None

        else:
            logger.warning("resolve_field_value: unknown data_source '%s'", data_source)
            return None

    except Exception as exc:
        logger.warning("resolve_field_value: error resolving '%s': %s", source_path, exc)
        return None


# ---------------------------------------------------------------------------
# 3c. DEFAULT_TOKEN_MAP + resolve_field_mapping_values
# Shared so EVERY creation path (CTA select-template AND the generic/unified
# replace_placeholders_in_docx flow) auto-resolves well-known placeholders the
# same way — a token that matches a known data column fills even when the
# template author never configured Field Mappings.
# ---------------------------------------------------------------------------
DEFAULT_TOKEN_MAP: Dict[str, str] = {
    # Site Profile
    "site_name": "site_profile.site_name",
    "hospital_name": "site_profile.hospital_name",
    "pi_name": "site_profile.pi_name",
    "pi_email": "site_profile.pi_email",
    "pi_phone": "site_profile.pi_phone",
    "pi_designation": "site_profile.pi_designation",
    "pi_department": "site_profile.pi_department",
    "primary_contracting_entity": "site_profile.primary_contracting_entity",
    "authorized_signatory_name": "site_profile.authorized_signatory_name",
    "authorized_signatory_email": "site_profile.authorized_signatory_email",
    "authorized_signatory_title": "site_profile.authorized_signatory_title",
    "address_line_1": "site_profile.address_line_1",
    "city": "site_profile.city",
    "state": "site_profile.state",
    "country": "site_profile.country",
    "postal_code": "site_profile.postal_code",
    "site_coordinator_name": "site_profile.site_coordinator_name",
    "site_coordinator_email": "site_profile.site_coordinator_email",
    "site_coordinator_phone": "site_profile.site_coordinator_phone",
    "site_address": "site_profile.address",  # composed in resolve_field_value
    # Study — the Study model only has: study_id (external code), name, description,
    # status. There is NO study_name/protocol_number/sponsor_name column, so map to the
    # real fields (name; study_id doubles as the protocol/external number). sponsor_name
    # has no schema home today → left unmapped (stays an editable placeholder).
    "study_name": "study.name",
    "study_protocol_number": "study.study_id",
    # IRB
    "irb_name": "irb.name",
    "irb_number": "irb.unique_code",
    "irb_country": "irb.country",
}


async def resolve_field_mapping_values(
    field_mappings: Optional[Dict[str, str]],
    agreement: Any,
    db: AsyncSession,
) -> Dict[str, str]:
    """Resolve EVERY known placeholder to its runtime value, keyed by UPPER-CASE token.

    Combines the template's own ``field_mappings`` (priority) with DEFAULT_TOKEN_MAP
    (fallback) and resolves each via ``resolve_field_value`` (all 4 sources: site_profile,
    agreement, study, irb). Returns ``{"SPONSOR_NAME": "Merck", ...}`` — only entries that
    resolve to a non-empty value. Case-insensitive: callers look up by ``token.upper()``.
    This is the single source of truth shared by the CTA select-template flow and the
    generic/unified replace_placeholders_in_docx flow."""
    lookup: Dict[str, str] = {}
    for k, v in (field_mappings or {}).items():
        if isinstance(k, str) and isinstance(v, str) and v.strip():
            lookup[k.strip().upper()] = v.strip()
    for k, v in DEFAULT_TOKEN_MAP.items():
        lookup.setdefault(k.upper(), v)

    resolved: Dict[str, str] = {}
    for key, source_path in lookup.items():
        try:
            value = await resolve_field_value(source_path, agreement, db)
        except Exception as exc:  # noqa: BLE001 — one bad mapping must not abort the rest
            logger.warning("resolve_placeholder_values: %s -> %s failed: %s", key, source_path, exc)
            continue
        if value is not None and str(value).strip():
            resolved[key] = str(value)
    return resolved


# ---------------------------------------------------------------------------
# 4. apply_placeholders_to_docx
# ---------------------------------------------------------------------------

async def apply_placeholders_to_docx(
    file_path: str,
    mapping: Dict[str, str],
) -> str:
    """Open the DOCX, find-replace each {{token}} -> value across paragraphs
    (including individual runs) and table cells. Saves as a new DOCX in the
    same folder with `_filled` suffix. Returns the new path.

    Only tokens present in *mapping* with a non-None value are replaced.
    """
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        raise RuntimeError("python-docx is not installed")

    doc = Document(file_path)

    # Build {raw_token_string: replacement} where raw_token_string is "{{token}}"
    replacements: Dict[str, str] = {}
    for token, value in mapping.items():
        if value is not None:
            replacements[f"{{{{{token}}}}}"] = str(value)
            # Also handle whitespace variants: {{ token }}
            replacements[f"{{{{ {token} }}}}"] = str(value)

    def _replace_in_paragraph(para: Any) -> None:
        """Replace tokens in a paragraph, handling tokens split across runs."""
        # Rebuild full text to find tokens that span runs
        full_text = para.text
        if not any(k in full_text for k in replacements):
            return
        # Simple single-run replacement first
        for run in para.runs:
            for token_str, replacement in replacements.items():
                if token_str in run.text:
                    run.text = run.text.replace(token_str, replacement)

        # Handle cross-run case by re-examining paragraph's full text
        # and reconstructing via XML if needed.  For most DOCX files the
        # run-level pass above handles it, but we do a full-text pass too.
        # We use a simple whole-paragraph text replacement as fallback.
        full_text_after = para.text
        for token_str, replacement in replacements.items():
            if token_str in full_text_after:
                # Attempt XML-level replacement via run text modification
                _xml_replace_in_paragraph(para, token_str, replacement)

    def _xml_replace_in_paragraph(para: Any, token_str: str, replacement: str) -> None:
        """Cross-run replacement: reassemble runs until token is found."""
        runs = para.runs
        if not runs:
            return
        # Build cumulative text positions
        texts = [r.text for r in runs]
        combined = "".join(texts)
        idx = combined.find(token_str)
        if idx == -1:
            return
        end = idx + len(token_str)
        # Find which runs cover [idx, end)
        pos = 0
        run_ranges = []
        for i, t in enumerate(texts):
            run_ranges.append((pos, pos + len(t), i))
            pos += len(t)

        affected = [(s, e, i) for s, e, i in run_ranges if e > idx and s < end]
        if not affected:
            return
        if len(affected) == 1:
            s, e, i = affected[0]
            local_start = idx - s
            local_end = end - s
            runs[i].text = texts[i][:local_start] + replacement + texts[i][local_end:]
            return
        # Multi-run: put replacement in first run, clear others
        s0, e0, i0 = affected[0]
        local_start = idx - s0
        runs[i0].text = texts[i0][:local_start] + replacement
        sl, el, il = affected[-1]
        local_end = end - sl
        runs[il].text = texts[il][local_end:]
        for _, _, im in affected[1:-1]:
            runs[im].text = ""

    # Walk all paragraphs: body
    for para in doc.paragraphs:
        _replace_in_paragraph(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)

    # Headers and footers
    for section in doc.sections:
        for container in (section.header, section.footer):
            for para in container.paragraphs:
                _replace_in_paragraph(para)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_in_paragraph(para)

    # Save to new path
    src = Path(file_path)
    out_path = src.parent / f"{src.stem}_filled{src.suffix}"
    doc.save(str(out_path))
    return str(out_path)


# ---------------------------------------------------------------------------
# 5. Source-backed field drift (#6) — GENERAL, document-type-agnostic.
#
# A filled placeholder lives in the DOCX as a tagged content control
# (w:sdt with w:tag = the UPPER token). That lets us read each field's CURRENT
# value back and compare it to the field's LIVE source value. A difference means
# the field "drifted" — the user edited it in the document (the edited value wins)
# OR the underlying source changed since fill. Either way we WARN and never
# silently overwrite; the caller (UI) decides per field whether to keep the edit
# or pull the source value. Works for ANY source-backed token, any agreement type.
# ---------------------------------------------------------------------------

def _iter_content_controls(doc):
    """Yield (tag_upper, value) for every tagged content control in a DOCX —
    body + table cells + section headers/footers. Value is the concatenated text of
    the control's content runs. General; no token/type assumptions."""
    from docx.oxml.ns import qn

    parts = [doc.element.body]
    for section in doc.sections:
        parts.extend([section.header._element, section.footer._element])

    seen_ids = set()
    for part in parts:
        if part is None or id(part) in seen_ids:
            continue
        seen_ids.add(id(part))
        for sdt in part.iter(qn('w:sdt')):
            tag_el = sdt.find('.//' + qn('w:sdtPr') + '/' + qn('w:tag'))
            if tag_el is None:
                continue
            tag = (tag_el.get(qn('w:val')) or '').strip()
            if not tag:
                continue
            content = sdt.find(qn('w:sdtContent'))
            if content is None:
                continue
            value = ''.join(t.text or '' for t in content.iter(qn('w:t')))
            yield tag.upper(), value


def read_filled_field_values(file_path: str) -> Dict[str, str]:
    """Read the current value of every tagged content-control field in a DOCX, keyed by
    UPPER-CASE token. The basis for source-drift detection. Returns {} if python-docx is
    unavailable or the file can't be opened (caller treats as "no fields to compare")."""
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        logger.warning("read_filled_field_values: python-docx not installed")
        return {}
    try:
        doc = Document(file_path)
    except Exception as exc:  # noqa: BLE001
        logger.warning("read_filled_field_values: could not open %s: %s", file_path, exc)
        return {}
    out: Dict[str, str] = {}
    for tag, value in _iter_content_controls(doc):
        # First occurrence wins (a token rarely repeats; if it does, they should match).
        out.setdefault(tag, value)
    return out


async def detect_source_drift(
    file_path: str,
    field_mappings: Optional[Dict[str, str]],
    agreement: Any,
    db: AsyncSession,
) -> Dict[str, list]:
    """Compare each source-backed field's CURRENT value in the DOCX against its LIVE
    source value. GENERAL across any source-backed token / agreement type.

    Returns ``{"in_sync": [TOKEN, ...], "drifted": [{token, doc_value, source_value}, ...]}``.
    ``drifted`` means the document value differs from the source — the field was edited
    (edited value wins) or the source moved on. The caller WARNS and never silently
    overwrites; the user decides per field whether to keep the edit or pull the source.
    Only tokens that are BOTH source-backed AND present as a filled field in the document
    are considered (free-text fields like SPONSOR_NAME, which have no source, are ignored)."""
    current = read_filled_field_values(file_path)                       # {TOKEN: doc_value}
    source = await resolve_field_mapping_values(field_mappings, agreement, db)  # {TOKEN: source_value}

    in_sync: list = []
    drifted: list = []
    for token, source_value in source.items():
        doc_value = current.get(token)
        if doc_value is None:
            continue  # source-backed token not present as a filled field in this doc
        if str(doc_value).strip() == str(source_value).strip():
            in_sync.append(token)
        else:
            drifted.append({
                "token": token,
                "doc_value": doc_value,
                "source_value": source_value,
            })
    return {"in_sync": in_sync, "drifted": drifted}


def set_filled_field_values(file_path: str, updates: Dict[str, str]) -> str:
    """Set the value of specific tagged content-control fields (by UPPER token) in a DOCX
    — the "pull source value into this field" action the user picks for a drifted field.
    Leaves every other field (and all free-text edits) untouched. Saves a new
    ``_resynced.docx`` next to the source and returns its path. GENERAL; no type logic."""
    from docx import Document  # type: ignore[import]
    from docx.oxml.ns import qn

    if not updates:
        return file_path
    norm = {str(k).strip().upper(): str(v) for k, v in updates.items()}

    doc = Document(file_path)
    parts = [doc.element.body]
    for section in doc.sections:
        parts.extend([section.header._element, section.footer._element])

    seen_ids = set()
    for part in parts:
        if part is None or id(part) in seen_ids:
            continue
        seen_ids.add(id(part))
        for sdt in part.iter(qn('w:sdt')):
            tag_el = sdt.find('.//' + qn('w:sdtPr') + '/' + qn('w:tag'))
            if tag_el is None:
                continue
            tag = (tag_el.get(qn('w:val')) or '').strip().upper()
            if tag not in norm:
                continue
            content = sdt.find(qn('w:sdtContent'))
            if content is None:
                continue
            t_elems = list(content.iter(qn('w:t')))
            if not t_elems:
                continue
            # Put the whole new value in the first text node; clear the rest so the
            # control shows exactly the source value.
            t_elems[0].text = norm[tag]
            t_elems[0].set(qn('xml:space'), 'preserve')
            for extra in t_elems[1:]:
                extra.text = ''

    src = Path(file_path)
    out_path = src.parent / f"{src.stem}_resynced{src.suffix}"
    doc.save(str(out_path))
    return str(out_path)

"""PDF export utilities for the agreements module.

Functions
---------
docx_to_pdf(docx_path) -> pdf_path
    Convert a DOCX to PDF. Primary: OnlyOffice doc-server conversion API.
    Fallback: subprocess LibreOffice headless.

render_visit_matrix_pdf(study_id, site_id, db) -> pdf_path | None
    Render the resolved visit matrix for a study+site to a PDF file via
    WeasyPrint. Returns None when WeasyPrint is unavailable or the matrix
    cannot be resolved (caller should use a one-page placeholder instead).

merge_pdfs(primary_pdf, appendix_pdf) -> new_pdf_path
    Append appendix pages after primary using pypdf.PdfWriter.

embed_signature_into_pdf(pdf_path, signature_png_base64, signer_name,
                          signer_title, page_num=-1) -> new_pdf_path
    Stamp a drawn signature PNG + name/title onto a page of the PDF.
"""
from __future__ import annotations

import asyncio
import base64
import io
import logging
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, Optional
from uuid import UUID

import aiofiles

from app.utils.log_sanitize import sfmt

logger = logging.getLogger(__name__)


def _make_named_tempfile(suffix: str, data: Optional[bytes] = None):
    """Create a NamedTemporaryFile(delete=False) with ``suffix``, optionally
    write ``data`` into it, and return the (closed) tempfile object — callers
    rely on its full-path ``.name``. Sync on purpose — call via
    ``asyncio.to_thread`` from async code so the blocking file I/O stays off
    the event loop."""
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    try:
        if data is not None:
            tmp.write(data)
    finally:
        tmp.close()
    return tmp


# ---------------------------------------------------------------------------
# docx_to_pdf
# ---------------------------------------------------------------------------

async def docx_to_pdf(docx_path: str) -> str:
    """Convert a DOCX file to PDF.

    Strategy 1: Call the OnlyOffice document server's conversion endpoint
    (if settings.onlyoffice_url is reachable and the server responds).

    Strategy 2: Spawn ``libreoffice --headless --convert-to pdf`` in the
    same directory as the source file.

    Returns the absolute path to the generated PDF.
    Raises RuntimeError when neither strategy succeeds.
    """
    from app.config import settings

    src = Path(docx_path).resolve()
    if not src.exists():
        raise FileNotFoundError(f"docx_to_pdf: source file not found: {src}")

    pdf_path = src.with_suffix(".pdf")

    # ── Strategy 1: OnlyOffice conversion API ──────────────────────────────
    try:
        pdf_path = await _onlyoffice_convert(src, settings.onlyoffice_url)
        logger.info("docx_to_pdf: converted via OnlyOffice → %s", pdf_path)
        return str(pdf_path)
    except Exception as oo_exc:
        logger.warning(
            "docx_to_pdf: OnlyOffice conversion failed (%s); falling back to LibreOffice",
            oo_exc,
        )

    # ── Strategy 2: LibreOffice headless ──────────────────────────────────
    try:
        pdf_path = await _libreoffice_convert(src)
        logger.info("docx_to_pdf: converted via LibreOffice → %s", pdf_path)
        return str(pdf_path)
    except Exception as lo_exc:
        raise RuntimeError(
            f"docx_to_pdf: both conversion strategies failed. "
            f"OnlyOffice error captured in logs; LibreOffice error: {lo_exc}"
        ) from lo_exc


async def _onlyoffice_convert(src: Path, onlyoffice_url: str) -> Path:
    """Call OnlyOffice document conversion API.

    Uses the ConvertService JSON endpoint: POST /ConvertService.ashx
    (available on all OnlyOffice Document Server deployments).
    """
    import asyncio
    import httpx

    # Serve the DOCX via a file:// workaround or upload it.
    # In a typical docker-compose deployment the backend and OnlyOffice share
    # a volume, so we can give an internal path URL.
    # We use the file upload approach: post the bytes directly.
    oo_base = onlyoffice_url.rstrip("/")
    convert_url = f"{oo_base}/ConvertService.ashx"

    docx_bytes = src.read_bytes()

    # Build a minimal conversion request. OnlyOffice requires a document URL
    # rather than raw bytes, so we fall back to a temp file served locally.
    # For environments where the backend is NOT reachable from OnlyOffice,
    # this will fail and we cascade to LibreOffice.
    from app.config import settings
    internal_base = settings.backend_internal_url.rstrip("/")
    # We don't have a served URL for an arbitrary local file; raise to skip.
    raise NotImplementedError(
        "OnlyOffice conversion requires a backend-served document URL; "
        "falling back to LibreOffice."
    )


async def _libreoffice_convert(src: Path) -> Path:
    """Convert DOCX to PDF using LibreOffice headless."""
    import asyncio

    out_dir = src.parent

    # Run in thread pool to avoid blocking the event loop
    def _run() -> Path:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(out_dir),
                str(src),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"libreoffice exited {result.returncode}: {result.stderr[:500]}"
            )
        pdf = out_dir / (src.stem + ".pdf")
        if not pdf.exists():
            raise FileNotFoundError(
                f"libreoffice ran successfully but PDF not found at {pdf}"
            )
        return pdf

    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _run)


# ---------------------------------------------------------------------------
# render_visit_matrix_pdf
# ---------------------------------------------------------------------------

async def render_visit_matrix_pdf(
    study_id: UUID,
    site_id: UUID,
    db,
) -> Optional[str]:
    """Render the visit matrix for study+site as a PDF using WeasyPrint.

    Returns the absolute path to a temporary PDF file on success.
    Returns None when WeasyPrint is unavailable or data cannot be loaded.
    Caller must handle None by using a placeholder page.
    """
    try:
        import weasyprint  # type: ignore[import]
    except ImportError:
        logger.warning(
            "render_visit_matrix_pdf: WeasyPrint not installed; returning None"
        )
        return None

    # Pull resolved visit matrix data from site_budgeting services
    try:
        matrix_data = await _load_visit_matrix_data(study_id, site_id, db)
    except Exception as exc:
        logger.warning(
            "render_visit_matrix_pdf: could not load matrix data for "
            "study=%s site=%s: %s",
            study_id,
            site_id,
            exc,
        )
        return None

    html = _render_matrix_html(matrix_data)

    try:
        # WeasyPrint is sync; run in executor to avoid blocking
        tmp = await asyncio.to_thread(_make_named_tempfile, "_visit_matrix.pdf")
        tmp_path = tmp.name

        def _write_pdf() -> None:
            weasyprint.HTML(string=html).write_pdf(tmp_path)

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, _write_pdf)
        logger.info("render_visit_matrix_pdf: wrote %s", tmp_path)
        return tmp_path

    except Exception as exc:
        logger.error("render_visit_matrix_pdf: WeasyPrint failed: %s", exc, exc_info=True)
        return None


async def _load_visit_matrix_data(study_id: UUID, site_id: UUID, db) -> dict:
    """Pull the visit matrix WITH COSTS, mirroring the Site Budgeting "Visit Matrix" UI.

    Reuses the SAME canonical resolvers the frontend uses (no reinvented math):
      * budget_service.resolve_budget_line_items → each line's factored/computed
        ``unit_cost`` (+ name, category, is_excluded) — cascade + currency + factors;
      * budget_service.resolve_visit_matrix       → per-cell ``units`` (cascade-merged).
    Cost model (identical to VisitMatrix.tsx): cell cost = units × unit_cost;
    per-patient (row) = Σ visits (units × unit_cost); visit total (column) = Σ rows;
    grand total = Σ all = total per-patient budget.

    Returns rows enriched with ``unit_cost`` and ``per_patient`` (cost), plus
    ``visit_totals`` and ``grand_total`` and ``currency``. Keeps the legacy ``total``
    key (= per_patient) so the HTML/PDF renderer keeps working.
    """
    from collections import defaultdict
    from sqlalchemy import or_, select
    from app.modules.site_budgeting.db_models import BudgetTemplate, VisitSchedule
    from app.modules.site_budgeting.services import budget_service

    empty = {"visits": [], "rows": [], "template_name": "Visit Matrix",
             "currency": None, "visit_totals": {}, "grand_total": 0.0}

    # 1. Pick the entry template for (trial, site): prefer this site's own template,
    #    else the most-specific by level. The canonical resolvers cascade up the
    #    parent chain from here for inherited lines/cells, so passing the most specific
    #    EXISTING template is correct even when the data lives at trial level.
    cand_rows = await db.execute(
        select(BudgetTemplate).where(
            BudgetTemplate.trial_id == study_id,
            or_(BudgetTemplate.site_id == site_id, BudgetTemplate.site_id.is_(None)),
        )
    )
    candidates = list(cand_rows.scalars().all())
    if not candidates:
        logger.info("_load_visit_matrix_data: no BudgetTemplate for trial_id=%s (site_id=%s)",
                    study_id, site_id)
        return empty

    def _spec_rank(t) -> int:
        if site_id is not None and str(t.site_id) == str(site_id):
            return 0
        lvl = (getattr(t, "template_level", "") or "").upper()
        return {"COUNTRY": 1, "TRIAL": 2}.get(lvl, 3)

    candidates.sort(key=_spec_rank)
    template = candidates[0]

    # 2. Canonical resolved lines (with computed/factored unit cost) + matrix cells.
    resolved_lines = await budget_service.resolve_budget_line_items(
        db, template.id, site_id=site_id)
    cells = await budget_service.resolve_visit_matrix(db, template.id)

    # 3. Visit name/order (trial-scoped, guide §3.5) — ordered columns.
    vs_rows = await db.execute(
        select(VisitSchedule)
        .where(VisitSchedule.trial_id == study_id, VisitSchedule.is_active.is_(True))
        .order_by(VisitSchedule.visit_order)
    )
    visits = [{"id": str(v.id), "name": v.visit_name, "order": int(v.visit_order or 0)}
              for v in vs_rows.scalars().all()]

    # 4. units keyed by (cost_element_id, visit_schedule_id) — robust across cascade
    #    (line_item_id can be the parent's for inherited rows; cost_element_id is stable).
    units_by_ce: dict = defaultdict(dict)
    for c in cells:
        units_by_ce[str(c["cost_element_id"])][str(c["visit_schedule_id"])] = float(c["units"])

    # 5. Build rows with unit cost + per-patient cost; accumulate column + grand totals.
    rows: list = []
    visit_totals = {v["id"]: 0.0 for v in visits}
    grand_total = 0.0
    for ln in resolved_lines:
        if ln.get("is_excluded"):
            continue
        ceid = str(ln.get("cost_element_id"))
        per_visit_units = units_by_ce.get(ceid, {})
        if not any(per_visit_units.values()):
            continue  # no per-visit units → not part of the visit matrix (e.g. PER_PATIENT/FIXED)
        unit_cost = float(ln.get("unit_cost") or ln.get("unit_cost_computed") or 0.0)
        per_visit = {v["id"]: per_visit_units.get(v["id"], 0.0) for v in visits}
        per_patient = sum(u * unit_cost for u in per_visit.values())
        for vid, u in per_visit.items():
            visit_totals[vid] += u * unit_cost
        grand_total += per_patient
        rows.append({
            "name": ln.get("name") or ln.get("code") or "Unnamed",
            "unit_cost": unit_cost,
            "per_visit": per_visit,
            "per_patient": per_patient,
            "total": per_patient,  # legacy key for the HTML/PDF renderer
        })

    return {
        "template_name": getattr(template, "name", None) or "Visit Matrix",
        "currency": getattr(template, "target_currency_code", None),
        "visits": visits,
        "rows": rows,
        "visit_totals": visit_totals,
        "grand_total": grand_total,
    }


def _render_matrix_html(data: dict) -> str:
    """Render visit matrix data as a minimal HTML table."""
    visits = data.get("visits", [])
    rows = data.get("rows", [])
    title = data.get("template_name", "Visit Matrix")
    cur = data.get("currency")
    visit_totals = data.get("visit_totals", {}) or {}
    grand_total = data.get("grand_total", 0.0)

    def _money(x) -> str:
        try:
            val = float(x or 0)
        except (TypeError, ValueError):
            val = 0.0
        if val == 0:
            return "—"
        return f"{cur} {val:,.2f}" if cur else f"{val:,.2f}"

    header_cells = "<th>Unit Cost</th>"
    header_cells += "".join(f"<th>{v['name']}</th>" for v in visits)
    header_cells += "<th>Per-Patient</th>"

    body_rows = ""
    for row in rows:
        cells = f"<td>{_money(row.get('unit_cost'))}</td>"
        cells += "".join(
            f"<td>{row['per_visit'].get(v['id'], 0.0):g}</td>" for v in visits
        )
        cells += f"<td>{_money(row.get('per_patient', row.get('total', 0.0)))}</td>"
        body_rows += f"<tr><td>{row['name']}</td>{cells}</tr>\n"

    if rows:
        foot = "<td><strong>Total cost</strong></td><td></td>"
        foot += "".join(f"<td><strong>{_money(visit_totals.get(v['id'], 0.0))}</strong></td>" for v in visits)
        foot += f"<td><strong>{_money(grand_total)}</strong></td>"
        body_rows += f"<tr>{foot}</tr>\n"

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; font-size: 11px; margin: 20px; }}
  h2 {{ font-size: 14px; margin-bottom: 8px; }}
  table {{ border-collapse: collapse; width: 100%; }}
  th, td {{ border: 1px solid #999; padding: 4px 8px; text-align: right; }}
  th {{ background: #e0e0e0; text-align: center; }}
  td:first-child {{ text-align: left; }}
</style>
</head>
<body>
<h2>{title} — Budget Appendix</h2>
<table>
  <thead>
    <tr><th>Cost Element</th>{header_cells}</tr>
  </thead>
  <tbody>
    {body_rows if body_rows else "<tr><td colspan='100'>No data</td></tr>"}
  </tbody>
</table>
</body>
</html>"""


# ---------------------------------------------------------------------------
# append_visit_matrix_to_docx — bake the budget into the DOCX itself so the
# internal reviewer sees CTA text + budget inside the OnlyOffice editor (a PDF
# appendix only helps at PDF generation time, not during review).
# ---------------------------------------------------------------------------

def _norm_marker(s: str) -> str:
    """Normalize a marker / paragraph text for matching: drop ALL whitespace and
    upper-case it. Word often splits a token like ``{{ANNEXURE_A}}`` across runs and
    may pad it with spaces, so an exact compare misses; the normalized compare doesn't."""
    return "".join((s or "").split()).upper()


def _find_marker_paragraph(doc, markers):
    """Return the FIRST paragraph whose text carries one of ``markers`` (normalized,
    whitespace/case-insensitive), plus the matched marker — or (None, None)."""
    norm = [(_norm_marker(m), m) for m in (markers or []) if m]
    for p in doc.paragraphs:
        nt = _norm_marker(p.text)
        if not nt:
            continue
        for nm, original in norm:
            if nm and nm in nt:
                return p, original
    return None, None


def _remove_or_clear_marker(marker_para, matched: str) -> None:
    """The marker line has done its job (the budget now sits before it). If the marker
    was the WHOLE paragraph, remove the paragraph entirely; otherwise strip just the
    marker token from it and keep the surrounding text."""
    if _norm_marker(marker_para.text) == _norm_marker(matched):
        marker_para._p.getparent().remove(marker_para._p)
        return
    # Marker embedded in other text: rewrite the paragraph without the token.
    cleaned = (marker_para.text or "").replace(matched, "")
    for r in list(marker_para.runs):
        r._r.getparent().remove(r._r)
    if cleaned.strip():
        marker_para.add_run(cleaned)


async def append_visit_matrix_to_docx(
    docx_path: str,
    study_id: UUID,
    site_id: UUID,
    db,
    markers=None,
) -> str:
    """Insert a "Site Budget — Visit Matrix" section (heading + table) into the DOCX.

    PLACEMENT: when ``markers`` is given (e.g. ["{{ANNEXURE_A}}", "{{BUDGET}}"]) and one
    is found in the document, the budget is inserted AT that marker and the marker line is
    removed — so the budget lands exactly where the template author put the placeholder
    (e.g. Section 4 "Budget (Annexure A)"), NOT dumped after the signatures. When no marker
    is given or none is found, the budget is appended at the END on its own page (legacy
    behaviour). Returns the new path (a `_with_budget` suffix file).

    Falls back to a one-line "Site Budget: not yet created." paragraph when the matrix is
    empty / unavailable so the section is at least signposted.
    """
    try:
        from docx import Document  # type: ignore[import]
        from docx.enum.text import WD_BREAK  # type: ignore[import]
        from docx.shared import Pt  # type: ignore[import]
    except ImportError:
        logger.warning("append_visit_matrix_to_docx: python-docx unavailable; skipping")
        return docx_path

    try:
        matrix = await _load_visit_matrix_data(study_id, site_id, db)
    except Exception as exc:
        logger.warning("append_visit_matrix_to_docx: data load failed: %s", exc)
        matrix = {"visits": [], "rows": [], "template_name": "Visit Matrix"}

    visits = matrix.get("visits", []) or []
    rows = matrix.get("rows", []) or []
    title = matrix.get("template_name", "Visit Matrix")

    try:
        doc = Document(docx_path)
        marker_para, matched = _find_marker_paragraph(doc, markers)

        if not rows or not visits:
            # No budget set up yet — keep it terse, no page break, no heading.
            para = doc.add_paragraph()
            run = para.add_run("Site Budget: not yet created.")
            run.italic = True
            if marker_para is not None:
                marker_para._p.addprevious(para._p)
                _remove_or_clear_marker(marker_para, matched)
        else:
            # Page break only for the END-appended appendix; a marker insertion stays
            # inline where the author placed the placeholder.
            if marker_para is None:
                break_para = doc.add_paragraph()
                break_para.add_run().add_break(WD_BREAK.PAGE)

            # Heading as a plain bold paragraph — NOT doc.add_heading(level=1), which
            # requires the document to already define a built-in "Heading 1" style.
            # Custom CTA templates often don't, so add_heading raises KeyError
            # ("no style with name 'Heading 1'"); the except below then swallows it and
            # returns the doc with NO budget appended — i.e. the budget silently vanishes.
            heading = doc.add_paragraph()
            heading_run = heading.add_run(f"Site Budget — {title}")
            heading_run.bold = True
            heading_run.font.size = Pt(16)

            # Cost-aware columns mirroring the Site Budgeting Visit Matrix UI:
            #   Cost Element | Unit Cost | <visit unit counts...> | Per-Patient (cost)
            # plus a "Total cost" footer (per-visit cost + grand total per patient).
            cur = matrix.get("currency")
            visit_totals = matrix.get("visit_totals", {}) or {}
            grand_total = matrix.get("grand_total", 0.0)

            def _money(x) -> str:
                try:
                    val = float(x or 0)
                except (TypeError, ValueError):
                    val = 0.0
                if val == 0:
                    return "—"  # unpriced / zero — matches the Visit Matrix UI convention
                return f"{cur} {val:,.2f}" if cur else f"{val:,.2f}"

            def _units(x) -> str:
                val = float(x or 0)
                return str(int(val)) if val == int(val) else f"{val:g}"

            n_cols = 2 + len(visits) + 1  # element + unit-cost + visits + per-patient
            table = doc.add_table(rows=1 + len(rows) + 1, cols=n_cols)  # +1 totals footer
            # Apply a table style defensively: a custom CTA template may not define
            # "Light Grid Accent 1" (or even "Table Grid"), and assigning a missing
            # style raises KeyError — which the except below swallows, dropping the
            # WHOLE budget. Try the preferred style, fall back to a plain grid, and if
            # neither exists leave the table unstyled (it still renders with content).
            for _style in ("Light Grid Accent 1", "Table Grid"):
                try:
                    table.style = _style
                    break
                except KeyError:
                    continue

            header_cells = table.rows[0].cells
            header_cells[0].text = "Cost Element"
            header_cells[1].text = "Unit Cost"
            for i, v in enumerate(visits, start=2):
                header_cells[i].text = str(v.get("name", f"Visit {i - 1}"))
            header_cells[-1].text = "Per-Patient"
            for c in header_cells:  # bold the header row
                for p in c.paragraphs:
                    for run in p.runs:
                        run.bold = True

            for r_idx, r in enumerate(rows, start=1):
                cells = table.rows[r_idx].cells
                cells[0].text = str(r.get("name", ""))
                cells[1].text = _money(r.get("unit_cost"))
                per_visit = r.get("per_visit", {}) or {}
                for i, v in enumerate(visits, start=2):
                    cells[i].text = _units(per_visit.get(v.get("id"), 0.0))
                cells[-1].text = _money(r.get("per_patient", r.get("total", 0.0)))

            # Totals footer: per-visit cost + grand total (= total per-patient budget).
            foot = table.rows[-1].cells
            foot[0].text = "Total cost"
            for i, v in enumerate(visits, start=2):
                foot[i].text = _money(visit_totals.get(v.get("id"), 0.0))
            foot[-1].text = _money(grand_total)
            for c in foot:  # bold the totals row
                for p in c.paragraphs:
                    for run in p.runs:
                        run.bold = True

            if marker_para is not None:
                # Move the heading + table (currently at the end of the body) to right
                # before the marker line, then drop the marker line itself.
                marker_para._p.addprevious(heading._p)
                marker_para._p.addprevious(table._tbl)
                _remove_or_clear_marker(marker_para, matched)
                logger.info("append_visit_matrix_to_docx: inserted budget at marker %r", matched)
            else:
                logger.info("append_visit_matrix_to_docx: no marker found; appended budget at end")

        src = Path(docx_path)
        out_path = src.parent / f"{src.stem}_with_budget{src.suffix}"
        doc.save(str(out_path))
        logger.info("append_visit_matrix_to_docx: wrote %s", out_path)
        return str(out_path)
    except Exception as exc:
        # Do NOT silently return the original (budget-less) doc — that is exactly what
        # hid the missing-style bugs ("Heading 1" / "Light Grid Accent 1") and made the
        # budget "vanish": the job reported appended=true while shipping a doc with no
        # budget. RAISE so the failure is loud (the job records the error / the
        # fill-placeholders caller logs+degrades via its own try/except). The empty-data
        # case is handled above (it returns a signposted doc, not via this except).
        logger.error(
            "append_visit_matrix_to_docx: failed to append budget to %s: %s",
            docx_path, exc, exc_info=True,
        )
        raise


# ---------------------------------------------------------------------------
# merge_pdfs
# ---------------------------------------------------------------------------

def merge_pdfs(primary_pdf: str, appendix_pdf: str) -> str:
    """Append appendix pages after primary using pypdf.PdfWriter.

    Returns the path to the merged PDF (saved next to the primary file with
    a ``_merged`` suffix).
    """
    try:
        from pypdf import PdfWriter, PdfReader  # type: ignore[import]
    except ImportError:
        try:
            from PyPDF2 import PdfWriter, PdfReader  # type: ignore[import]
        except ImportError:
            raise RuntimeError(
                "merge_pdfs: neither pypdf nor PyPDF2 is installed. "
                "Add 'pypdf>=4' to requirements.txt."
            )

    writer = PdfWriter()

    for path in (primary_pdf, appendix_pdf):
        reader = PdfReader(path)
        for page in reader.pages:
            writer.add_page(page)

    src = Path(primary_pdf)
    out_path = src.parent / f"{src.stem}_merged.pdf"
    with open(out_path, "wb") as fh:
        writer.write(fh)

    logger.info("merge_pdfs: merged %s pages → %s", len(writer.pages), out_path)
    return str(out_path)


# ---------------------------------------------------------------------------
# Placeholder-based placement: insert an appendix AT an author-placed {{MARKER}}
# (the same braces convention as {{SIGNATURE_<ROLE>}}). DETERMINISTIC — the author
# controls placement by where they put the marker; nothing AI decides position.
# ---------------------------------------------------------------------------

def find_placeholder_page(pdf_path: str, token: str) -> Optional[int]:
    """Return the 0-based page index where `token` (e.g. '{{ANNEXURE_A}}') first
    appears in the PDF, or None when it's absent. Whitespace- and case-insensitive —
    reuses the SAME pdfplumber word-extraction + whitespace-stripped concat search the
    signature placeholder anchor (_resolve_placeholder_anchor) uses, so author-placed
    markers are located the same deterministic way across templates (LibreOffice can
    split a token across words; stripping whitespace re-joins it)."""
    if not (token and token.strip()):
        return None
    try:
        import pdfplumber  # type: ignore[import]
    except Exception:  # noqa: BLE001 — no pdfplumber -> caller falls back to end-append
        return None
    target = re.sub(r"\s+", "", token.upper())
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for pidx, page in enumerate(pdf.pages):
                words = page.extract_words(use_text_flow=True, keep_blank_chars=False) or []
                concat = re.sub(r"\s+", "", "".join(w["text"] for w in words).upper())
                if target in concat:
                    return pidx
    except Exception as exc:  # noqa: BLE001 — search failure -> caller falls back to end
        logger.warning("find_placeholder_page: search for %r failed: %s", token, exc)
    return None


def insert_pdf_after_page(primary_pdf: str, appendix_pdf: str, after_page_index: int) -> str:
    """Assemble primary[0..after] + appendix + primary[after+1..] into a new PDF, so the
    appendix lands right after the page carrying the marker. Reuses the SAME pypdf
    writer as merge_pdfs — page ORDERING only, no new rendering. Returns the new path
    (``_placed`` suffix next to the primary)."""
    try:
        from pypdf import PdfReader, PdfWriter  # type: ignore[import]
    except ImportError:
        from PyPDF2 import PdfReader, PdfWriter  # type: ignore[import]

    primary = PdfReader(primary_pdf)
    appendix = PdfReader(appendix_pdf)
    n = len(primary.pages)
    cut = max(-1, min(int(after_page_index), n - 1))  # clamp; -1 => appendix goes first
    writer = PdfWriter()
    for i in range(0, cut + 1):
        writer.add_page(primary.pages[i])
    for pg in appendix.pages:
        writer.add_page(pg)
    for i in range(cut + 1, n):
        writer.add_page(primary.pages[i])

    src = Path(primary_pdf)
    out_path = src.parent / f"{src.stem}_placed.pdf"
    with open(out_path, "wb") as fh:
        writer.write(fh)
    logger.info("insert_pdf_after_page: appendix after page %s → %s (%s pages)",
                cut, out_path, len(writer.pages))
    return str(out_path)


# ---------------------------------------------------------------------------
# embed_signature_into_pdf — role-aware, label-based placement
# ---------------------------------------------------------------------------

_STACK_ROW = 15.0  # vertical offset (pt) between stacked signatures in the same block


class SignaturePlacementError(RuntimeError):
    """Raised when a signer's signature cannot be placed in the document.

    FAIL LOUD: signing must never report success while leaving a
    ``{{SIGNATURE_<ROLE>}}`` placeholder empty. The caller is expected to
    translate this into a clear user-facing error AND to avoid advancing the
    workflow / marking the signer complete. General across templates, roles and
    signer counts — the message carries the role and the candidate tokens that
    were searched so the misconfiguration is obvious."""


# Role-id synonyms that collapse to the two well-known block families. Anything that
# matches NEITHER set keeps its OWN family (its primary token) — so witness, legal, a
# second sponsor signer, etc. are no longer mis-bucketed. Matched as whole word tokens
# of the role id (not substrings) to avoid false hits.
_SPONSOR_ROLE_SYNONYMS = frozenset(("sponsor", "vp", "cro", "company", "dizzaroo"))
_SITE_ROLE_SYNONYMS = frozenset((
    "site", "hospital", "centre", "center", "clinic", "institute",
    "investigator", "pi", "principal",
))


def _role_tokens(role: Optional[str]) -> list:
    """Whole-word tokens of a role id, with signing-convention noise removed.
    'site_director_signer' -> ['site', 'director']; 'vp_signer' -> ['vp']."""
    return [t for t in re.split(r"[^a-z0-9]+", (role or "").lower())
            if t and t not in ("signer", "sign")]


def _block_family(role: Optional[str]) -> str:
    """The signature-block family a signer's role belongs to — GENERIC across any number
    and type of signer (no more two-bucket collapse). Sponsor-side synonyms map to
    'sponsor', site-side synonyms to 'site' (preserving the established behavior for
    those roles), and ANY other role keeps its own primary token as a distinct family
    (e.g. 'witness' -> 'witness', 'legal' -> 'legal', 'designated' -> 'designated') so it
    is matched and stacked on its own rather than dumped into 'site'."""
    toks = _role_tokens(role)
    tset = set(toks)
    if tset & _SPONSOR_ROLE_SYNONYMS:
        return "sponsor"
    if tset & _SITE_ROLE_SYNONYMS:
        return "site"
    return toks[0] if toks else "site"


def _anchor_terms(role: Optional[str]):
    """(family, header_keywords, role_label_phrases) for matching a role's signature
    block in PDF text. REUSES the existing _ROLE_BLOCK_KEYWORDS table from the
    aggregator (deferred import avoids a circular import) — no new role mapping."""
    fam = _block_family(role)
    try:
        from app.modules.agreements.aggregator import _ROLE_BLOCK_KEYWORDS
        kw = tuple(_ROLE_BLOCK_KEYWORDS.get(fam, ()))
    except Exception:  # noqa: BLE001 — fall back to minimal keywords
        kw = ("sponsor", "cro") if fam == "sponsor" else ("site", "hospital", "investigator")
    r = (role or "").replace("_", " ").strip().lower()
    role_labels = [f"{r} signature", f"signature of {r}", f"for {r}"] if r else []
    return fam, kw, role_labels


def _family_of_header(txt: str):
    """If `txt` is a block header ('For <party>'), return its family; else None. GENERIC:
    sponsor synonyms -> 'sponsor', site synonyms -> 'site', otherwise the party word
    itself (so 'For Witness' -> 'witness', 'For Legal Counsel' -> 'legal-ish' via its
    first word) — matching the generic families produced by _block_family so any signer
    type lines up with its own block. Sponsor checked first so 'For Sponsor' never falls
    through. Matched on whole words to avoid substring false hits."""
    if not txt.startswith("for "):
        return None
    words = [w for w in re.split(r"[^a-z0-9]+", txt) if w and w != "for"]
    wset = set(words)
    if wset & _SPONSOR_ROLE_SYNONYMS:
        return "sponsor"
    if wset & _SITE_ROLE_SYNONYMS:
        return "site"
    return words[0] if words else None


def _resolve_signature_anchor(pdf_path: str, role: Optional[str]):
    """Locate the signature LINE for `role` in the PDF. Returns (page_index, x, y_pdf)
    where (x, y_pdf) is just right of the block's 'Signature:' label (y from the page
    bottom), or None when no matching block/line is found. Uses pdfplumber word
    coordinates and is general across templates:

      * Walks lines top-to-bottom tracking the CURRENT block via 'For <party>' headers
        (mapped to a family with the existing role keyword table), then returns the
        first true 'Signature:' label line (starts with 'signature' + has a colon —
        so a section heading like '10. Signatures' is NOT mistaken for the line) that
        falls inside the requested family's block.
      * Also matches an explicit role-labeled line ('<Role> Signature', 'Signature of
        <Role>') anywhere, for templates that label per signer instead of per block.
    """
    if not role:
        return None
    fam, kw, role_labels = _anchor_terms(role)
    try:
        import pdfplumber  # type: ignore[import]
    except Exception as exc:  # noqa: BLE001
        # FAIL LOUD: without pdfplumber we cannot locate ANY anchor (placeholder or
        # label line). Returning None silently degraded to a corner block and left
        # the placeholder empty. Surface the environment problem instead.
        raise SignaturePlacementError(
            "pdfplumber is unavailable — cannot locate signature placement anchors. "
            "This is an environment/dependency problem, not a missing placeholder."
        ) from exc

    def _anchor_from_word(pidx, ph, ws):
        sig_word = next((x for x in ws if "signature" in x["text"].lower()), ws[-1])
        return (pidx, float(sig_word["x1"]) + 8.0, ph - float(sig_word["bottom"]))

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for pidx, page in enumerate(pdf.pages):
                ph = float(page.height)
                words = page.extract_words(use_text_flow=True, keep_blank_chars=False) or []
                if not words:
                    continue
                # Group words into visual lines by rounded top coordinate (top ascending).
                lines: dict = {}
                for w in words:
                    lines.setdefault(round(float(w["top"]) / 3.0) * 3.0, []).append(w)
                rows = [(top, " ".join(x["text"] for x in sorted(ws, key=lambda z: z["x0"])).lower(), ws)
                        for top, ws in sorted(lines.items())]

                current_family = None
                for _top, txt, ws in rows:
                    hf = _family_of_header(txt)
                    if hf is not None:
                        current_family = hf
                        continue
                    # Explicit per-signer label ('<role> signature' / 'signature of <role>').
                    if "signature" in txt and any(rl in txt for rl in role_labels if not rl.startswith("for ")):
                        return _anchor_from_word(pidx, ph, ws)
                    # A true 'Signature:' label line inside the requested family's block.
                    is_sig_label = txt.startswith("signature") and ":" in txt
                    if is_sig_label and current_family == fam:
                        return _anchor_from_word(pidx, ph, ws)
        return None
    except Exception as exc:  # noqa: BLE001
        # Distinguish a genuine probe FAILURE (corrupt PDF / pdfplumber crash) from a
        # legitimate "anchor not found" (the `return None` above). A crash must be
        # loud, not silently turned into a corner-block fallback.
        raise SignaturePlacementError(
            f"signature anchor resolution failed for role={role!r}: {exc}"
        ) from exc


# ---------------------------------------------------------------------------
# Per-signer placeholder placement: {{SIGNATURE_<ROLE>}}
# ---------------------------------------------------------------------------

def _placeholder_tokens_for_role(role: Optional[str]):
    """Candidate per-signer placeholder tokens (UPPER, no braces) for `role`, in the
    {{SIGNATURE_<ROLE>}} convention, ordered MOST SPECIFIC FIRST.

    A workflow signing slot id often carries convention noise / qualifiers, e.g.
    'site_director_signer', 'vp_signer', 'cro_legal_signer', 'vp_signer_site'. The
    template, however, uses the bare role token ({{SIGNATURE_SITE_DIRECTOR}},
    {{SIGNATURE_VP}}, {{SIGNATURE_CRO}}). So we generate candidates by:
      1. the exact slot id (in case a template literally uses the _SIGNER token),
      2. the slot id with the 'signer'/'sign' convention segment removed, then
      3. progressively dropping the trailing qualifier segment down to the primary
         role family — so 'site_director_signer' -> SITE_DIRECTOR -> SITE,
         'cro_legal_signer' -> CRO_LEGAL -> CRO, 'vp_signer_site' -> VP_SITE -> VP.
    This makes any '<role>[_qualifier]_signer' slot land at its real placeholder
    without per-def tweaks. The resolver tries these in order and the EXACT token
    wins; the generic family token is only reached when no specific one exists. Where
    two distinct slots collapse to one shared template token (e.g. both CRO signers ->
    SIGNATURE_CRO, when the template has a single CRO block) they share that block.

    We also keep the EXISTING cross-naming aliases from _SIGNATURE_PLACEHOLDERS_BY_ROLE
    (deferred import) so SIGNATURE_PI / SIGNATURE_INVESTIGATOR remain interchangeable —
    the role map is reused, not rebuilt. Nothing else about placement changes; the
    fallback tiers (label line, labelled corner) remain the safety net."""
    r = (role or "").strip()
    if not r:
        return []
    norm = re.sub(r"[^A-Z0-9]+", "_", r.upper()).strip("_")
    segs = [s for s in norm.split("_") if s]
    if not segs:
        return []
    core = [s for s in segs if s not in ("SIGNER", "SIGN")]  # drop convention noise

    out: list[str] = []

    def add(name: str) -> None:
        t = "SIGNATURE_" + name
        if name and t not in out:
            out.append(t)

    add("_".join(segs))                         # 1) exact slot id (e.g. SITE_DIRECTOR_SIGNER)
    for n in range(len(core), 0, -1):           # 2-3) signer-stripped, specific -> family
        add("_".join(core[:n]))                 #      SITE_DIRECTOR, then SITE

    # Existing cross-naming aliases — looked up by the original role AND the primary
    # family segment, so a 'pi_signer' slot still benefits from PI<->INVESTIGATOR etc.
    try:
        from app.modules.agreements.aggregator import _SIGNATURE_PLACEHOLDERS_BY_ROLE
        alias_keys = {r.lower()}
        if core:
            alias_keys.add(core[0].lower())
        for key in alias_keys:
            for entry in _SIGNATURE_PLACEHOLDERS_BY_ROLE.get(key, []):
                u = entry.upper()
                if u.endswith("_SIGNATURE_BLOCK"):
                    continue
                for suf in ("_SIGNATURE", "_SIGN"):
                    if u.endswith(suf):
                        stem = u[: -len(suf)]
                        if stem:
                            add(stem)
                        break
    except Exception:  # noqa: BLE001
        pass
    return out


def _resolve_placeholder_anchor(pdf_path: str, role: Optional[str]):
    """Locate this signer's exact {{SIGNATURE_<ROLE>}} placeholder in the PDF. Returns
    (page_index, x0, x1, y_bottom, y_top) — the placeholder's bounding box in PDF
    coordinates (y from the page bottom) — or None when no per-signer placeholder for
    this role is present. The box lets the stamper COVER the literal token text and
    draw the mark exactly where the author put it. General across templates/roles.

    Matching is anchored on the BRACED token ('{{SIGNATURE_<ROLE>}}') and returns the
    MINIMAL consecutive-word span that covers it, computed from per-word character
    offsets. This is deliberate:
      * it never swallows a preceding label ('Signature:' has no braces, so it can't
        be part of the match), and
      * it always includes the trailing '}}' even when LibreOffice splits the token
        across words (e.g. '{{SIGNATURE_SITE' + '}}') — so the whole token, braces
        included, is covered (the earlier word-accumulation logic could stop early and
        leave the closing braces showing through)."""
    tokens = _placeholder_tokens_for_role(role)
    if not tokens:
        return None
    try:
        import pdfplumber  # type: ignore[import]
    except Exception as exc:  # noqa: BLE001
        # FAIL LOUD — see _resolve_signature_anchor. A missing dependency is an
        # environment problem, not "placeholder absent".
        raise SignaturePlacementError(
            "pdfplumber is unavailable — cannot locate signature placeholders. "
            "This is an environment/dependency problem, not a missing placeholder."
        ) from exc

    def nows(s):
        # Uppercase, drop only whitespace — keep braces/underscores so the braced
        # token stays distinguishable from a bare 'Signature' label.
        return re.sub(r"\s+", "", s.upper())

    # Candidate tokens, MOST SPECIFIC FIRST: braced/precise forms before bare forms,
    # each in the specificity order from _placeholder_tokens_for_role.
    targets = [nows("{{" + t + "}}") for t in tokens] + [nows(t) for t in tokens]
    _seen: set = set()
    targets = [t for t in targets if t and not (t in _seen or _seen.add(t))]

    try:
        with pdfplumber.open(pdf_path) as pdf:
            # Cache each line's concatenation + per-word [start,end) char ranges ONCE,
            # so candidates can be searched in specificity order across the WHOLE doc.
            cached: list = []  # (pidx, page_height, concat, ranges)
            for pidx, page in enumerate(pdf.pages):
                ph = float(page.height)
                words = page.extract_words(use_text_flow=True, keep_blank_chars=False) or []
                if not words:
                    continue
                lines: dict = {}
                for w in words:
                    lines.setdefault(round(float(w["top"]) / 3.0) * 3.0, []).append(w)
                for _top, ws in sorted(lines.items()):
                    ws = sorted(ws, key=lambda z: z["x0"])
                    concat, ranges = "", []
                    for w in ws:
                        s = nows(w["text"])
                        ranges.append((len(concat), len(concat) + len(s), w))
                        concat += s
                    cached.append((pidx, ph, concat, ranges))
            # Specificity-first: the exact per-signer token (SIGNATURE_SITE_DIRECTOR)
            # wins over the family token (SIGNATURE_SITE) even if the family token
            # happens to appear earlier in the document. Returns the minimal
            # consecutive-word span covering the matched token (unchanged).
            for tgt in targets:
                for pidx, ph, concat, ranges in cached:
                    pos = concat.find(tgt)
                    if pos < 0:
                        continue
                    end = pos + len(tgt)
                    covering = [w for (a, b, w) in ranges if a < end and b > pos]
                    if not covering:
                        continue
                    x0 = min(float(w["x0"]) for w in covering)
                    x1 = max(float(w["x1"]) for w in covering)
                    t0 = min(float(w["top"]) for w in covering)
                    b1 = max(float(w["bottom"]) for w in covering)
                    return (pidx, x0, x1, ph - b1, ph - t0)
        return None
    except Exception as exc:  # noqa: BLE001
        # Genuine probe failure (corrupt PDF / pdfplumber crash) must be loud, not
        # silently degraded to "placeholder not found".
        raise SignaturePlacementError(
            f"signature placeholder resolution failed for role={role!r}: {exc}"
        ) from exc


async def embed_signature_into_pdf(
    pdf_path: str,
    signature_png_base64: str,
    signer_name: str,
    signer_title: str,
    page_num: int = -1,
    *,
    role: Optional[str] = None,
    stack_offset: int = 0,
    signed_on: Optional[str] = None,
) -> str:
    """Stamp the signature PNG + name/date onto the PDF.

    PLACEMENT — general across templates/roles, tried in this order (this is the only
    thing that changed; engine/compliance/hash/manifestation are untouched):
      1. EXACT per-signer placeholder: this signer's {{SIGNATURE_<ROLE>}} token. The
         literal token is COVERED and the mark drawn precisely over it.
      2. LABEL-based block line: the role's 'For Site' / 'For Sponsor' Signature line
         (reusing the role tables + pdfplumber); several signers in one block STACK
         downward by `stack_offset`.
      3. LABELLED fallback block ("Electronically signed by {name} — OTP verified —
         {date}") on the last page, logged (template may want standardized anchors).
      * When `role` is None (legacy CTA callers), the original fixed-corner behavior
        is kept unchanged.
    Returns the new PDF path (`_signed`). Compliance is unaffected — the caller hashes
    the produced bytes exactly as before; only (x, y) of the visible mark changes."""
    try:
        from pypdf import PdfWriter, PdfReader  # type: ignore[import]
    except ImportError:
        try:
            from PyPDF2 import PdfWriter, PdfReader  # type: ignore[import]
        except ImportError:
            raise RuntimeError("embed_signature_into_pdf: pypdf not installed")
    try:
        from reportlab.pdfgen import canvas as rl_canvas  # type: ignore[import]
    except ImportError:
        raise RuntimeError("embed_signature_into_pdf: reportlab is not installed")

    reader = PdfReader(pdf_path)
    total_pages = len(reader.pages)
    if total_pages == 0:
        raise ValueError("embed_signature_into_pdf: source PDF has no pages")

    # ---- Decide placement (precedence: placeholder -> label line -> labelled corner).
    ph_anchor = _resolve_placeholder_anchor(pdf_path, role) if role else None
    placeholder_mode = ph_anchor is not None
    anchor = None if placeholder_mode else (_resolve_signature_anchor(pdf_path, role) if role else None)
    line_mode = anchor is not None

    if placeholder_mode:
        target_idx, box_x0, box_x1, box_ybot, box_ytop = ph_anchor
    elif line_mode:
        target_idx, sig_x, sig_y = anchor
    else:
        target_idx = page_num if page_num >= 0 else total_pages - 1
    target_page = reader.pages[target_idx]
    page_width = float(target_page.mediabox.width)
    page_height = float(target_page.mediabox.height)
    if not (placeholder_mode or line_mode):
        sig_x = page_width * 0.55
        sig_y = 80.0
    # Stack within the same block so multiple signers don't overlap. Signature lines
    # often sit low in the block, so keep stacked rows compact and clamp to the page.
    if line_mode and stack_offset:
        sig_y = max(20.0, sig_y - stack_offset * _STACK_ROW)
    if role and not (placeholder_mode or line_mode):
        # FAIL LOUD. Previously we silently drew a corner block here, leaving the
        # {{SIGNATURE_<ROLE>}} placeholder visibly empty while the workflow reported
        # "signed". Refuse to stamp instead, so signing cannot succeed without a real
        # placement. General: applies to ANY role / template / number of signers.
        candidates = _placeholder_tokens_for_role(role)
        looked_for = ", ".join("{{" + t + "}}" for t in candidates) or "(none derivable)"
        raise SignaturePlacementError(
            f"No signature placement found for signer role={role!r}. "
            f"Searched for placeholder tokens [{looked_for}] and a labeled "
            f"'For <party>' signature line, but the document has neither. "
            f"Add a signature placeholder for this signer's role to the template."
        )

    overlay_buf = io.BytesIO()
    c = rl_canvas.Canvas(overlay_buf, pagesize=(page_width, page_height))

    if placeholder_mode:
        # COVER the literal {{SIGNATURE_<ROLE>}} token, then stamp exactly over it so
        # the final doc shows the signature — not the placeholder text.
        pad = 1.5
        c.setFillColorRGB(1, 1, 1)
        c.rect(box_x0 - pad, box_ybot - pad, (box_x1 - box_x0) + 2 * pad,
               (box_ytop - box_ybot) + 2 * pad, stroke=0, fill=1)
        c.setFillColorRGB(0, 0, 0)
        sig_x = box_x0
        box_h = max(box_ytop - box_ybot, 10.0)
        sig_w = min(max(box_x1 - box_x0, 70.0), 130.0)
        sig_h = min(box_h + 4.0, 18.0)
        img_y = box_ybot
    else:
        # Sizing: small mark on a line (so several stack without colliding); larger
        # block for the corner/fallback.
        sig_w, sig_h = (76.0, 11.0) if line_mode else (page_width * 0.35, 60.0)
        img_y = sig_y
    try:
        png_bytes = base64.b64decode(signature_png_base64)
        tmp_png = await asyncio.to_thread(_make_named_tempfile, ".png", png_bytes)
        tmp_png_path = tmp_png.name
        c.drawImage(tmp_png_path, sig_x, img_y, width=sig_w, height=sig_h, mask="auto")
        os.unlink(tmp_png_path)
    except Exception as img_exc:
        logger.warning("embed_signature_into_pdf: could not draw signature image (%s); text placeholder", img_exc)
        c.setFont("Helvetica-Oblique", 12)
        c.drawString(sig_x, img_y + 24, f"[Signature: {signer_name}]")

    if placeholder_mode or line_mode:
        # Compact "name — date" caption beside/under the mark.
        c.setFont("Helvetica", 6)
        caption = signer_name + (f" — {signed_on}" if signed_on else "")
        if placeholder_mode:
            c.drawString(box_x0, max(2.0, box_ybot - 7.0), caption)
        else:
            c.drawString(sig_x + sig_w + 4.0, img_y + 2.0, caption)
    else:
        # Corner/fallback block: a clear "electronically signed" header (only when
        # role-aware, i.e. unified) + the existing Name/Title lines.
        if role:
            c.setFont("Helvetica-Bold", 9)
            hdr = f"Electronically signed by {signer_name} — OTP verified" + (f" — {signed_on}" if signed_on else "")
            c.drawString(sig_x, sig_y + sig_h + 4, hdr)
        c.setFont("Helvetica", 10)
        c.drawString(sig_x, sig_y - 15, f"Name: {signer_name}")
        c.drawString(sig_x, sig_y - 27, f"Title: {signer_title}")

    c.save()
    overlay_buf.seek(0)

    overlay_page = PdfReader(overlay_buf).pages[0]
    target_page.merge_page(overlay_page)

    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    src = Path(pdf_path)
    out_path = src.parent / f"{src.stem}_signed.pdf"
    out_buf = io.BytesIO()
    writer.write(out_buf)
    async with aiofiles.open(out_path, "wb") as fh:
        await fh.write(out_buf.getvalue())

    mode = "placeholder" if placeholder_mode else ("line" if line_mode else "fallback")
    logger.info("embed_signature_into_pdf: stamped %s (role=%s, mode=%s, stack=%s) → %s",
                sfmt(signer_name), sfmt(role), mode, stack_offset, sfmt(out_path))
    return str(out_path)

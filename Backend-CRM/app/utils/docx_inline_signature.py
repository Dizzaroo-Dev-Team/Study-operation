"""
Inject visible electronic signatures into CTA/CDA signature blocks inside DOCX templates.

Fills the sponsor column (e.g. VERILEX BIOSCIENCES CORPORATION) for reviewer/internal
signatures and the site column (e.g. [NAME OF MEMBER INSTITUTION]) for site signatures.
"""
from __future__ import annotations

import io
import re
from typing import Any, Dict, Iterable, List, Optional, Sequence

from docx.document import Document as DocumentType
from docx.shared import Inches
from docx.table import _Cell

from app.config import settings
from app.utils.signature_image import render_signature_png_bytes

# Sponsor / internal reviewer signature anchors
SPONSOR_MARKERS: tuple[str, ...] = (
    "VERILEX BIOSCIENCES CORPORATION",
    "VERILEX BIOSCIENCES",
    "VERILEX",
    "{{SPONSOR_SIGNATURE_BLOCK}}",
    "SPONSOR_SIGNATURE_BLOCK",
)

# Site signature anchors
SITE_MARKERS: tuple[str, ...] = (
    "[NAME OF MEMBER INSTITUTION]",
    "NAME OF MEMBER INSTITUTION",
    "{{SITE_SIGNATURE_BLOCK}}",
    "SITE_SIGNATURE_BLOCK",
)


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").upper()).strip()


def _signer_display_name(signature_data: Dict[str, Any]) -> str:
    if signature_data.get("signer_name"):
        return str(signature_data["signer_name"]).strip()
    email = str(signature_data.get("email") or "").strip()
    if email and "@" in email:
        local = email.split("@", 1)[0]
        return re.sub(r"[._]+", " ", local).strip().title() or email
    return email or "Signer"


def _sign_date_only(signature_data: Dict[str, Any]) -> str:
    raw = str(signature_data.get("timestamp_display") or signature_data.get("timestamp") or "")
    if not raw:
        return ""
    return raw.split()[0] if raw else ""


def _set_paragraph_text(paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _add_signature_image(paragraph, signer_display: str) -> None:
    png_bytes = render_signature_png_bytes(signer_display)
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(2.75))


def _cell_full_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs)


def _paragraph_has_drawing(paragraph) -> bool:
    for run in paragraph.runs:
        if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
            return True
    return False


def _unique_row_cells(row) -> List[_Cell]:
    cells: List[_Cell] = []
    seen_tc: set[int] = set()
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id in seen_tc:
            continue
        seen_tc.add(tc_id)
        cells.append(cell)
    return cells


def _table_is_party_signature_table(table) -> bool:
    """True for 2-column witness/signature tables (sponsor left, site right)."""
    parts: List[str] = []
    for row in table.rows:
        for cell in _unique_row_cells(row):
            parts.append(_normalize(_cell_full_text(cell)))
    combined = " ".join(parts)
    if not combined:
        return False
    has_sponsor = "VERILEX" in combined or "BIOSCIENCES CORPORATION" in combined
    has_site = "MEMBER INSTITUTION" in combined or "IN WITNESS WHEREOF" in combined
    has_fields = any(
        hint in combined for hint in ("BY:", "NAME:", "SIGNATURE", "DATE:", "TITLE:")
    )
    return has_sponsor and has_site and has_fields


def _cell_has_signature_field_labels(cell: _Cell) -> bool:
    text = _normalize(_cell_full_text(cell))
    return any(
        hint in text for hint in ("BY:", "NAME:", "SIGNATURE", "DATE:", "TITLE:")
    )


def _inject_witness_table_columns(
    doc: DocumentType,
    role: str,
    signature_data: Dict[str, Any],
    profile: Any = None,
) -> bool:
    """
    CTA/CDA templates use a 2-column table: left = sponsor, right = site.
    Header and field rows are often split — inject only the target column.
    """
    role_l = (role or "").strip().lower()
    if role_l in {"reviewer", "sponsor", "internal"}:
        col_idx = 0
    elif role_l == "site":
        col_idx = 1
    else:
        return False

    replace_institution = None
    if role_l == "site" and profile:
        replace_institution = (
            getattr(profile, "hospital_name", None)
            or getattr(profile, "site_name", None)
            or None
        )

    changed = False
    for table in doc.tables:
        if not _table_is_party_signature_table(table):
            continue
        seen_target_header = False
        for row in table.rows:
            cells = _unique_row_cells(row)
            if len(cells) < 2:
                continue
            target = cells[col_idx] if col_idx < len(cells) else cells[-1]
            norm = _normalize(_cell_full_text(target))

            if col_idx == 0 and ("MEMBER INSTITUTION" in norm or "PRINCIPAL INVESTIGATOR" in norm):
                continue
            if col_idx == 1 and ("VERILEX" in norm or "BIOSCIENCES CORPORATION" in norm):
                continue
            if col_idx == 1 and "PRINCIPAL INVESTIGATOR" in norm:
                continue

            if col_idx == 0 and ("VERILEX" in norm or "BIOSCIENCES" in norm):
                seen_target_header = True
            if col_idx == 1 and (
                "MEMBER INSTITUTION" in norm
                or (replace_institution and replace_institution.upper() in norm)
            ):
                seen_target_header = True
                for para in target.paragraphs:
                    if "[NAME OF MEMBER INSTITUTION]" in (para.text or ""):
                        _set_paragraph_text(
                            para,
                            (para.text or "").replace(
                                "[NAME OF MEMBER INSTITUTION]",
                                replace_institution or "[NAME OF MEMBER INSTITUTION]",
                            ),
                        )
                        changed = True

            if not (seen_target_header or _cell_has_signature_field_labels(target)):
                continue

            if inject_signature_into_paragraphs(target.paragraphs, signature_data):
                changed = True
    return changed


def _is_section_break(text: str) -> bool:
    t = _normalize(text)
    if not t:
        return False
    if t.startswith("PRINCIPAL INVESTIGATOR"):
        return True
    if len(t) > 60 and "PAGE " in t and " OF " in t:
        return True
    return False


def _paragraph_matches_field(text: str, prefix: str) -> bool:
    return _normalize(text).startswith(prefix)


def _field_value_empty(text: str, prefix: str) -> bool:
    """True when 'Name:' / 'Date:' etc. has no value after the label."""
    raw = (text or "").strip()
    if not raw:
        return True
    parts = raw.split(":", 1)
    if len(parts) < 2:
        return True
    return not parts[1].strip() or parts[1].strip() in {"_", "—", "-"}


def inject_signature_into_paragraphs(
    paragraphs: Sequence,
    signature_data: Dict[str, Any],
    *,
    replace_institution_placeholder: Optional[str] = None,
) -> bool:
    """
    Fill Name / signature image / Title / Date lines in a signature block.
    Returns True if at least one field was updated.
    """
    signer_name = _signer_display_name(signature_data)
    signer_display = str(signature_data.get("email") or signer_name).strip() or signer_name
    sign_date = _sign_date_only(signature_data)
    title = str(signature_data.get("title") or "").strip()
    changed = False

    for i, para in enumerate(paragraphs):
        text = para.text or ""
        upper = _normalize(text)

        if replace_institution_placeholder and "[NAME OF MEMBER INSTITUTION]" in text:
            _set_paragraph_text(
                para,
                text.replace("[NAME OF MEMBER INSTITUTION]", replace_institution_placeholder),
            )
            changed = True
            continue

        if _paragraph_matches_field(text, "NAME") and (
            _field_value_empty(text, "NAME") or upper == "NAME:"
        ):
            _set_paragraph_text(para, f"Name: {signer_name}")
            changed = True
            continue

        if title and _paragraph_matches_field(text, "TITLE") and _field_value_empty(text, "TITLE"):
            _set_paragraph_text(para, f"Title: {title}")
            changed = True
            continue

        if _paragraph_matches_field(text, "DATE") and _field_value_empty(text, "DATE"):
            _set_paragraph_text(para, f"Date: {sign_date}")
            changed = True
            continue

        if "(SIGNATURE)" in upper or upper in {"SIGNATURE:", "SIGNATURE"}:
            if _paragraph_has_drawing(para):
                continue
            _clear_paragraph(para)
            _add_signature_image(para, signer_display)
            changed = True
            continue

        if upper == "BY:" or upper.startswith("BY:"):
            for j in range(i + 1, min(i + 5, len(paragraphs))):
                nxt = paragraphs[j]
                if _paragraph_has_drawing(nxt):
                    changed = True
                    break
                nxt_upper = _normalize(nxt.text)
                if "(SIGNATURE)" in nxt_upper or nxt_upper in {"SIGNATURE:", "SIGNATURE"}:
                    _clear_paragraph(nxt)
                    _add_signature_image(nxt, signer_display)
                    changed = True
                    break
                if not (nxt.text or "").strip():
                    _clear_paragraph(nxt)
                    _add_signature_image(nxt, signer_display)
                    changed = True
                    break
            continue

    return changed


def _markers_for_role(
    role: str,
    profile: Any = None,
) -> List[str]:
    role_l = (role or "").strip().lower()
    markers: List[str] = []

    if role_l in {"reviewer", "sponsor", "internal"}:
        markers.extend(SPONSOR_MARKERS)
        if settings.sponsor_signatory_name:
            markers.append(settings.sponsor_signatory_name.upper())
        entity = getattr(profile, "primary_contracting_entity", None) if profile else None
        if entity:
            markers.append(str(entity).upper())
    elif role_l == "site":
        markers.extend(SITE_MARKERS)
        hospital = getattr(profile, "hospital_name", None) if profile else None
        if hospital:
            markers.append(str(hospital).upper())
        site_name = getattr(profile, "site_name", None) if profile else None
        if site_name:
            markers.append(str(site_name).upper())
    else:
        markers.extend(SPONSOR_MARKERS)
        markers.extend(SITE_MARKERS)

    return markers


def _container_matches_markers(text: str, markers: Iterable[str]) -> bool:
    upper = _normalize(text)
    for marker in markers:
        m = _normalize(marker)
        if m and m in upper:
            return True
    return False


def _inject_into_cell(
    cell: _Cell,
    signature_data: Dict[str, Any],
    markers: List[str],
    *,
    replace_institution: Optional[str] = None,
) -> bool:
    cell_text = _cell_full_text(cell)
    if not _container_matches_markers(cell_text, markers):
        return False
    return inject_signature_into_paragraphs(
        cell.paragraphs,
        signature_data,
        replace_institution_placeholder=replace_institution,
    )


def _inject_into_body_paragraphs(
    doc: DocumentType,
    signature_data: Dict[str, Any],
    markers: List[str],
    *,
    replace_institution: Optional[str] = None,
) -> bool:
    changed = False
    paras = doc.paragraphs
    for idx, para in enumerate(paras):
        if not _container_matches_markers(para.text, markers):
            continue
        block = [para]
        for j in range(idx + 1, min(idx + 14, len(paras))):
            nxt = paras[j]
            if _is_section_break(nxt.text):
                break
            block.append(nxt)
        if inject_signature_into_paragraphs(
            block,
            signature_data,
            replace_institution_placeholder=replace_institution,
        ):
            changed = True
    return changed


def inject_inline_signature_into_docx(
    doc: DocumentType,
    role: str,
    signature_data: Dict[str, Any],
    profile: Any = None,
) -> bool:
    """
    Write signature image + Name/Title/Date into the matching signature block.
    Returns True when at least one block was updated.
    """
    role_l = (role or "").strip().lower()
    markers = _markers_for_role(role_l, profile)
    replace_institution = None
    signer_name = _signer_display_name(signature_data)
    signer_title = str(signature_data.get("title") or "").strip()

    if role_l == "site" and profile:
        replace_institution = (
            getattr(profile, "hospital_name", None)
            or getattr(profile, "site_name", None)
            or None
        )
        if getattr(profile, "authorized_signatory_name", None):
            signature_data = {
                **signature_data,
                "signer_name": profile.authorized_signatory_name,
            }
        if getattr(profile, "authorized_signatory_title", None) and not signer_title:
            signature_data = {
                **signature_data,
                "title": profile.authorized_signatory_title,
            }
    elif role_l in {"reviewer", "sponsor", "internal"}:
        if settings.sponsor_signatory_name:
            signature_data = {
                **signature_data,
                "signer_name": settings.sponsor_signatory_name,
            }

    changed = False
    witness_changed = _inject_witness_table_columns(
        doc, role_l, signature_data, profile=profile
    )
    if witness_changed:
        changed = True
    else:
        for table in doc.tables:
            if _table_is_party_signature_table(table):
                continue
            for row in table.rows:
                for cell in _unique_row_cells(row):
                    if _inject_into_cell(
                        cell,
                        signature_data,
                        markers,
                        replace_institution=replace_institution,
                    ):
                        changed = True

        if _inject_into_body_paragraphs(
            doc,
            signature_data,
            markers,
            replace_institution=replace_institution,
        ):
            changed = True

    # Explicit placeholder blocks from template library
    if role_l in {"reviewer", "sponsor", "internal"}:
        for para in doc.paragraphs:
            if "{{SPONSOR_SIGNATURE_BLOCK}}" in para.text:
                from app.utils.docx_placeholder_replace import (
                    _generate_sponsor_signature_block_text,
                    _replace_multiline_placeholder_in_paragraph,
                )

                block_text = _generate_sponsor_signature_block_text(
                    profile,
                    settings.sponsor_signatory_name,
                    settings.sponsor_signatory_email or signature_data.get("email"),
                    signature_data.get("email"),
                )
                block_text += f"\n\nSigned electronically: {signer_name}\nDate: {_sign_date_only(signature_data)}"
                if _replace_multiline_placeholder_in_paragraph(
                    para, "{{SPONSOR_SIGNATURE_BLOCK}}", block_text
                ):
                    changed = True

    if role_l == "site" and profile:
        for para in doc.paragraphs:
            if "{{SITE_SIGNATURE_BLOCK}}" in para.text:
                from app.utils.docx_placeholder_replace import (
                    _generate_site_signature_block_text,
                    _replace_multiline_placeholder_in_paragraph,
                )

                block_text = _generate_site_signature_block_text(profile)
                block_text += f"\n\nSigned electronically: {_signer_display_name(signature_data)}\nDate: {_sign_date_only(signature_data)}"
                if _replace_multiline_placeholder_in_paragraph(
                    para, "{{SITE_SIGNATURE_BLOCK}}", block_text
                ):
                    changed = True

    return changed
